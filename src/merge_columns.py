#!/usr/bin/env python3
"""Merge a second engine's translation column into a primary docx as an extra column.

This is the reusable core of the ``google+deepl`` combined engine (see
``combine_google_deepl.py``): translate the source once with Google (primary) and once with
DeepL (secondary) — two identical-geometry 3-column tables — then append DeepL's translation
column (col 2) onto the Google docx as a new 4th column, so the reader sees both engines side
by side:

    col 0 (index) . col 1 (source) . col 2 (Google) . col 3 (DeepL, new)

Pure ``python-docx`` / lxml — no Selenium, no globals, fully unit-testable. The new column's
cells are a deepcopy of the secondary's ``<w:tc>`` elements, so RTL direction, the run style,
and the destination font carry over verbatim.

Row geometry is engine-independent (it derives from the source phrase grouping), so the two
passes produce identical row counts; the merge hard-asserts that and raises ``ValueError`` on
any drift or on an irregular (merged-cell) grid, so the caller can degrade to Google-only
rather than emit scrambled side-by-side cells.

Run standalone:
    python merge_columns.py google_out.docx deepl_out.docx combined.docx [DeepL]
"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_TC = qn("w:tc")
_GRIDCOL = qn("w:gridCol")
_TBLGRID = qn("w:tblGrid")


def _row_tc_count(tr):
    """Number of direct ``<w:tc>`` children of a ``<w:tr>`` (cells before span expansion)."""
    return len(tr.findall(_TC))


def _pick_translation_table(doc):
    """Pick the table the pipeline wrote the translation into: the LARGEST table with at least
    3 columns. A docx that also carries a small header/metadata table is handled correctly,
    instead of the naive ``tables[0]`` (which could be that header table)."""
    best = None
    for t in doc.tables:
        ncols = len(t.rows[0].cells) if t.rows else 0
        if ncols >= 3 and (best is None or len(t.rows) > len(best.rows)):
            best = t
    if best is not None:
        return best
    return doc.tables[0] if doc.tables else None


def merge_second_engine_column(
    primary_path,
    secondary_path,
    out_path,
    primary_label="Google",
    secondary_label="DeepL",
    landscape=True,
):
    """Append the secondary docx's translation column (col 2) onto the primary as a new column.

    Both docx must be regular >=3-column tables with identical row counts. Returns ``out_path``.
    Raises ``ValueError`` on row-count drift or an irregular grid (merged cells) -- the caller is
    expected to fall back to serving the primary (Google) docx alone.

    ``primary_label`` / ``secondary_label`` -- each engine's name, written in the cell DIRECTLY
    BELOW the destination-language label. The pipeline writes the language name into cell
    (1, 2); the cell directly below it (row 2) is a blank separator, so labelling there is safe
    and never overwrites a subtitle line (only an empty cell is filled). Set either to None to
    skip that label.

    ``landscape`` -- rotate the page to landscape so the four side-by-side columns are
    comfortably visible. Pass False to keep the source's portrait orientation
    (combine_google_deepl.py exposes this as ``--no-landscape``).
    """
    primary = Document(primary_path)
    secondary = Document(secondary_path)

    pt = _pick_translation_table(primary)
    st = _pick_translation_table(secondary)
    if pt is None or st is None:
        raise ValueError("merge_second_engine_column: a docx has no table")

    p_rows = pt.rows
    s_rows = st.rows
    if len(p_rows) != len(s_rows):
        raise ValueError(
            "row-count drift: primary=%d secondary=%d -- refusing to merge mismatched columns"
            % (len(p_rows), len(s_rows))
        )

    # Regular-grid guard: every row of BOTH tables must currently carry the same number of
    # cells. gridSpan / vMerge (merged cells) make a ragged grid that a naive append would
    # corrupt -- and a ragged secondary would also make the "col 2" pick ambiguous.
    p_counts = {_row_tc_count(r._tr) for r in p_rows}
    if len(p_counts) != 1:
        raise ValueError("irregular primary grid (cells/row = %s) -- merged cells unsupported" % sorted(p_counts))
    s_counts = {_row_tc_count(r._tr) for r in s_rows}
    if len(s_counts) != 1:
        raise ValueError("irregular secondary grid (cells/row = %s) -- merged cells unsupported" % sorted(s_counts))

    # 1) widen the table grid by one column (copy the last <w:gridCol> to keep a sensible width).
    grid = pt._tbl.find(_TBLGRID)
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        pt._tbl.insert(0, grid)
    existing = grid.findall(_GRIDCOL)
    if existing:
        grid.append(deepcopy(existing[-1]))
    else:
        grid.append(OxmlElement("w:gridCol"))

    # 2) per row, deepcopy the secondary's translation cell (<w:tc> col 2) onto the primary row.
    for ri in range(len(p_rows)):
        # Use the LITERAL <w:tc> children (not python-docx's span-expanded .cells, which can
        # repeat a spanned cell) so col 2 is unambiguously the secondary's translation cell.
        s_tcs = s_rows[ri]._tr.findall(_TC)
        src_tc = s_tcs[2] if len(s_tcs) >= 3 else s_tcs[-1]
        new_tc = deepcopy(src_tc)
        p_rows[ri]._tr.append(new_tc)

    # 3) label each engine in the cell directly BELOW the language name (row 2): the primary
    #    (Google) under col 2, the secondary (DeepL) under the new last col. Only an EMPTY cell
    #    is written, so a subtitle line is never overwritten.
    _label_engines_below_language(pt, primary_label, secondary_label)

    # 4) landscape so the four side-by-side columns fit (caller-toggleable).
    if landscape:
        _set_landscape(primary)

    out = Path(out_path)
    primary.save(str(out))
    return out


def _label_engines_below_language(table, primary_label, secondary_label):
    """Write each engine's name in the cell directly below the destination-language label.

    The language name sits in cell (1, 2); row 2 is the blank separator beneath it. Only an
    EMPTY target cell is written, so a real subtitle line is never overwritten."""
    rows = table.rows
    if len(rows) < 3:
        return
    tcs = rows[2]._tr.findall(_TC)
    targets = []
    if primary_label and len(tcs) >= 3:
        targets.append((tcs[2], primary_label))      # below col 2 (primary / Google)
    if secondary_label and len(tcs) >= 1:
        targets.append((tcs[-1], secondary_label))   # below the new last col (secondary / DeepL)
    for tc, label in targets:
        existing = "".join((t.text or "") for t in tc.iter(qn("w:t")))
        if existing.strip() == "":
            _set_cell_text(tc, label)


def _set_landscape(doc):
    """Rotate every section to landscape (swap page width/height). Idempotent: a section that
    is already landscape (width >= height) keeps its dimensions."""
    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        w, h = section.page_width, section.page_height
        if w is not None and h is not None and w < h:
            section.page_width, section.page_height = h, w
        section.orientation = WD_ORIENT.LANDSCAPE


def _set_cell_text(tc, text):
    """Replace all text in a ``<w:tc>`` with a single run carrying ``text`` (keeps the cell's
    first paragraph + its properties, drops extra paragraphs/runs)."""
    paras = tc.findall(qn("w:p"))
    for extra in paras[1:]:
        tc.remove(extra)
    p = paras[0] if paras else OxmlElement("w:p")
    if not paras:
        tc.append(p)
    for r in p.findall(qn("w:r")):
        p.remove(r)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    p.append(run)


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("usage: python merge_columns.py <primary.docx> <secondary.docx> <out.docx> [label]")
        raise SystemExit(2)
    label = sys.argv[4] if len(sys.argv) > 4 else "DeepL"
    result = merge_second_engine_column(sys.argv[1], sys.argv[2], sys.argv[3], secondary_label=label)
    print("merged -> %s" % result)
