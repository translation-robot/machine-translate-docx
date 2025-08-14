#!/usr/bin/python3
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
import platform

# - *- coding: utf- 8 - *-
PROGRAM_VERSION="2025-08-14"
json_configuration_url='https://raw.githubusercontent.com/translation-robot/machine-translate-docx/main/src/configuration/configuration.json'
# Day 0 is October 3rd 2017


print("*********************************************************")
print("*  machine-translate-docx program version : %s" % (PROGRAM_VERSION))
print("*********************************************************")

print("Python programming language %s\n" % (platform.python_version()))

import gc
import pprint
from pprint import pprint
import traceback
import shlex
import subprocess
import os
#from googletrans import Translator
import re
import time
import codecs
import urllib
import urllib.request
import requests
import json


from inspect import currentframe, getframeinfo
import chardet
import getpass
import datetime

import zipfile
import xml.dom.minidom
# used to get elements in XML, shading in docx for example
from lxml import etree

# This library automatically downloads chrome driver
# pyderman was replaced with webdriver_manager
# then selenium 4.11.2 managed downloading the drivers
#import pyderman
# For selenium 3

#from selenium import webdriver
import undetected_chromedriver as uc

from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.remote_connection import LOGGER
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.keys import Keys 

from screeninfo import get_monitors


#from selenium.webdriver import Firefox, FirefoxOptions
from time import sleep
import argparse
import clipboard
#import pyperclip

import psutil

#import winsound

import docx
from docx import Document
from docx import oxml
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

# For japanese
#import tinysegmenter

#from tinysegmenter import TinySegmenter
# For Thai
#from thai_tokenizer import Tokenizer as thai_tokenizer_tokenizer
#from thai_tokenizer import data
#from thai_tokenizer.data import bpe_merges
#import thai_tokenizer 
#import thai-segmenter

import timeit
import datetime
import progressbar


from timeit import default_timer as timer

import re
import inspect

from xlsx_translation_memory import xlsx_translation_memory

import html

from urllib.parse import urlencode, quote_plus

from openpyxl import load_workbook
from openpyxl import Workbook

from bs4 import BeautifulSoup

# pip install pycryptodome
# used for passwords (deepl, etc)

# Load configuration from a json file on internet (github for example)

import json

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import glob

from langcodes import *

import math

import shutil


#from parsivar import Normalizer
my_hazm_normalizer = None

#from hazm import Normalizer
#import hazm

# Get key value from an array of json strings, ['deepl','account','email'] for example
# The first json object containing the key value is returned or default_when_none value.
def validate_json_string(json_string):
    try:
        if type(json_string) is str:
            json_string
        elif type(json_string) is bytes:
            pass # OK
        else:
            return False
        json_obj = json.loads(json_string)
        if json_obj is None:
            return False
        return True
        
    except:
        var = traceback.format_exc()
        print(var)
        return False

# Get key value from an array of json strings, ['deepl','account','email'] key for example
# The first json object containing the key value is returned or default_when_none value.
def get_nested_value_from_json_array(json_array, keys, default_when_none = None):
    try:
        for json_str in json_array:
            try:
                json_obj = json.loads(json_str)
                value = None
                
                for key in keys:
                    if key in json_obj:
                        json_obj = json_obj[key]
                    else:
                        json_obj = None
                if json_obj is not None:
                    return json_obj
            except:
                pass
        
        return default_when_none
        
    except json.JSONDecodeError:
        print("Invalid JSON input")
        return default_when_none

# Remove the _ a the end of the email and password key name
DefaultJsonConfiguration = """{
    "local_configuration":{
        "json_filename_path": "configuration.json"
    },
	"deepl": {
		"account": {
			"email_": "********@gmail.com",
			"password_": "********",
			"type": "free",
			"maximum_character_block": 1500
		},
		"no_account": {
			"maximum_character_block": 1500
		},
        "maximum_clear_cache_retry" : 20
	},
	"google": {
		"javascript_translation": {},
		"html_translation_form": {
			"maximum_character_block": 5000
		}
	},
	"perplexity": {
		"account": {
			"maximum_character_block": 2000
		}
	},
	"statistics": {
		"html_statistics_form_url": "https://contactdirectavecdieu.net/robot-stats.php",
		"google_sheets_statistics": {
			"account": {
				"google_account": "to be determined"
			}
		}
	},
	"version_checker": {
	  "javascript_json_version_checker_url": "https://translation-robot.github.io/machine-translate-docx/src/robot_js_query.html",
      "sleep_seconds_on_update": 30
	},
	"location": {
	  "primary_country_checker_url": "http://ip-api.com/json/",
	  "secondary_country_checker_url": "https://www.contactdirectavecdieu.net/geoip/index.php",
      "http_query_timeout" : 3
	},
	"chrome_driver": {
      "restricted_countries": ["North Korea", "Iran", "Syria", "Sudan", "Cuba", "Crimea"],
      "mirror_url": "https://www.contactdirectavecdieu.net/known-good-versions-with-downloads.php"
    }
}"""

def test_internet(host="8.8.8.8", port=53, timeout=3):
    """
    Host: 8.8.8.8 (google-public-dns-a.google.com)
    OpenPort: 53/tcp
    Service: domain (DNS/TCP)
    """
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error as ex:
        print(ex)
        return False

try:
    json_online_configuration = requests.get(json_configuration_url).content
except:
    print("Warning, unable to get configuration from internet at {json_configuration_url}")
    if not test_internet():
        print("Warning, internet connection seems to be down, google name servers don't respond")
        time.sleep(1)
        
    json_online_configuration = "{}"


# Find default configuration file name from other configuration files
local_configuration_json_path_key = ["local_configuration", "json_filename_path"]
local_configuration_json_path = get_nested_value_from_json_array([json_online_configuration,DefaultJsonConfiguration], local_configuration_json_path_key)

# determine if application is a script file or frozen exe
if getattr(sys, 'frozen', False):
    application_path = os.path.dirname(sys.executable)
elif __file__:
    application_path = os.path.dirname(__file__)

configuration_file_full_path = os.path.join(application_path, local_configuration_json_path)

try:
    if os.path.isfile(configuration_file_full_path):
        with open(configuration_file_full_path) as configuration_file:
          local_json_contents = configuration_file.read()
    else:
        #print(f"Optional local json configuration file not found at {configuration_file_full_path}, ignoring")
        local_json_contents = None
except:
    local_json_contents = None
          
json_configuration_array = [local_json_contents,json_online_configuration,DefaultJsonConfiguration]

# Find default sleep time for update message
version_checker_sleep_seconds_on_update_key = ["version_checker", "sleep_seconds_on_update"]
version_checker_sleep_seconds_on_update = get_nested_value_from_json_array(json_configuration_array,
    version_checker_sleep_seconds_on_update_key, default_when_none=30)
# We assume the program does not need update. Value of 1 is for update needed.
str_needs_update = "0"



process_platform = platform.system()
if platform.system() == 'Windows':
    from win32com.client import Dispatch

tried_login_in_deepl = False
logged_into_deepl = False

from_text_table = [''] *1
from_text_is_greyed_table = [0] *1
from_text_is_red_color_table = [0] *1
from_text_is_end_of_line_table = [0] *1
from_text_is_beginning_of_line_table = [0] *1
from_text_is_empty_line_table = [0] *1
from_text_is_conditional_end_of_line_table = [0] *1
from_text_by_phrase_separator_table = [''] *1
from_text_by_phrase_table = [''] *1
#number of lines in per phrase
from_text_nb_lines_in_phrase = [0] *1
from_text_nb_lines_in_cell = [0] *1
#
to_text_by_phrase_separator_table = [''] *1
to_text_by_phrase_separator_removed_table = [''] *1
to_text_splited_table1 = [''] *1
to_text_by_phrase_table = [''] *1
to_text_table = [''] *1
to_raw_translated_table = [''] *1
to_text_removed_line_separator = [''] *1
translation_result_using_separator = [''] *1
translation_result_phrase_array = [[]] *1
translation_result = [''] *1
from_text_is_read = [0] *1
word_translation_table_length = 0
table = None

table_cells = [['' for i in range(1)] for j in range(1)]

docxfile_table_number_of_phrases = 0

selenium_chrome_machine_translate_once = None

numerrors_deepl = 0

# We have found zero phrase up to now
docxfile_table_number_of_characters = 0
docxfile_table_number_of_phrases = 0
docxfile_table_number_of_words = 0

numrows = 0
numcols = 0

E_mail_str = 'sm' + 'tv' + '.' + 'bot' + '@g' + 'mail' + '.' + 'c' + 'o' + 'm'

#import pandas as pd
#import multiprocessing

cf = currentframe()
filename = getframeinfo(cf).filename

start_time = datetime.datetime.now()

for m in get_monitors():
    #print(str(m))
    break

parser = argparse.ArgumentParser()

#parser = argparse.ArgumentParser(description = "Translate everything!")
#parser.add_argument('--source-language', required = True, choices = Languages, help="Specify the source language!")
parser.add_argument('--srclang', '-sl', required = False, help="Specify the default source language, en is default (hi,ja,ru,de,ru,hi,ja,in, etc)", default='en')
parser.add_argument('--destlang', '--dl', required = False, help="Specify the destination language with 2 letter code (hi,ja,ru,de,ru,hi,ja,in, etc)")
parser.add_argument('--engine', '-e', required = False, help="Specify the translation engine (google, deepl, yandex, chatgpt, perplexity)")
parser.add_argument('--enginemethod', '-m', required = False, help="Specify the method (javascript, phrasesblock, singlephrase, xlsxfile, textfile )")
parser.add_argument('--docxfile', '-d', required = False, help="Input file name")
parser.add_argument('--xlsxreplacefile', '-x', required = False, help="Excel xlsx search and replace file")
parser.add_argument('--destfont', '-f', required = False, help="Destination font name")
parser.add_argument('--useapi', '-a', required = False, help="Use api to get translation, lower quality but faster", action='store_true')
parser.add_argument('--split', '-s', required = False, help="Split web translation into cells", action='store_true')
parser.add_argument('--splitonly', required = False, help="Split translation into lines only, do not translate.", action='store_true')
parser.add_argument('--showbrowser', '-b', required = False, help="Show browser", action='store_true')
parser.add_argument('--exitonsuccess', '-t', required = False, help="Exit progream on success", action='store_true')
parser.add_argument('--viewdocx', '-l', required = False, help="Open the docx file with the default application after completion.", action='store_true')
parser.add_argument('--silent', '-q', required = False, help="Silent, do not ask question and exit silently", action='store_true')
parser.add_argument("--verbose", '-v', help="increase output verbosity", action="store_true")
#parser.add_argument('--destination-file', required = True, help="Output file name")
#args = parser.parse_args()
parser.add_argument('--version', required = False, help="Show program version", action='store_true')

try:
    args = parser.parse_args()
except:
    #print("Waiting for the input_element...")
    var = traceback.format_exc()
    print(var)
    #input ("Type enter to continue")

show_version = args.version
silent = args.silent
if show_version:

    print("\nDeveloper: %s\n" %(E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    sys.exit(0)

if args.docxfile is None:
    parser.print_help()
    print("\nDeveloper: smtv.bot@gmail.com\n")
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    else:
        print("Program ended with errors")
    sys.exit(1)

use_html = False

google_translate_lang_codes = {
    'af': 'Afrikaans',
    'sq': 'Albanian',
    'am': 'Amharic',
    'ar': 'Arabic',
    'hy': 'Armenian',
    'az': 'Azerbaijani',
    'eu': 'Basque',
    'be': 'Belarusian',
    'bn': 'Bengali',
    'bs': 'Bosnian',
    'bg': 'Bulgarian',
    'ca': 'Catalan',
    'ceb': 'Cebuano',
    'zh': 'Chinese (Simplified)',
    'zh-CN': 'Chinese (Simplified)',
    'zh-TW': 'Chinese (Traditional)',
    'co': 'Corsican',
    'hr': 'Croatian',
    'cs': 'Czech',
    'da': 'Danish',
    'nl': 'Dutch',
    'en': 'English',
    'eo': 'Esperanto',
    'et': 'Estonian',
    'fi': 'Finnish',
    'fr': 'French',
    'fy': 'Frisian',
    'gl': 'Galician',
    'ka': 'Georgian',
    'de': 'German',
    'el': 'Greek',
    'gu': 'Gujarati',
    'ht': 'Haitian Creole',
    'ha': 'Hausa',
    'haw': 'Hawaiian',
    'iw': 'Hebrew',
    'hi': 'Hindi',
    'hmn': 'Hmong',
    'hu': 'Hungarian',
    'is': 'Icelandic',
    'ig': 'Igbo',
    'id': 'Indonesian',
    'ga': 'Irish',
    'it': 'Italian',
    'ja': 'Japanese',
    'jv': 'Javanese',
    'kn': 'Kannada',
    'kk': 'Kazakh',
    'km': 'Khmer',
    'ko': 'Korean',
    'ku': 'Kurdish',
    'ky': 'Kyrgyz',
    'lo': 'Lao',
    'la': 'Latin',
    'lv': 'Latvian',
    'lt': 'Lithuanian',
    'lb': 'Luxembourgish',
    'mk': 'Macedonian',
    'mg': 'Malagasy',
    'ms': 'Malay',
    'ml': 'Malayalam',
    'mt': 'Maltese',
    'mi': 'Maori',
    'mr': 'Marathi',
    'mn': 'Mongolian',
    'my': 'Myanmar (Burmese)',
    'ne': 'Nepali',
    'no': 'Norwegian',
    'ny': 'Nyanja (Chichewa)',
    'ps': 'Pashto',
    'fa': 'Persian',
    'pl': 'Polish',
    'pt': 'Portuguese (Portugal, Brazil)',
    'pa': 'Punjabi',
    'ro': 'Romanian',
    'ru': 'Russian',
    'sm': 'Samoan',
    'gd': 'Scots Gaelic',
    'sr': 'Serbian',
    'st': 'Sesotho',
    'sn': 'Shona',
    'sd': 'Sindhi',
    'si': 'Sinhala (Sinhalese)',
    'sk': 'Slovak',
    'sl': 'Slovenian',
    'so': 'Somali',
    'es': 'Spanish',
    'su': 'Sundanese',
    'sw': 'Swahili',
    'sv': 'Swedish',
    'tl': 'Tagalog (Filipino)',
    'tg': 'Tajik',
    'ta': 'Tamil',
    'te': 'Telugu',
    'th': 'Thai',
    'tr': 'Turkish',
    'uk': 'Ukrainian',
    'ur': 'Urdu',
    'uz': 'Uzbek',
    'vi': 'Vietnamese',
    'cy': 'Welsh',
    'xh': 'Xhosa',
    'yi': 'Yiddish',
    'yo': 'Yoruba',
    'zu': 'Zulu'
    }

deepl_translate_lang_codes = {
    'ar': 'Arabic',
    'bg': 'Bulgarian',
    'cs': 'Czech',
    'da': 'Danish',
    'de': 'German',
    'el': 'Greek',
    'en': 'English',
    'en-us': 'English (American)',
    'en-gb': 'English (British)',
    'es': 'Spanish',
    'et': 'Estonian',
    'fi': 'Finnish',
    'fr': 'French',
    'he': 'Hebrew',
    'hu': 'Hungarian',
    'id': 'Indonesian',
    'it': 'Italian',
    'ja': 'Japanese',
    'ko': 'Korean',
    'lt': 'Lithuanian',
    'lv': 'Latvian',
    'nb': 'Norwegian',
    'nl': 'Dutch',
    'pl': 'Polish',
    'pt': 'Portuguese',
    'pt-br': 'Portuguese (Brazilian)',
    'pt-pt': 'Portuguese (all Portuguese variants excluding Brazilian Portuguese)',
    'ro': 'Romanian',
    'ru': 'Russian',
    'sk': 'Slovak',
    'sl': 'Slovenian',
    'sv': 'Swedish',
    'tr': 'Turkish',
    'uk': 'Ukrainian',
    'vi': 'Vietnamese',
    'zh': 'Chinese (Simplified)',
    'zh-hans': 'Chinese (Simplified)',
    'zh-hant': 'Chinese (Traditional)',
}

# This is to set docx document language for spelling
office_language_tags = {
    'ar': 'ar-SA',
    'bg': 'bg-BG',
    'zh': 'zh-CN',
    'zh': 'zh-TW',
    'hr': 'hr-HR',
    'cs': 'cs-CZ',
    'da': 'da-DK',
    'nl': 'nl-NL',
    'en': 'en-US',
    'et': 'et-EE',
    'fi': 'fi-FI',
    'fr': 'fr-FR',
    'de': 'de-DE',
    'el': 'el-GR',
    'he': 'he-IL',
    'hi': 'hi-IN',
    'hu': 'hu-HU',
    'id': 'id-ID',
    'it': 'it-IT',
    'ja': 'ja-JP',
    'kk': 'kk-KZ',
    'ko': 'ko-KR',
    'lv': 'lv-LV',
    'lt': 'lt-LT',
    'ms': 'ms-MY',
    'nb': 'nb-NO',
    'pl': 'pl-PL',
    'pt': 'pt-BR',
    'pt': 'pt-PT',
    'ro': 'ro-RO',
    'ru': 'ru-RU',
    'sr': 'sr-latn-RS',
    'sk': 'sk-SK',
    'sl': 'sl-SI',
    'es': 'es-ES',
    'sv': 'sv-SE',
    'th': 'th-TH',
    'tr': 'tr-TR',
    'uk': 'uk-UA',
    'vi': 'vi-VN'
}

right_to_left_languages_list = {
    'am': 'Amharic',
    'ar': 'Arabic',
    'az': 'Azerbaijani',
    'iw': 'Hebrew',
    'ku': 'Kurdish',
    'fa': 'Persian',
    'ur': 'Urdu'
    }

script_folder = os.path.dirname(os.path.realpath(__file__))
os_path = os.environ["PATH"]
new_os_path = "%s;%s" %(script_folder, os_path)

#print("\nAdding %s to path.\n" % (script_folder))
os.environ["PATH"] = new_os_path

use_translation_api = False


#line_separator_str = ' () '
line_separator_str = ' '
#line_separator_nospace_str = '()'
line_separator_nospace_str = '()'
line_separator_regex_str = ' ?\(\) ?'

#pp = pprint.PrettyPrinter(indent=4)

# ... = \u2026 Horizontal ellipsis
# ”  = \u201D Right double quotation mark
eol_array = ['\. {0,}$', '\! {0,}$', '\? {0,}$',  '[\.\!\?\'] ?["”\'\)] {0,}$', u'\u2026 {0,}$',
    '। {0,}$', # Hindi period
    '。 {0,}$', '？ {0,}$', '！ {0,}$', # Chinese and Japanese period
    '։ {0,}$', #ARMENIAN FULL STOP	U+0589	ARMENIAN FULL STOP	
    '։ {0,}$', #full stop, georgian	U+0589	ARMENIAN FULL STOP	
    '։ {0,}$', #FULL STOP, ARMENIAN	U+0589	ARMENIAN FULL STOP	
    '։ {0,}$', #georgian full stop	U+0589	ARMENIAN FULL STOP	
    '۔ {0,}$', #ARABIC FULL STOP	U+06D4	ARABIC FULL STOP	
    '۔ {0,}$', #FULL STOP, ARABIC	U+06D4	ARABIC FULL STOP	
    '܁ {0,}$', #SYRIAC SUPRALINEAR FULL STOP	U+0701	SYRIAC SUPRALINEAR FULL STOP	
    '܂ {0,}$', #SYRIAC SUBLINEAR FULL STOP	U+0702	SYRIAC SUBLINEAR FULL STOP	
    '። {0,}$', #ETHIOPIC FULL STOP	U+1362	ETHIOPIC FULL STOP	
    '። {0,}$', #FULL STOP, ETHIOPIC	U+1362	ETHIOPIC FULL STOP	
    '᙮ {0,}$', #CANADIAN SYLLABICS FULL STOP	U+166E	CANADIAN SYLLABICS FULL STOP	
    '᙮ {0,}$', #FULL STOP, CANADIAN SYLLABICS	U+166E	CANADIAN SYLLABICS FULL STOP	
    '᙮ {0,}$', #SYLLABICS FULL STOP, CANADIAN	U+166E	CANADIAN SYLLABICS FULL STOP	
    '᠃ {0,}$', #MONGOLIAN FULL STOP	U+1803	MONGOLIAN FULL STOP	
    '᠃ {0,}$', #FULL STOP, MONGOLIAN	U+1803	MONGOLIAN FULL STOP	
    '᠉ {0,}$', #MONGOLIAN MANCHU FULL STOP	U+1809	MONGOLIAN MANCHU FULL STOP	
    '᠉ {0,}$', #FULL STOP, MONGOLIAN MANCHU	U+1809	MONGOLIAN MANCHU FULL STOP	
    '᠉ {0,}$', #MANCHU FULL STOP, MONGOLIAN	U+1809	MONGOLIAN MANCHU FULL STOP	
    '⒈ {0,}$', #DIGIT ONE FULL STOP	U+2488	DIGIT ONE FULL STOP	
    '⒉ {0,}$', #DIGIT TWO FULL STOP	U+2489	DIGIT TWO FULL STOP	
    '⒊ {0,}$', #DIGIT THREE FULL STOP	U+248A	DIGIT THREE FULL STOP	
    '⒋ {0,}$', #DIGIT FOUR FULL STOP	U+248B	DIGIT FOUR FULL STOP	
    '⒌ {0,}$', #DIGIT FIVE FULL STOP	U+248C	DIGIT FIVE FULL STOP	
    '⒍ {0,}$', #DIGIT SIX FULL STOP	U+248D	DIGIT SIX FULL STOP	
    '⒎ {0,}$', #DIGIT SEVEN FULL STOP	U+248E	DIGIT SEVEN FULL STOP	
    '⒏ {0,}$', #DIGIT EIGHT FULL STOP	U+248F	DIGIT EIGHT FULL STOP	
    '⒐ {0,}$', #DIGIT NINE FULL STOP	U+2490	DIGIT NINE FULL STOP	
    '⒑ {0,}$', #NUMBER TEN FULL STOP	U+2491	NUMBER TEN FULL STOP	
    '⒒ {0,}$', #NUMBER ELEVEN FULL STOP	U+2492	NUMBER ELEVEN FULL STOP	
    '⒓ {0,}$', #NUMBER TWELVE FULL STOP	U+2493	NUMBER TWELVE FULL STOP	
    '⒔ {0,}$', #NUMBER THIRTEEN FULL STOP	U+2494	NUMBER THIRTEEN FULL STOP	
    '⒕ {0,}$', #NUMBER FOURTEEN FULL STOP	U+2495	NUMBER FOURTEEN FULL STOP	
    '⒖ {0,}$', #NUMBER FIFTEEN FULL STOP	U+2496	NUMBER FIFTEEN FULL STOP	
    '⒗ {0,}$', #NUMBER SIXTEEN FULL STOP	U+2497	NUMBER SIXTEEN FULL STOP	
    '⒘ {0,}$', #NUMBER SEVENTEEN FULL STOP	U+2498	NUMBER SEVENTEEN FULL STOP	
    '⒙ {0,}$', #NUMBER EIGHTEEN FULL STOP	U+2499	NUMBER EIGHTEEN FULL STOP	
    '⒚ {0,}$', #NUMBER NINETEEN FULL STOP	U+249A	NUMBER NINETEEN FULL STOP	
    '⒛ {0,}$', #NUMBER TWENTY FULL STOP	U+249B	NUMBER TWENTY FULL STOP	
    '⳹ {0,}$', #COPTIC OLD NUBIAN FULL STOP	U+2CF9	COPTIC OLD NUBIAN FULL STOP	
    '⳾ {0,}$', #COPTIC FULL STOP	U+2CFE	COPTIC FULL STOP	
    '⸼ {0,}$', #STENOGRAPHIC FULL STOP	U+2E3C	STENOGRAPHIC FULL STOP	
    '。 {0,}$', #IDEOGRAPHIC FULL STOP	U+3002	IDEOGRAPHIC FULL STOP	
    '。 {0,}$', #FULL STOP, IDEOGRAPHIC	U+3002	IDEOGRAPHIC FULL STOP	
    '꓿ {0,}$', #LISU PUNCTUATION FULL STOP	U+A4FF	LISU PUNCTUATION FULL STOP	
    '꘎ {0,}$', #VAI FULL STOP	U+A60E	VAI FULL STOP	
    '꛳ {0,}$', #BAMUM FULL STOP	U+A6F3	BAMUM FULL STOP	
    '︒ {0,}$', #PRESENTATION FORM FOR VERTICAL IDEOGRAPHIC FULL STOP	U+FE12	PRESENTATION FORM FOR VERTICAL IDEOGRAPHIC FULL STOP	
    '﹒ {0,}$', #SMALL FULL STOP	U+FE52	SMALL FULL STOP	
    '． {0,}$', #FULLWIDTH FULL STOP	U+FF0E	FULLWIDTH FULL STOP	
    '｡ {0,}$', #HALFWIDTH IDEOGRAPHIC FULL STOP	U+FF61	HALFWIDTH IDEOGRAPHIC FULL STOP	
    '! {0,}$', #EXCLAMATION MARK	U+0021	EXCLAMATION MARK	
    '¡ {0,}$', #INVERTED EXCLAMATION MARK	U+00A1	INVERTED EXCLAMATION MARK	
    '¡ {0,}$', #EXCLAMATION MARK, INVERTED	U+00A1	INVERTED EXCLAMATION MARK	
    'ǃ {0,}$', #latin letter exclamation mark	U+01C3	LATIN LETTER RETROFLEX CLICK	
    'ǃ {0,}$', #exclamation mark, latin letter	U+01C3	LATIN LETTER RETROFLEX CLICK	
    'ǃ {0,}$', #LATIN LETTER EXCLAMATION MARK	U+01C3	LATIN LETTER RETROFLEX CLICK	
    '՜ {0,}$', #ARMENIAN EXCLAMATION MARK	U+055C	ARMENIAN EXCLAMATION MARK	
    '՜ {0,}$', #EXCLAMATION MARK, ARMENIAN	U+055C	ARMENIAN EXCLAMATION MARK	
    '߹ {0,}$', #NKO EXCLAMATION MARK	U+07F9	NKO EXCLAMATION MARK	
    '᥄ {0,}$', #LIMBU EXCLAMATION MARK	U+1944	LIMBU EXCLAMATION MARK	
    '᥄ {0,}$', #EXCLAMATION MARK, LIMBU	U+1944	LIMBU EXCLAMATION MARK	
    '‼ {0,}$', #DOUBLE EXCLAMATION MARK	U+203C	DOUBLE EXCLAMATION MARK	
    '‼ {0,}$', #EXCLAMATION MARK, DOUBLE	U+203C	DOUBLE EXCLAMATION MARK	
    '⁈ {0,}$', #QUESTION EXCLAMATION MARK	U+2048	QUESTION EXCLAMATION MARK	
    '⁈ {0,}$', #EXCLAMATION MARK, QUESTION	U+2048	QUESTION EXCLAMATION MARK	
    '❕ {0,}$', #WHITE EXCLAMATION MARK ORNAMENT	U+2755	WHITE EXCLAMATION MARK ORNAMENT	
    '❕ {0,}$', #EXCLAMATION MARK ORNAMENT, WHITE	U+2755	WHITE EXCLAMATION MARK ORNAMENT	
    '❗ {0,}$', #HEAVY EXCLAMATION MARK SYMBOL	U+2757	HEAVY EXCLAMATION MARK SYMBOL	
    '❢ {0,}$', #HEAVY EXCLAMATION MARK ORNAMENT	U+2762	HEAVY EXCLAMATION MARK ORNAMENT	
    '❢ {0,}$', #EXCLAMATION MARK ORNAMENT, HEAVY	U+2762	HEAVY EXCLAMATION MARK ORNAMENT	
    '❣ {0,}$', #HEAVY HEART EXCLAMATION MARK ORNAMENT	U+2763	HEAVY HEART EXCLAMATION MARK ORNAMENT	
    '⹓ {0,}$', #MEDIEVAL EXCLAMATION MARK	U+2E53	MEDIEVAL EXCLAMATION MARK	
    'ꜝ {0,}$', #MODIFIER LETTER RAISED EXCLAMATION MARK	U+A71D	MODIFIER LETTER RAISED EXCLAMATION MARK	
    'ꜞ {0,}$', #MODIFIER LETTER RAISED INVERTED EXCLAMATION MARK	U+A71E	MODIFIER LETTER RAISED INVERTED EXCLAMATION MARK	
    'ꜟ {0,}$', #MODIFIER LETTER LOW INVERTED EXCLAMATION MARK	U+A71F	MODIFIER LETTER LOW INVERTED EXCLAMATION MARK	
    '︕ {0,}$', #PRESENTATION FORM FOR VERTICAL EXCLAMATION MARK	U+FE15	PRESENTATION FORM FOR VERTICAL EXCLAMATION MARK	
    '﹗ {0,}$', #SMALL EXCLAMATION MARK	U+FE57	SMALL EXCLAMATION MARK	
    '！ {0,}$', #FULLWIDTH EXCLAMATION MARK	U+FF01	FULLWIDTH EXCLAMATION MARK	
    '！ {0,}$', #FULLWIDTH EXCLAMATION MARK	U+FF01	FULLWIDTH EXCLAMATION MARK	
    '; {0,}$', #question mark, greek	U+003B	SEMICOLON	
    '\; {0,}$', #greek question mark	U+003B	SEMICOLON	
    '\? {0,}$', #QUESTION MARK	U+003F	QUESTION MARK	
    '¿ {0,}$', #INVERTED QUESTION MARK	U+00BF	INVERTED QUESTION MARK	
    '¿ {0,}$', #question mark, turned	U+00BF	INVERTED QUESTION MARK	
    '¿ {0,}$', #QUESTION MARK, INVERTED	U+00BF	INVERTED QUESTION MARK	
    '¿ {0,}$', #turned question mark	U+00BF	INVERTED QUESTION MARK	
    '; {0,}$', #GREEK QUESTION MARK	U+037E	GREEK QUESTION MARK	
    '; {0,}$', #QUESTION MARK, GREEK	U+037E	GREEK QUESTION MARK	
    '՞ {0,}$', #ARMENIAN QUESTION MARK	U+055E	ARMENIAN QUESTION MARK	
    '՞ {0,}$', #QUESTION MARK, ARMENIAN	U+055E	ARMENIAN QUESTION MARK	
    '؟ {0,}$', #ARABIC QUESTION MARK	U+061F	ARABIC QUESTION MARK	
    '؟ {0,}$', #QUESTION MARK, ARABIC	U+061F	ARABIC QUESTION MARK	
    '፧ {0,}$', #ETHIOPIC QUESTION MARK	U+1367	ETHIOPIC QUESTION MARK	
    '፧ {0,}$', #QUESTION MARK, ETHIOPIC	U+1367	ETHIOPIC QUESTION MARK	
    '᥅ {0,}$', #LIMBU QUESTION MARK	U+1945	LIMBU QUESTION MARK	
    '᥅ {0,}$', #QUESTION MARK, LIMBU	U+1945	LIMBU QUESTION MARK	
    '⁇ {0,}$', #DOUBLE QUESTION MARK	U+2047	DOUBLE QUESTION MARK	
    '⁇ {0,}$', #QUESTION MARK, DOUBLE	U+2047	DOUBLE QUESTION MARK	
    '⁉ {0,}$', #EXCLAMATION QUESTION MARK	U+2049	EXCLAMATION QUESTION MARK	
    '⁉ {0,}$', #QUESTION MARK, EXCLAMATION	U+2049	EXCLAMATION QUESTION MARK	
    '❓ {0,}$', #BLACK QUESTION MARK ORNAMENT	U+2753	BLACK QUESTION MARK ORNAMENT	
    '❓ {0,}$', #QUESTION MARK ORNAMENT, BLACK	U+2753	BLACK QUESTION MARK ORNAMENT
    '❔ {0,}$', #WHITE QUESTION MARK ORNAMENT	U+2754	WHITE QUESTION MARK ORNAMENT	
    '❔ {0,}$', #QUESTION MARK ORNAMENT, WHITE	U+2754	WHITE QUESTION MARK ORNAMENT	
    '⩻ {0,}$', #LESS-THAN WITH QUESTION MARK ABOVE	U+2A7B	LESS-THAN WITH QUESTION MARK ABOVE	
    '⩼ {0,}$', #GREATER-THAN WITH QUESTION MARK ABOVE	U+2A7C	GREATER-THAN WITH QUESTION MARK ABOVE	
    '⳺ {0,}$', #COPTIC OLD NUBIAN DIRECT QUESTION MARK	U+2CFA	COPTIC OLD NUBIAN DIRECT QUESTION MARK	
    '⳻ {0,}$', #COPTIC OLD NUBIAN INDIRECT QUESTION MARK	U+2CFB	COPTIC OLD NUBIAN INDIRECT QUESTION MARK	
    '⸮ {0,}$', #REVERSED QUESTION MARK	U+2E2E	REVERSED QUESTION MARK	
    '⹔ {0,}$', #MEDIEVAL QUESTION MARK	U+2E54	MEDIEVAL QUESTION MARK	
    '꘏ {0,}$', #VAI QUESTION MARK	U+A60F	VAI QUESTION MARK	
    '꛷ {0,}$', #BAMUM QUESTION MARK	U+A6F7	BAMUM QUESTION MARK	
    '︖ {0,}$', #PRESENTATION FORM FOR VERTICAL QUESTION MARK	U+FE16	PRESENTATION FORM FOR VERTICAL QUESTION MARK	
    '﹖ {0,}$', #SMALL QUESTION MARK	U+FE56	SMALL QUESTION MARK	
    '？ {0,}$', #FULLWIDTH QUESTION MARK	U+FF1F	FULLWIDTH QUESTION MARK	
]
eol_conditional_array = ['\" {0,}$', u'\u201D {0,}$', u'\)']
bol_array = ['^[A-Z]']

# Colors : grey and pink backgroud to ignore
# https://learn.microsoft.com/en-us/office/vba/api/word.wdcolor
shading_color_ignore_text = ['FFD320', 'D9D9D9', 'BFBFBF', 'A6A6A6', '808080', 'FF00FF', 'FF0000', 'F3F3F3', 'E6E6E6', 'E0E0E0', 'CCCCCC', 'C0C0C0', 'B3B3B3', 'A0A0A0', '999999', '8C8C8C',  '737373', '666666', '606060', '595959', '4C4C4C', '404040', '333333', '262626', '202020', '191919', '0C0C0C']

html_file_path = ''

nb_character_total = 0
MAX_LINE_SIZE = 36
COUNTRY_QUERY_HTTP_TIMEOUT = 3

# Maximum 5000 characters on the free version
# but only 1500 if not logged on to deepl with free version
deepl_max_char_bloc_size_key = ['deepl', 'no_account','maximum_character_block']
deepl_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, deepl_max_char_bloc_size_key)

deepl_sleep_wait_translation_seconds = 0.1
translation_errors_count = 0

word_file_to_translate = args.docxfile

viewdocx = args.viewdocx

xlsxreplacefile = args.xlsxreplacefile
dest_font = args.destfont
split_translation = args.split
use_api = args.useapi
#use_browser = args.useapi

showbrowser = args.showbrowser
exitonsuccess = args.exitonsuccess
splitonly = args.splitonly
if splitonly:
    split_translation = True

driver = None

src_lang = args.srclang
dest_lang = args.destlang
if dest_lang is not None:
    dest_lang = dest_lang.lower()
else:
    dest_lang = ""
    if not splitonly:
        dest_lang = input ("Please enter language translation code (fr,de,ru,hi,etc.)")

cjk_segmenter = None 
if dest_lang == 'zh-cn':
    dest_lang = 'zh-CN'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'zh-tw':
    dest_lang = 'zh-TW'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'th':
    from newmm_tokenizer.tokenizer import word_tokenize
if dest_lang == 'zh' or dest_lang == 'ja' or dest_lang == 'ko':
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'fa':
    from hazm import Normalizer
    my_hazm_normalizer = Normalizer()

valid_online_json = validate_json_string(json_online_configuration)
if not valid_online_json == True:
    print(f"json_online_configuration={json_online_configuration}")
    print(f"Warning: Json file at {json_configuration_url} is not valid. Ignoring this configuration file.")
else:
    #print(f"Using JSON configuration file at {json_configuration_url} : OK")
    pass
    
valid_local_json = validate_json_string(local_json_contents)
if os.path.isfile(configuration_file_full_path):
    if not valid_local_json == True:
        print(f"Warning: Json file at {configuration_file_full_path} is not valid. Ignoring this configuration file.")
    else:
        print(f"Using JSON configuration file at {configuration_file_full_path}")

print("")


src_lang_name = (google_translate_lang_codes.get(src_lang))
if src_lang_name is None:
    src_lang_name = ""
    if not splitonly:
        print("Source language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Source language name for '%s' : %s" % (src_lang, src_lang_name))

dest_lang_name = (google_translate_lang_codes.get(dest_lang))

if dest_lang_name is None:
    dest_lang_name = deepl_translate_lang_codes.get(dest_lang)
    if not splitonly:
        print("Target language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Target language name for '%s' : %s" % (dest_lang, dest_lang_name))

dest_lang_tag = ""
try:
    dest_lang_tag = office_language_tags[dest_lang]
except:
    pass

translation_engine = args.engine

if translation_engine is not None:
    translation_engine = translation_engine.lower()
else:
    translation_engine = ""

if translation_engine == 'yandex' or translation_engine == 'perplexity':
    showbrowser = True

                                     
elif translation_engine == 'deepl':
    translation_engine = 'deepl'
elif translation_engine == 'chatgpt':
    translation_engine = 'chatgpt'
elif translation_engine == 'perplexity':
    translation_engine = 'perplexity'
else:
    translation_engine = 'google'


perplexity_max_char_bloc_size_key = ['perplexity', 'account','maximum_character_block']
perplexity_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, perplexity_max_char_bloc_size_key)

if translation_engine == 'perplexity':
    MAX_TRANSLATION_BLOCK_SIZE = perplexity_maximum_character_block
else:
    MAX_TRANSLATION_BLOCK_SIZE = deepl_maximum_character_block
# Override MAX_TRANSLATION_BLOCK_SIZE value after logging on Deepl


engine_method = args.enginemethod
engine_method = "%s" % engine_method
engine_method = engine_method.strip().lower()

translation_array = []

if splitonly:
    engine_method = ''
elif translation_engine == 'google':
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
    elif engine_method  == 'singlephrase':
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    elif engine_method =='xlsxfile':
        engine_method = 'xlsxfile'
        # There is a bug on xlsxfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'textfile':
        engine_method = 'textfile'
        # There is a bug on textfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'javascript':
        engine_method = 'javascript'
    else:
        engine_method = 'javascript'
elif translation_engine == 'deepl':
    if engine_method == 'singlephrase' or use_api == True:
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    else:
        engine_method = 'phrasesblock'
elif translation_engine == 'chatgpt':
    engine_method = 'phrasesblock'
elif translation_engine == 'perplexity':
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
    else:
        engine_method = 'phrasesblock'
else:
    engine_method = "web"

def lineno():
    """Returns the current line number in our program."""
    return inspect.currentframe().f_back.f_lineno
    
def linux_distribution():
    try:
        return platform.linux_distribution()
    except:
        return "N/A"

def print_os_info():

    print("""Python version: %s
    dist: %s
    linux_distribution: %s
    system: %s
    machine: %s
    platform: %s
    uname: %s
    version: %s
    mac_ver: %s
    """ % (
    sys.version.split('\n'),
    str(platform.dist()),
    linux_distribution(),
    platform.system(),
    platform.machine(),
    platform.platform(),
    platform.uname(),
    platform.version(),
    platform.mac_ver(),
    ))


if not os.path.exists(word_file_to_translate) :
    print("ERROR: File not found: %s" % (word_file_to_translate))
    sys.exit(1)

splitted_filename = os.path.splitext(os.path.basename(word_file_to_translate))

# number of segment separated by dot in the docx filename
splitted_filename_size = len(splitted_filename)

docx_file_name =  "%s%s" % (splitted_filename[splitted_filename_size-2], splitted_filename[splitted_filename_size-1])

if splitted_filename_size > 1:
    word_file_to_translate_extension = splitted_filename[splitted_filename_size-1].lower()

if word_file_to_translate_extension == ".docx":
    try:
        docxdoc = docx.Document(word_file_to_translate)
    except:
        print(f"Error, file {word_file_to_translate} does not appear to be a valid Microsoft Word docx file.")
        print("Please check that the file is a valid document and rerun on a valid Microsoft docx document.\n")
        
        print("\nDeveloper: %s" % (E_mail_str))
        print("Program version: %s\n" % (PROGRAM_VERSION))
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(2)
    styles = docxdoc.styles
    
    if dest_lang_tag != '':
        styles_element = docxdoc.styles.element
        try:
            # Some office suite like WPS does not handle language tag in a document, ignore it
            rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
            lang_default = rpr_default.xpath('w:lang')[0]
            lang_default.set(docx.oxml.shared.qn('w:val'),dest_lang_tag)
        except Exception:
            # Ignore the language tag of the document, it is not supported by some office suites
            pass

    # Create Right to Left Style if it is not found
    try:
        rtlstyle = styles['rtlstyle']
    except Exception:
        rtlstyle = docxdoc.styles.add_style('rtlstyle', WD_STYLE_TYPE.CHARACTER)
    if dest_lang == "" or dest_lang is None:
        dest_lang_name_from_cell = docxdoc.tables[0].cell(1, 2).text
        print("Lang cell: %s" % (dest_lang_name_from_cell))
        for lang_code, lang_name in google_translate_lang_codes.items():
            #print("%s : %s" % (lang_code, lang_name))
            if dest_lang_name_from_cell == lang_name:
                dest_lang = lang_code

print("File: %s" %(args.docxfile))
print("Language code: %s" %(dest_lang))
print("Language name: %s" %(dest_lang_name))

print("Destination font: %s" %(dest_font))
print("Split: %s" %(split_translation))
print("Splitonly: %s" %(splitonly))


tmx_file_path = "%s\%s_%s.tmx" % (os.path.dirname(word_file_to_translate), os.path.splitext(os.path.basename(word_file_to_translate))[0],dest_lang)
#print(tmx_file_path)


print("Extension: %s" % (word_file_to_translate_extension))

if not os.path.exists(word_file_to_translate):
    print("ERROR: docxfile '%s' not found, exiting." % (word_file_to_translate))


if word_file_to_translate_extension != ".docx":
    print("ERROR: not a docx file. Please convert to docx first then run on docx file. Exiting.")
    if not silent:
        input("Enter to close program")
    else:
        print("Program ended with errors")
    os._exit(3)

print("")


location_primary_country_checker_url_key = ["local_configuration", "json_filename_path"]
location_primary_country_checker_url_key = ["location", "primary_country_checker_url"]
location_primary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_primary_country_checker_url_key)

location_secondary_country_checker_url_key = ["location", "secondary_country_checker_url"]
location_secondary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_secondary_country_checker_url_key)

location_http_query_timeout_key = ["location", "http_query_timeout"]
location_http_query_timeout = get_nested_value_from_json_array(json_configuration_array, location_http_query_timeout_key)

# Check if location_http_query_timeout is not an integer > 0
if not isinstance(location_http_query_timeout, int) or location_http_query_timeout <= 0:
    location_http_query_timeout = COUNTRY_QUERY_HTTP_TIMEOUT  # Set to 3 if the condition is not met

chrome_driver_restricted_countries_key = ["chrome_driver", "restricted_countries"]
chrome_driver_restricted_countries = get_nested_value_from_json_array(json_configuration_array, chrome_driver_restricted_countries_key)

chrome_driver_mirror_url_key = ["chrome_driver", "mirror_url"]
chrome_driver_mirror_url = get_nested_value_from_json_array(json_configuration_array, chrome_driver_mirror_url_key)

#print(f"location_primary_country_checker_url = {location_primary_country_checker_url}")
#print(f"location_secondary_country_checker_url = {location_secondary_country_checker_url}")
#print(f"chrome_driver_restricted_countries = {chrome_driver_restricted_countries}")
#print(f"chrome_driver_mirror_url = {chrome_driver_mirror_url}")

def fetch_country_data(url):
    """Fetch country data from the specified URL."""
    try:
        response = requests.get(url, timeout=location_http_query_timeout)
        response.raise_for_status()  # Check if the request was successful (status code 200)
        
        # Parse the JSON response
        data = response.json()
        
        # Check if the status is success and return the country name
        if data.get("status") == "success":
            return data.get('country')
        else:
            print(f"Failed to retrieve IP information: {data.get('message')}")
            return None
            
    except requests.exceptions.RequestException as e:
        print(f"HTTP request failed: {e}")
    except json.JSONDecodeError:
        print("Failed to parse the JSON response.")
    return None

def check_mirror_url(url):
    """Check if the mirror URL responds with HTTP 200 or 400 status codes."""
    try:
        response = requests.get(url, timeout=location_http_query_timeout)
        return response.status_code in [200, 400]
    except requests.exceptions.RequestException as e:
        print(f"Mirror URL check failed: {e}")
        return False

def set_SE_DRIVER_MIRROR_URL_if_needed(country_name, mirror_url):
    """Set the SE_DRIVER_MIRROR_URL environment variable if the country is restricted and mirror URL is valid."""
    if country_name in chrome_driver_restricted_countries:
        print(f"The host country ({country_name}) is restricted from downloading Google Chrome Driver, using proxy to bypass restrictions...")
        
        # Check the mirror URL and set environment variable if it responds with HTTP 200 or 400
        if check_mirror_url(mirror_url):
            os.environ['SE_DRIVER_MIRROR_URL'] = mirror_url
            print(f"SE_DRIVER_MIRROR_URL set to: {os.environ['SE_DRIVER_MIRROR_URL']}")
        else:
            print(f"Mirror URL ({mirror_url}) did not respond with HTTP 200 or 400.")
    else:
        print(f"Using Google Chrome Driver from {country_name}...")


# Set chrome driver download proxy URL for restricted countries
country_name = fetch_country_data(location_primary_country_checker_url)

# If primary URL fails or does not return a valid country name, fallback to the secondary URL
if not country_name:
    print("Falling back to secondary URL...")
    country_name = fetch_country_data(location_secondary_country_checker_url)

# Set environment variable if needed
set_SE_DRIVER_MIRROR_URL_if_needed(country_name, chrome_driver_mirror_url)

# Set up Chrome options
# Set the user-data-dir to the parent of the profiles

#chrome_options.add_argument(f"--user-data-dir={user_data_dir}") 
#chrome_options.add_argument(r'--profile-directory=Default')


user_data_dir = fr"C:\Temp\Chrome"
# Set the user-data-dir to the parent of the profiles



chrome_options = Options()
chrome_options.add_argument("--disable-web-security")
chrome_options.add_argument("--disable-xss-auditor")
chrome_options.add_argument("--lang=en-GB")
#chrome_options.add_argument("--verbose")
chrome_options.add_argument("--log-level=3")  # fatal
chrome_options.add_argument("--password-store=basic")


if  translation_engine.lower() == "chatgpt":
    print(f"Using Chrome profile")
    print(f"Using user data dir: {user_data_dir}")
    chrome_options.add_argument(f"--user-data-dir={user_data_dir}") 
    chrome_options.add_argument(r'--profile-directory=Default')

#chrome_options.add_argument("load-extension=C:\\Users\Patriot\\AppData\\Local\\Google\\Chrome\\User Data\\Default\\Extensions\\mooikfkahbdckldjjndioackbalphokd\\3.17.0_0")

if not showbrowser and translation_engine.lower() != "deepl" and translation_engine.lower() != "chatgpt" :
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--lang=en-GB")
                              

if translation_engine.lower() == "chatgpt":
    # Get the Windows username
    try:
        username = os.getlogin()
        home_dir = os.path.expanduser("~")

        # Construct the Chrome user data directory path
        user_data_dir = os.path.join(home_dir, "AppData", "Local", "Google", "Chrome", "User Data")
        
        user_data_dir = fr"C:\Temp\Chrome"
        print(f"Using Chrome user data directory: {user_data_dir}")

        # Set up ChromeOptions
        print(f"Using Chrome user data directory: {user_data_dir}")
        chrome_options.add_argument(f"--user-data-dir={user_data_dir}")  # Path to the user data directory
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_argument("--profile-directory=Default")  # Use the "Default" profile

        print(f"Using Chrome user data directory: {user_data_dir}")
    #word_file_to_translate = r'X:\travail\smtv-hindi\NWN 584 sf2 - table fix1.doc'
    except:
        var = traceback.format_exc()
        print(var)
        print("Failed to add chrome options")


# Used to tokenize thai
#thai_segmenter = thai_tokenizer_tokenizer()
#word_tokenize(text)

#translator = Translator(service_urls=['translate.google.com'], user_agent='Mozilla/5.0 (Windows NT 6.1; Win64; x64; rv:47.0) Gecko/20100101 Firefox/47.0')
#user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36')

#driver = uc.Chrome()
#driver.set_window_position(0, 350)
#driver.set_window_size(400, 650)

#driver.get("https://translate.google.com/#%s/%s/" % (src_lang,dest_lang))
#driver = uc.Firefox()

def get_translated_cells_content(lineno, to_translate):
    print("get_translation for line %d" % (lineno))
    print("from_text_nb_lines_in_phrase %d" % (from_text_nb_lines_in_phrase[lineno]))
    translation = ""

    if dest_lang.lower() == 'ja' or dest_lang.lower() == 'zh-cn' or dest_lang.lower() == 'zh-tw' or dest_lang.lower() == 'ko':
        cell_space = ''
    else:
        cell_space = ' '

    last_row_n = from_text_nb_lines_in_phrase[lineno] + lineno
    for row_n in range(lineno, last_row_n):
        #cell_text = docxdoc.tables[0].cell(row_n, 2).text
        cell_text = table_cells[row_n][2].text
        cell_text = cell_text.strip()
        if cell_text != "":
            print("adding cell %s " % (cell_text))
            if row_n == lineno:
                translation = cell_text
            else:
                translation = translation + cell_space + cell_text
    return translation

found_google_cookies_consent_button = False
google_translate_first_page_loaded = False

def selenium_chrome_translate_get_from_text_array(to_translate, index):
    #print("to_translate   :%s" % to_translate)
    #print("translation %3d:%s" % (index, translation_array[index - 1]))
    #input("selenium_chrome_translate_get_from_text_array")
    #input("In selenium_chrome_translate_get_from_text_array")
    return translation_array[index - 1]

def selenium_chrome_google_translate(to_translate):
    global found_google_cookies_consent_button
    global google_translate_first_page_loaded
    try:
        translation = ''
        #to_translate_encoding = chardet.detect(bytes(to_translate))['encoding']
        #print("to_translate_encoding=%s" % (to_translate_encoding))
        
        if not google_translate_first_page_loaded:
            selenium_chrome_google_click_cookies_consent_button()
        #print("HERE **********")
        #input("Here")
            
        #driver.get("https://translate.google.com/?sl=%s&tl=%s&text=%s&op=translate" % (src_lang,dest_lang,""))
        #driver.get("https://translate.google.com/?sl=%s&tl=%s&op=translate&text=%s" % (src_lang,dest_lang,html.escape(to_translate)))
        
        #print ("to_translate before using in url:")
        #print(to_translate)
        #print ("***************")
        
        
        #to_translate_escaped = html.escape(to_translate)
        to_translate_escaped = to_translate
        # This is necessary because & in the text to be translated will from the remaining
        # from the text fild in the URL and will not translate after the &
        
        to_translate_escaped = to_translate_escaped.replace('%','%25')
        to_translate_escaped = to_translate_escaped.replace('&','%26')
        
        query_translation = {
            "sl" : src_lang,
            "tl" : dest_lang,
            "docxfile" : to_translate
        }
        
        base_url = "https://translate.google.com/"
        encoded_params = urlencode(query_translation, quote_via=quote_plus)
        url = f"{base_url}?{encoded_params}"
        
        to_translate_add_new_line = '%0A '.join(to_translate_escaped.split('\n'))
        translation_url = "https://translate.google.com/?sl=%s&tl=%s&op=translate&text=%s" % (src_lang,dest_lang,to_translate_add_new_line)
        driver.get(translation_url)
        #print("---------------------------------------------------------")
        #print("translation_url")
        #print(translation_url)
        #print("---------------------------------------------------------")
        #print("*********************************************************")
        #print("url")
        #print(url)
        #print("*********************************************************")
        
        #print("to_translate_add_new_line:")
        #print(to_translate_add_new_line)
        #print ("***************")
        #input("Here")

        #input("Here")
        
        input_element = "//div[@id='input-wrap']/textarea"
        input_element = "//textarea[@id='source']"
        input_element = "//textarea"
        
        #input_button = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, input_element)))))
        #driver.find_element(By.XPATH, input_element).clear()
        #to_translate_utf8 = to_translate.decode('utf-8')
        to_translate_utf8 = to_translate
        #input_button.send_keys (to_translate_utf8)

        time.sleep(0.2)
        #driver.execute_script('document.getElementById(\'//textarea[@id=\\\'source\\\']\').setAttribute(\'value\', \'Hello world !\');')
        #driver.execute_script("arguments[0].value = arguments[1];", input_button, to_translate_utf8)

        try:
            copy_translation_element = "//div[4]/div[4]/div"
            copy_translation_element = "//i[contains(.,'content_copy')]"
            
            copy_translation_button = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, copy_translation_element)))
        except :
            var = traceback.format_exc()
            print(var)
            
        # EditTranslationElement = "xpath=//button/div[2]"
        # EditTranslationButton = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, EditTranslationElement)))))
        
        # driver.execute_script("scrollBy(0,-1000);")
        # actions.move_to_element(EditTranslationButton).perform()
        # sleep(0.1)
        #driver.set_window_size(800, 700)
        # EditTranslationButton.click()
        
        res_element_xpath = "//textarea[@lang='%s']" % (dest_lang)
        
        regex_still_translating_str = '$Translation'
        pos_separator_phonetic = 0
        #if re.search(regex_still_translating_str, to_translate_utf8):
        if re.search(regex_still_translating_str, to_translate):
            time.sleep(4)
            try:
                #res_element_xpath = "//c-wiz/div/div[2]/c-wiz/div[2]/c-wiz/div[1]/div[2]/div[2]/c-wiz[2]/div[5]/div/div[3]/div[1]/div/div[1]/div[1]/textarea"
                
                result_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, res_element_xpath)))
                translation = result_element.get_attribute('innerHTML')
                #translation = html.unescape(translation)
            except :
                var = traceback.format_exc()
                print(var)
        else:
            try:
                #res_element_xpath = "//c-wiz/div/div[2]/c-wiz/div[2]/c-wiz/div[1]/div[2]/div[2]/c-wiz[2]/div[5]/div/div[3]/div[1]/div/div[1]/div[1]/textarea"
                #res_element_xpath = "xpath=//div[6]/div/div/span/span/span"
                #res_element_xpath = "//textarea[@lang='%s']" % (dest_lang)
                
                result_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, res_element_xpath)))
                translation = result_element.get_attribute('innerHTML')
                #translation = html.unescape(translation)
                
                #print("translation:")
                #print(translation)
                #print("*******************")
                
                #page_source_str = driver.page_source
                #print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                #with open('before.html', 'w', encoding="utf-8") as f:
                #    f.write(page_source_str)
                #f.close()
                #input("wait here")
                
            except Exception:
                #page_source_str = driver.page_source
                #print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                #with open('before.html', 'w', encoding="utf-8") as f:
                #    f.write(page_source_str)
                #f.close()
                #input("wait here")
                
                var = traceback.format_exc()
                print(var)
            while re.search(regex_still_translating_str, translation):
                print("")
                print("Still waiting for translation........")
                print("")
                time.sleep(1)
                try:
                    #result_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, ".result"))))
                    result_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, res_element_xpath)))
                    
                    translation = result_element.get_attribute('innerHTML')
                    #translation = html.unescape(translation)
                    
                    #print("translation:")
                    #print(translation)
                    #print("*******************")
                except Exception:
                    var = traceback.format_exc()
                    print(var)

        #pos_new_line = translation.find("\n")
        #if pos_new_line > 0:
        #    translation = translation[0:pos_new_line]
        #    #print("\n\npos_new_line=%s\n\n\n" % (pos_new_line))

    except Exception:
        var = traceback.format_exc()
        print(var)
        sys.exit(4)
   
    page_source_str = driver.page_source
    #print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
    #with open('before.html', 'w', encoding="utf-8") as f:
    #    f.write(page_source_str)
    #f.close()
    #input("wait here")
                
    #print("returned translation:")
    #print("#########################")
    #print(translation)
    #print("#########################")
    #input("press enter")
    return translation


def selenium_chrome_translate_maxchar_blocks():
    global found_google_cookies_consent_button
    global google_translate_first_page_loaded
    global docxfile_table_number_of_phrases
    global translation_errors_count
    global deepl_sleep_wait_translation_seconds
    global selenium_chrome_machine_translate_once
    global service, driver, chrome_options

    translation_succeded = True

    try:
        
        blocks_nchar_max_to_translate_array_len = len(blocks_nchar_max_to_translate_array)
        
        text_translated_document_str = ""
        #for current_block_str in blocks_nchar_max_to_translate_array :
        print("")

        for current_block_no in range(0, blocks_nchar_max_to_translate_array_len):
            current_block_str = blocks_nchar_max_to_translate_array[current_block_no]
            current_block_str_len = len(current_block_str)
            print("Translating block %d/%d, %d characters..." % (current_block_no + 1, blocks_nchar_max_to_translate_array_len, current_block_str_len))
            translation = ""
            translation_try_count = 1
            if translation_engine == 'perplexity':
                max_try_count = 6
            else:
                max_try_count = 4
            
            if ((current_block_no + 1) % 4) == 0:
                print("Cleaning up perplexity cookies...")
                driver.delete_all_cookies()
                
            #print("--index-- : %d" % index)
            #to_translate_str = str(to_translate)
            to_translate = current_block_str
            try:
                while translation_try_count < max_try_count and translation == "":
                    if translation_try_count > 1:
                        print("Retrying to translate again (%d)..." % (translation_try_count))
                        translation_errors_count = translation_errors_count + 1
                        deepl_sleep_wait_translation_seconds  = deepl_sleep_wait_translation_seconds * 1.1
                        print("%d translation retry so far..." % (translation_errors_count))
                    if translation_engine == 'deepl':
                        #print(to_translate)
                        #print(translation_try_count)
                        translation_succeded, translation = selenium_chrome_deepl_translate(to_translate, translation_try_count - 1)
                        if translation_succeded == False:
                            print("Deepl translation permited limit exeeded")
                            return translation_succeded, []
                    if translation_engine == 'chatgpt':
                        #print(to_translate)
                        #print(translation_try_count)
                        translation_succeded, translation = selenium_chrome_chatgpt_translate(to_translate, translation_try_count - 1)
                        if translation_succeded == False:
                            print("Chatgpt translation failed")
                            return translation_succeded, []
                    if translation_engine == 'perplexity':
                        #print(to_translate)
                        #print(translation_try_count)
                        #print("Translating with Perplexity AI")
                        #translation_succeded, translation = selenium_chrome_perplexity_translate(to_translate, translation_try_count - 1)
                        if engine_method == 'api':
                            translation_succeded, translation = perplexity_api_translate(to_translate, translation_try_count - 1)
                        else:
                            translation_succeded, translation = selenium_chrome_perplexity_translate(to_translate, translation_try_count - 1)
                            # Sleep enough time to let perplexity page display a new converstation
                            time.sleep(0.25)
                        
                        if translation_succeded == False:
                            print("Perplexity translation failed")
                            return translation_succeded, []
                    elif translation_engine == 'google':
                        if engine_method == 'xlsxfile':
                            translation = selenium_chrome_machine_translate_once(to_translate, index)
                        if engine_method == 'textfile':
                            translation = selenium_chrome_machine_translate_once(to_translate, index)
                        else:
                            #print("to_translate")   
                            #print(to_translate)            
                            translation = selenium_chrome_google_translate(to_translate)
                    translation_try_count = translation_try_count + 1
            except:
                print("Error in selenium_chrome_translate_maxchar_blocks...")
                var = traceback.format_exc()
                print(var)
                print("Error in selenium_chrome_machine_translate")
            if current_block_no == 0:
                text_translated_document_str = translation
            else:
                text_translated_document_str = text_translated_document_str + "\n" + translation

            #print("Translation block (%d):" % (len(translation.split('\n'))))
            #print(translation)
            #input("Enter go on after viewing translated block")
            
        
        #print("text_translated_document_str:///////////////////////////////////////")
        #print(text_translated_document_str)
        #print("/////////////////////////////////////////")
        #input("Press enter")
        
        #text_translated_document_str = html.unescape(text_translated_document_str)
        
        translation_array = text_translated_document_str.split('\n')
        
        text_translated_document_str_nb_lines = len(translation_array)
        
        #print ("text_translated_document_str_nb_linestext_translated_document_str_nb_lines: %s" % text_translated_document_str_nb_lines)
        #print ("docxfile_table_number_of_phrases: %s" % docxfile_table_number_of_phrases)
        
        if docxfile_table_number_of_phrases == text_translated_document_str_nb_lines:
            #print("OK, we got the right number of translated lines !")
            #input("Here")
            pass
        else:
            print("oups ! we got %s translated lines out of %s" % (text_translated_document_str_nb_lines, docxfile_table_number_of_phrases))
            translation_succeded = False

        
        #input("finished translating")
        
        #print("text_translated_document_str:")
        #print(text_translated_document_str)
        
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(5)
    return translation_succeded, translation_array

def selenium_chrome_google_click_cookies_consent_button():
    global found_google_cookies_consent_button
    global google_translate_first_page_loaded
    global chrome_options
    global driver
    global chromedriverpath
    try:
        translation = ''
        browse_file_element_xpath = "//label[contains(.,'Browse your computer')]"

        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            print("Opening google translation page...")
            if engine_method == 'textfile' or engine_method == 'xlsxfile':
                driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (src_lang,dest_lang))
                (driver.page_source).encode('utf-8')
            else:
                driver.get("https://translate.google.com/?sl=%s&tl=%s&op=translate" % (src_lang,dest_lang))
                (driver.page_source).encode('utf-8')
        
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            print("Waiting for cookies consent button...")
            try:
                browse_file_element = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
                found_google_cookies_consent_button = True
            except:
                pass
                
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            consent_cookies_element = "//form/div/div/button/span"
            try:
                consent_cookies_button = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH, consent_cookies_element)))
                consent_cookies_button.click()
                browse_file_element = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
                found_google_cookies_consent_button = True
                
            except:
                pass
                
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            consent_cookies_element = "//button/span"
            try:
                consent_cookies_button = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH, consent_cookies_element)))
                consent_cookies_button.click()
                browse_file_element = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
                found_google_cookies_consent_button = True
                
            except:
                pass
                
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            consent_cookies_element = "//button/div[2]"
            try:
                consent_cookies_button = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH, consent_cookies_element)))
                consent_cookies_button.click()
                browse_file_element = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
                found_google_cookies_consent_button = True
                
            except:
                pass
        
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            try:
                browse_file_element = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
                found_google_cookies_consent_button = True
            except:
                pass
                
        if not found_google_cookies_consent_button and not google_translate_first_page_loaded:
            try:
                                
                if not showbrowser:
                    chrome_options = Options()
                    chrome_options.add_argument("--disable-web-security")
                    chrome_options.add_argument("--disable-xss-auditor")
                    #chrome_options.add_argument("--verbose")
                    chrome_options.add_argument("--log-level=3")  # fatal
                    chrome_options.add_argument("--lang=en-GB")
                    chrome_options.add_argument("--password-store=basic")
                    #options.add_argument(r'--user-data-dir=C:\Users\Patriot\AppData\Local\Google\Chrome\User Data') #e.g. C:\Users\You\AppData\Local\Google\Chrome\User Data
                    #options.add_argument(r'--profile-directory=C:\Users\Patriot\AppData\Local\Google\Chrome\User Data\Default') #e.g. Profile 3
                    #input("profile options added")
                    driver = uc.Chrome(executable_path=chromedriverpath, options=chrome_options)
            
                driver.maximize_window()
                driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (src_lang,dest_lang))
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                driver.switch_to.window(driver.current_window_handle)
                print('\nCLICK ON I Agree TO CONTINUE\n')
                browse_file_element = WebDriverWait(driver, 500).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))
            except:
                pass      
            
        google_translate_first_page_loaded = True
        print("Cookies consent button cliqued...")
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(6)
        
        
def selenium_chrome_google_translate_text_file(text_file_path):
    global found_google_cookies_consent_button
    global google_translate_first_page_loaded
    global docxfile_table_number_of_phrases
    try:
        
        if not google_translate_first_page_loaded:
            selenium_chrome_google_click_cookies_consent_button()
        
        driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (src_lang,dest_lang))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        
        browse_file_element_xpath = "//label[contains(.,'Browse your computer')]"
        
        #browse_file_element = WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))))
        
        #browse_file_element.click()
        
        # Waiting for URL : https://translate.googleusercontent.com/translate_f
        
        print("Selecting file %s for uploading..." % (text_file_path))
        text_file_element_xpath = "//input[@name='file']"
        text_file_element_xpath = "//input[@id='i37']"
        text_file_element_xpath = "//div[3]/input"
        text_file_element = WebDriverWait(driver, 925).until(EC.presence_of_element_located((By.XPATH, text_file_element_xpath)))
        
        text_file_element.send_keys(text_file_path)

        #text_file_translate_button_xpath = "//div[2]/div[2]/button/span"
        text_file_translate_button_xpath = "//div[2]/div/button/span"
        
        
       
        
        
        text_file_translate_button = WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.XPATH, text_file_translate_button_xpath)))
        
        print("Clicking on Translate button...")
        text_file_translate_button.click()
        
        # Wait for result text translation page to be loaded
        loop_wait_translation_count = 200
        
        print("Waiting for translation result...")
        while(driver.current_url != 'https://translate.googleusercontent.com/translate_f' and loop_wait_translation_count > 0):
            time.sleep(0.1)
            loop_wait_translation_count = loop_wait_translation_count - 1
        
        #Wait for page status loaded to be complete
        WebDriverWait(driver, 15).until(lambda driver: driver.execute_script('return document.readyState') == 'complete')
        
        print("Translation received.")
        
        #print("loop_wait_translation_count: %d" % (200 - loop_wait_translation_count))
        #print("URL: %s" % (driver.current_url))
        
        html_translation = driver.page_source
        text_translated_document_str = html_translation.replace('<html><head><meta charset="UTF-8"></head><body><pre>', '')
        text_translated_document_str = text_translated_document_str.replace('</pre></body></html>', '')
        text_translated_document_str = html.unescape(text_translated_document_str)
        
        translation_array = text_translated_document_str.split('\n')
        
        text_translated_document_str_nb_lines = len(translation_array)
        
        #print ("text_translated_document_str_nb_linestext_translated_document_str_nb_lines: %s" % text_translated_document_str_nb_lines)
        print ("docxfile_table_number_of_phrases: %s" % docxfile_table_number_of_phrases)
        
        if docxfile_table_number_of_phrases == text_translated_document_str_nb_lines:
            #print("OK, we got the right number of translated lines !")
            pass
        else:
            print("oups ! we got %s translated lines out of %s" % (text_translated_document_str_nb_lines, docxfile_table_number_of_phrases))
            translation_succeded = False
        
        #print("text_translated_document_str:")
        #print(text_translated_document_str)
        
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(7)
    return translation_array
    
    
def selenium_chrome_google_translate_html_javascript_file(html_file_path):
    global my_hazm_normalizer, showbrowser, driver
    html_file_path_escaped = html_file_path.replace('#','%23')
    file_url = 'file://' + html_file_path_escaped
    
    nb_retry = 3
    
    while nb_retry > 0:
        nb_retry = nb_retry -1
        try:
            driver.get(file_url)
            
            print("Reading translation")

            try:
                scrollHeight = driver.execute_script("return document.body.scrollHeight")
                innerHeight = driver.execute_script("return window.innerHeight")
                bar = progressbar.ProgressBar(max_value=scrollHeight)
                bar.update(0)
            except:
                var = traceback.format_exc()
                print(var)

            paragraphs = driver.find_elements(by=By.XPATH, value='//p[@class="translation"]')

            try:
                
                # How to detect a paragraph is translated is that it has the string below
                #translated_substring_old = '<font style="vertical-align: inherit;">'
                #translated_substring_new = '<font dir="auto" style="vertical-align: inherit;"><font dir="auto" style="vertical-align: inherit;">'
                re_translated_substring = re.compile('^[ \t\r\n]{0,}<font ')
                scroll_offset_paragraph = 60
                
                for index, paragraph in enumerate(paragraphs, start=1):
                
                    #input("Time out here next line ")
                    viewport_top = driver.execute_script("return window.pageYOffset;")
                    viewport_bottom = viewport_top + driver.execute_script("return window.innerHeight;")

                    # Get the coordinates of the element
                    element_top = paragraph.location['y']
                    element_bottom = element_top + paragraph.size['height']
                    
                    paragraph_html = paragraph.get_attribute('innerHTML')
                    #print(f"{paragraph_html}")
                    
                    location = paragraph.location
                    x = location['x']
                    y = location['y']
                    scroll_position = location['y'] - scroll_offset_paragraph
                    
                    wait_translation_sleep_sec = 0.05
                    
                    # or viewport_top <= element_bottom <= viewport_bottom
                    if viewport_top <= element_top <= viewport_bottom:
                        #print("The element is displayed at the current scroll position.")
                        pass
                    else:
                        #print("The element is not displayed at the current scroll position.")
                        try:
                            driver.execute_script(f"window.scrollTo(0, {scroll_position});")
                            time.sleep(wait_translation_sleep_sec)
                            #print("The element should NOW be displayed at the current scroll position.")
                        except Exception as e:
                            #Ignore and continue if there is an error
                            #print(f"Error scrolling paragraph {index}: {str(e)}")
                            pass
                    
                    # Wait until the translation is available        
                    
                    wait_translation_max_sleep_sec = 30
                    loop_wait_translation_count = wait_translation_max_sleep_sec / wait_translation_sleep_sec
                    match_translated_tag = re_translated_substring.match(paragraph_html)
                    while not match_translated_tag and loop_wait_translation_count > 0:
                        #print(f"Sleeping in Paragraph {index}")
                        time.sleep(wait_translation_sleep_sec)
                        paragraph_html = paragraph.get_attribute('innerHTML')
                        match_translated_tag = re_translated_substring.match(paragraph_html)
                        #print(f"\n{paragraph_html}\n")
                        loop_wait_translation_count = loop_wait_translation_count - 1
                    
                    try:
                        bar.update(scroll_position + scroll_offset_paragraph)    
                    except:
                        # Ignore progressbar errors at the end
                        pass
            except:
                var = traceback.format_exc()
                print(var)
            
            bar.update(scrollHeight)
            progressbar.streams.flush()
            bar.finish()
            
            # Read translation from HTML
            html_translation = driver.page_source
            #soup = BeautifulSoup(html_translation)
            soup = BeautifulSoup(html_translation, features="lxml")
            pTags = soup.find_all('p', {'class':"translation"})
            translation_array = []
            for pTranstlation in pTags:
                pData = pTranstlation.text
                if dest_lang.lower() == 'fa':
                   pData =  my_hazm_normalizer.normalize(text=pData)
                translation_array.append(pData)

            return (translation_array)
  
        except:
            var = traceback.format_exc()
            print(var)
                        
            
            print("Here do something exit with session failed ")
                
            chrome_options = Options()
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--disable-xss-auditor")
            chrome_options.add_argument("--log-level=3")  # fatal
            chrome_options.add_argument("--lang=en-GB")
            chrome_options.add_argument("--password-store=basic")
            
            if not showbrowser:
                chrome_options.add_argument("--headless")
            
            docxfile_table_number_of_lines = numrows
            if use_api or splitonly:
                print("\nCreating a new browser for stats")
                                                      
                service = Service()                                
                driver = uc.Chrome(service=service, options=chrome_options)
                                    
                          

                              
    
def getDownLoadedFileNameFirefox(waitTime):
    driver.execute_script("window.open()")
    WebDriverWait(driver,10).until(EC.new_window_is_opened)
    driver.switch_to.window(driver.window_handles[-1])
    driver.get("about:downloads")

    endTime = time.time()+waitTime
    while True:
        try:
            fileName = driver.execute_script("return document.querySelector('#contentAreaDownloadsView .downloadMainArea .downloadContainer description:nth-of-type(1)').value")
            if fileName:
                return fileName
        except:
            pass
        time.sleep(2)
        if time.time() > endTime:
            break


# method to get the downloaded file name
def getDownLoadedFileNameChrome(waitTime):
    driver.execute_script("window.open()")
    # switch to new tab
    driver.switch_to.window(driver.window_handles[-1])
    # navigate to chrome downloads
    driver.get('chrome://downloads')
    # define the endTime
    endTime = time.time()+waitTime
    while True:
        try:
            # get downloaded percentage
            downloadPercentage = driver.execute_script(
                "return document.querySelector('downloads-manager').shadowRoot.querySelector('#downloadsList downloads-item').shadowRoot.querySelector('#progress').value")
            # check if downloadPercentage is 100 (otherwise the script will keep waiting)
            if downloadPercentage == 100:
                # return the file name once the download is completed
                return driver.execute_script("return document.querySelector('downloads-manager').shadowRoot.querySelector('#downloadsList downloads-item').shadowRoot.querySelector('div#content  #file-link').text")
        except:
            pass
        time.sleep(1)
        if time.time() > endTime:
            break


# function to wait for download to finish and then rename the latest downloaded file
def get_last_downloaded_file_path():
    # function to wait for all chrome downloads to finish
    def chrome_downloads(drv):
        if not "chrome://downloads" in drv.current_url: # if 'chrome downloads' is not current tab
            drv.execute_script("window.open('');") # open a new tab
            drv.switch_to.window(driver.window_handles[1]) # switch to the new tab
            drv.get("chrome://downloads/") # navigate to chrome downloads
        dld_file_paths = drv.execute_script("""
            return document.querySelector('downloads-manager')
            .shadowRoot.querySelector('#downloadsList')
            .items.filter(e => e.state === 'COMPLETE')
            .map(e => e.filePath || e.file_path || e.fileUrl || e.file_url);
            """)
        print("dld_file_paths=%s" % (dld_file_paths))
        #input("dld_file_paths press enter")
        return dld_file_paths
    # wait for all the downloads to be completed
    dld_file_paths = []
    while len(dld_file_paths) == 0:
        dld_file_paths = chrome_downloads(driver)
        print("len dld_file_paths=%d" % (len(dld_file_paths)))
        print("res dld_file_paths=%s" % (dld_file_paths))
        #input("Opened download status page")
        #WebDriverWait(driver, 120, 1).until(chrome_downloads) # returns list of downloaded file paths)
    # Close the current tab (chrome downloads)
    if "chrome://downloads" in driver.current_url:
        driver.close()
    # Switch back to original tab
    driver.switch_to.window(driver.window_handles[0]) 
    # get latest downloaded file name and path
    dlFilename = dld_file_paths[0] # latest downloaded file from the list
    # wait till downloaded file appears in download directory
    time_to_wait = 20 # adjust timeout as per your needs
    time_counter = 0
    while not os.path.isfile(dlFilename):
        print ("We have dlFilename=%s" % (dlFilename))
        time.sleep(5)
        time_counter += 1
        if time_counter > time_to_wait:
            break
    # rename the downloaded file
    print("dlFilename=%s" %(dlFilename))
    return dlFilename
    #shutil.move(dlFilename, os.path.join(download_dir,newFilename))
    return

def selenium_chrome_google_translate_xlsx_file(xlsx_file_path):
    global found_google_cookies_consent_button
    global google_translate_first_page_loaded
    global docxfile_table_number_of_phrases
    
    try:
        
        if not google_translate_first_page_loaded:
            selenium_chrome_google_click_cookies_consent_button()
        
        driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (src_lang,dest_lang))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        
        browse_file_element_xpath = "//label[contains(.,'Browse your computer')]"
        
        #browse_file_element = WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))))
        
        #browse_file_element.click()
        
        # Waiting for URL : https://translate.googleusercontent.com/translate_f
        
        print("Selecting file %s for uploading..." % (xlsx_file_path))
        xlsx_file_element_xpath = "//input[@name='file']"
        xlsx_file_element_xpath = "//input[@id='i37']"
        #xlsx_file_element_xpath = "//span[contains(.,'Browse your computer')]"
        #xlsx_file_element_xpath = '//button[normalize-space()="Browse your computer"]'
        xlsx_file_element_xpath = "//div[3]/input"
        xlsx_file_element = WebDriverWait(driver, 925).until(EC.presence_of_element_located((By.XPATH, xlsx_file_element_xpath)))
        
        xlsx_file_element.send_keys(xlsx_file_path)
        #input(" HERE : xlsx_file_path")

        #xlsx_file_translate_button_xpath = "//div[2]/div[2]/button/span"
        xlsx_file_translate_button_xpath = "//div[2]/div/button/span"
        
        xlsx_file_translate_button = WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.XPATH, xlsx_file_translate_button_xpath)))
        
        print("Clicking on Translate button...")
        #input("BEFORE : xlsx_file_translate_button")
        xlsx_file_translate_button.click()
        #input(" HERE : xlsx_file_translate_button")
        
        # Wait for result text translation page to be loaded
        loop_wait_translation_count = 200
        
         
        res_downloaded_xlsx_translation  = False
        print("Waiting for translation result...")
        while(driver.current_url != 'https://translate.googleusercontent.com/translate_f' and loop_wait_translation_count > 0 and res_downloaded_xlsx_translation == False):
            time.sleep(0.1)
            loop_wait_translation_count = loop_wait_translation_count - 1
                
            if ("https://www.google.com/sorry" in driver.current_url):
                print("We found a CAPTCHA window")
                input("Sorry... found")
                
            download_button_xpath = '//button[normalize-space()="Download translation"]'
            download_button_xpath = '//div[2]/div/button/span'
            download_button_xpath = '//div/button[2]/span[2]'
            
            
            try:
                download_button = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, download_button_xpath)))
                download_button.click()
                print("We found a download button")
                print("Waiting for download to finish")
                #xlsx_tranlated_path = getDownLoadedFileNameChrome(15)
                downloaded_xlsx_translation_path = get_last_downloaded_file_path()
                if len(downloaded_xlsx_translation_path) > 0:
                    print("downloaded_xlsx_translation_path=%s" %(downloaded_xlsx_translation_path))
                    res_downloaded_xlsx_translation = True
                
            except:
                pass
        
        if res_downloaded_xlsx_translation:
            print("Translation xlsx file downloaded at : %s" %(downloaded_xlsx_translation_path))
            #input("Press enter")
        
        #input("After download button")
        #Wait for page status loaded to be complete
        WebDriverWait(driver, 10).until(EC.presence_of_element_located(driver.execute_script('return document.readyState') == 'complete'))
        
        print("Translation received.")
        #input("HERE AGAIN .")
        
        #print("loop_wait_translation_count: %d" % (200 - loop_wait_translation_count))
        #print("URL: %s" % (driver.current_url))
        
        html_translation = driver.page_source
        print("html_translation")
        print(html_translation)
        print("\n________________\nhtml_to_text")
        print("BeautifulSoup to text:")
        soup = BeautifulSoup(html_translation, features="lxml")
        soup = BeautifulSoup(html_translation)
        tdTag = soup.find_all("td")
        translation_array = []
        for td in tdTag:
            pData = td.text
            translation_array.append(td.text)
            #res = soup.get_text()
            print(pData)
        #input("after pData")
        print(translation_array)
        print("__________________________")
        
        # text_translated_document_str = html_translation.replace('<html><head></head><body><pre>', '')
        # text_translated_document_str = text_translated_document_str.replace('</pre></body></html>', '')
        # text_translated_document_str = html.unescape(text_translated_document_str)
        
        #translation_array = text_translated_document_str.split('\n')
        
        text_translated_document_str_nb_lines = len(translation_array)
        
        #print ("text_translated_document_str_nb_linestext_translated_document_str_nb_lines: %s" % text_translated_document_str_nb_lines)
        print ("docxfile_table_number_of_phrases: %s" % docxfile_table_number_of_phrases)
        
        if docxfile_table_number_of_phrases == text_translated_document_str_nb_lines:
            #print("OK, we got the right number of translated lines !")
            pass
        else:
            print("oups ! we got %s translated lines out of %s" % (text_translated_document_str_nb_lines, docxfile_table_number_of_phrases))
            translation_succeded = False
        
        #print("text_translated_document_str:")
        #print(text_translated_document_str)
        
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(8)
    return translation_array



def selenium_chrome_yandex_translate(to_translate):
    try:
        #https://translate.yandex.com/?lang=en-hu
        #driver.get("https://translate.yandex.com/?lang=%s-%s&text=%s" % (src_lang,dest_lang,to_translate))
        driver.get("https://translate.yandex.com/?lang=%s-%s" % (src_lang,dest_lang))
        #(driver.page_source).encode('utf-8')

        timeout_captcha = 200
        alert_captcha = False
        #while 'Unfortunately, it looks like the search request sent from your IP address are automated' in driver.page_source and timeout_captcha > 0:
        while 'IP address' in driver.page_source and timeout_captcha > 0:
            if not alert_captcha:
                print("--------------------------------------------------------------------")
                print("\nCAPTCHA : please fill requested information on browser to continue")
                print("--------------------------------------------------------------------")
            sleep(1)
            timeout_captcha -=1
            alert_captcha = True

        input_element = "#textarea"
        input_button = WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fakeArea")))

        input_button.send_keys (to_translate)

        copy_translation_element = "//span[@id='copyButton']"
        copy_translation_button = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, copy_translation_element)))

        copy_translation_button_class_attribute = copy_translation_button.get_attribute("class")
        #print("val class=%s" % (copy_translation_button_class_attribute))

        sleep(0.2)

        timeout_copybutton_disabled = 60
        while 'state-disabled' in copy_translation_button_class_attribute and timeout_copybutton_disabled > 0:
            sleep(0.2)
            copy_translation_button_class_attribute = copy_translation_button.get_attribute("class")
            #print("val class=%s (%d)" % (copy_translation_button_class_attribute, timeout_copybutton_disabled))
            timeout_copybutton_disabled -=1

        actions = ActionChains(driver)
        driver.set_window_size(800, 700)
        actions.move_to_element(copy_translation_button).perform()
        actions.move_to_element(copy_translation_button).perform()

        sleep(0.1)

        translation_result_element = "translation"
        translation_result_box = WebDriverWait(driver, 60).until(EC.presence_of_element_located((translation_result_element)))
        translation = translation_result_box.text
        res = translation

        translation = res

    except Exception:
        var = traceback.format_exc()
        print(var)
        sys.exit(9)
    return translation

def remove_span_tag(text):  
    search_opening_html_span_tag = r'(?i)<span class="[a-zA-Z]+">'
    search_replace_opening_span = re.compile(search_opening_html_span_tag)
                
    subn_result = search_replace_opening_span.subn("", text)
    subn_count = subn_result[1]
    if subn_count > 0:
        #print ("Replaced '%s' by '%s' %d times." % (search_opening_html_span_tag, "", subn_count))
        text = subn_result[0]
        #if subn_count > 0:
        #    print ("Replaced span %d times" % (subn_count))
            
    search_closing_html_span_tag = r'(?i)</span>'
    search_replace_closing_span = re.compile(search_closing_html_span_tag)
                
    subn_result = search_replace_closing_span.subn("", text)
    subn_count = subn_result[1]
    if subn_count > 0:
        #print ("Replaced '%s' by '%s' %d times." % (search_opening_html_span_tag, "", subn_count))
        text = subn_result[0]
        #if subn_count > 0:
        #    input ("Replaced span %d times" % (subn_count))
                
    return text



def selenium_chrome_deepl_log_in():
    global json_configuration_array, MAX_TRANSLATION_BLOCK_SIZE
    
    deepl_account_email_key = ['deepl', 'account', 'email']
    deepl_account_email = get_nested_value_from_json_array(json_configuration_array, deepl_account_email_key)
    
    deepl_account_password_key = ['deepl', 'account', 'password']
    deepl_account_password = get_nested_value_from_json_array(json_configuration_array, deepl_account_password_key)
        
    deepl_account_enabled_key = ['deepl', 'account', 'enabled']
    deepl_account_enabled = get_nested_value_from_json_array(json_configuration_array, deepl_account_enabled_key)
    
    driver.set_window_size(600, 600)
    #driver.maximize_window()

    try:
        driver.get("https://www.deepl.com/translator")
        
        try:
        
            try:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Accept')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 1).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                driver.execute_script("arguments[0].scrollIntoView();", deepl_accept_cookies_button)    
                deepl_accept_cookies_button.click()
                
            except:
                pass

            # Close the cookies message box if it is there
            try:
                if closed_cookies_accept_message_bool == False:
                    # Accept cookies
                    deepl_accept_cookies_element = "//button[contains(.,'Close')]"
                    deepl_accept_cookies_button = WebDriverWait(driver, 1).until(
                        EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                    deepl_accept_cookies_button.click()
                    closed_cookies_accept_message_bool = True
            except:
                pass
                
            try:
                # close install extension message
                driver.get("https://www.deepl.com/translator")
                deepl_close_deepl_extension_element = ".w-6 > .flex"
                deepl_close_deepl_extension_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, deepl_close_deepl_extension_element)))
                deepl_close_deepl_extension_button.click()
            except:
                pass
        
            # End function if no email or password are provided
            if (deepl_account_email is None) or (deepl_account_email is None):
                return False
            elif deepl_account_enabled == False:
                return False
            
            driver.set_window_position(0, 50)
            driver.set_window_size(800, 700)
            driver.get("https://www.deepl.com/en/login/")

            try:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Accept all cookies')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                driver.execute_script("arguments[0].scrollIntoView();", deepl_accept_cookies_button)    
                deepl_accept_cookies_button.click()
                
            except:
                pass       
            
            # Fill username 
            deepl_login_email_element = "//input[@name='email']"
            deepl_login_email_field = WebDriverWait(driver, 2).until(
                EC.presence_of_element_located((By.XPATH, deepl_login_email_element)))
            deepl_login_email_field.send_keys(deepl_account_email)
            
            # Fill password
            deepl_login_password_element = "//input[@name='password']"
            deepl_login_password_field = WebDriverWait(driver, 1).until(
                EC.presence_of_element_located((By.XPATH, deepl_login_password_element)))
            deepl_login_password_field.send_keys(deepl_account_password)
            sleep(1)

            # Close the cookies message box if it is there
            try:
                if closed_cookies_accept_message_bool == False:
                    # Accept cookies
                    deepl_accept_cookies_element = "//button[contains(.,'Close')]"
                    deepl_accept_cookies_button = WebDriverWait(driver, 0.5).until(
                        EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                    deepl_accept_cookies_button.click()
                    closed_cookies_accept_message_bool = True
            except:
                pass
                
            try:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Accept all cookies')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                driver.execute_script("arguments[0].scrollIntoView();", deepl_accept_cookies_button)    
                deepl_accept_cookies_button.click()
                
            except:
                pass       
            
            # Submit login
            deepl_login_submit_element = "//form/button"
            deepl_login_submit_element = "//input[@name='submit']"
            deepl_login_submit_element = "//button[contains(.,'Log in')]"
            deepl_login_submit_button = WebDriverWait(driver, 3).until(
                EC.presence_of_element_located((By.XPATH, deepl_login_submit_element)))
            driver.execute_script("arguments[0].scrollIntoView();", deepl_login_submit_button)    
            sleep(1.5)
            try:
                deepl_login_submit_button.click()
            except:
                pass
            
            try:
                # Check account button exist
                deepl_login_menu_element = ".dl_header_menu_v2__buttons__opener"
                deepl_login_menu_button = WebDriverWait(driver, 3).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, deepl_login_menu_element)))
                deepl_login_menu_button.click()
                # Close the opener dialog, not required but cleaner
                sleep(0.1)
                deepl_login_menu_button.click()
            except:
                pass
            
            try:
                # Close the annoying plugin for deepl if displayed - bug : it does not find this element
                deepl_plugin_dialog_element = ".w-6 path"
                deepl_plugin_dialog_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, deepl_plugin_dialog_element)))
                deepl_plugin_dialog_button.click()
            except:
                # Just ignore if this plugin dialog does not appear
                pass
            
            # Success change block size if value exists
            deepl_max_char_bloc_size_key = ['deepl', 'account','maximum_character_block']
            deepl_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, deepl_max_char_bloc_size_key)
            
            if isinstance(deepl_maximum_character_block, int):
                if deepl_maximum_character_block > MAX_TRANSLATION_BLOCK_SIZE:
                    MAX_TRANSLATION_BLOCK_SIZE = deepl_maximum_character_block
                    print("\nRobot is now logged in Deepl using %s account." % (deepl_account_email))
                    print("Changing the value of maximum number of characters per block: %s\n" % (MAX_TRANSLATION_BLOCK_SIZE))
                
            return True
            
        except:
            var = traceback.format_exc()
            print(var)
            print("Failed to login into Deepl, continuing without being logged on.")
            driver.set_window_size(800, 700)
            return False

    except:
        var = traceback.format_exc()
        print(var)
        print("Failed to login into Deepl, continuing without being logged on.")
        driver.set_window_size(800, 700)
        return False


def selenium_chrome_perplexity_wait_log_in():
    global json_configuration_array, MAX_TRANSLATION_BLOCK_SIZE
    
    driver.set_window_size(600, 600)
    #driver.maximize_window()

    loop_count = 200
    sleep_wait_sec = 5

    while True:
        try:
            driver.get("https://www.perplexity.ai/")
            
            # Wait up to 10 seconds for the signed-in avatar to appear
            WebDriverWait(driver, 3).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, '[data-testid="sidebar-popover-trigger-signed-in"]'))
            )
            print("✅ User is logged in perplexity.")
            return True

        except:
            var = traceback.format_exc()
            print(var)
            
    return False



def selenium_chrome_deepl_log_off():
    global json_configuration_array, MAX_TRANSLATION_BLOCK_SIZE
        
    driver.set_window_size(800, 700)

    try:
        driver.get("https://www.deepl.com/")
        
        try:
            
            # Open account menu by clicking the account button
            deepl_login_menu_element = ".dl_header_menu_v2__buttons__opener"
            deepl_login_menu_button = WebDriverWait(driver, 9).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, deepl_login_menu_element)))
            deepl_login_menu_button.click()
            
            try:
                # Open account menu by clicking the account button
                deepl_logout_menu_element = "//button[contains(.,'Log out')]"
                deepl_logout_menu_button = WebDriverWait(driver, 1).until(
                    EC.presence_of_element_located((By.XPATH, deepl_logout_menu_element)))
                deepl_logout_menu_button.click()
                print("\nRobot is now logged off Deepl account.")
                
            except:
                # Just ignore if this plugin dialog does not appear
                print("Unable to log off from Deepl, this can be ignored.")
                pass
                
            return True
            
        except:
            var = traceback.format_exc()
            print(var)
            print("Failed of Deepl, this can be ignored")
            return False

    except:
        var = traceback.format_exc()
        print(var)
        print("Failed of Deepl, this can be ignored")
        return False




def selenium_chrome_deepl_translate(to_translate, retry_count):
    global logged_into_deepl
    translation = ""
    Translated = False
    # Progress bar to show only when deepl also shows it on the browser
    bar = None
    global closed_cookies_accept_message_bool, close_install_extension_message_bool, deepl_nb_clear_cached_times
    global engine_method, end_time, elapsed_time, json_configuration_array
    
    deepl_maximum_clear_cache_retry_key = ['deepl', 'maximum_clear_cache_retry']
    deepl_maximum_clear_cache_retry = get_nested_value_from_json_array(json_configuration_array, deepl_maximum_clear_cache_retry_key)
    
    # Set variable to false if they are not globally defined
    try:
        tmp_var = closed_cookies_accept_message_bool
        tmp_var = close_install_extension_message_bool
    except:
        closed_cookies_accept_message_bool = False
        close_install_extension_message_bool = False

    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)

    driver.set_window_size(800, 700)

    try:
        translation_page_openeing_loop_count = 4
        translation_page_opened = False
        
        # Open Deepl translation page
        while translation_page_opened == False and translation_page_openeing_loop_count > 0:
            #print(f"{translation_page_openeing_loop_count} trying left")
            try:
                # driver.get("https://www.deepl.com/translator#%s/%s/%s" % (src_lang,dest_lang, to_translate))
                # Deepl has a bug for / in text to be translated
                # must be replaced by %5C%2F
                driver.get("https://www.deepl.com/translator#%s/%s/%s" % (
                src_lang, dest_lang, urllib.parse.quote(to_translate).replace("%5C", "%5C%5C").replace("/", "%5C%2F").replace("%7C", "%5C%7C")))

                translation_page_opened = True
            except:
                print("Waiting for https://www.deepl.com/ ...")
                sleep(1)
            translation_page_openeing_loop_count = translation_page_openeing_loop_count - 1
            # driver.get("https://www.deepl.com/translator#%s/%s/Hello" % (src_lang,dest_lang))

        # Wait for page to be loaded
        try:
            (driver.page_source).encode('utf-8')
            WebDriverWait(driver, 15).until(lambda driver: driver.execute_script('return document.readyState') == 'complete')
            
            try:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Accept')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 0.01).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                driver.execute_script("arguments[0].scrollIntoView();", deepl_accept_cookies_button)    
                deepl_accept_cookies_button.click()
                
            except:
                pass
            
            try:
                # Accept cookies
                deepl_one_click_navigation_element = "//div[@id='react-joyride-step-0']/div/div/div/div[3]/div/button/span/span"
                deepl_accept_cookies_element = "//button[contains(.,'Got it')]"
                deepl_one_click_navigation_button = WebDriverWait(driver, 0.01).until(
                    EC.presence_of_element_located((By.XPATH, deepl_one_click_navigation_element)))
                driver.execute_script("arguments[0].scrollIntoView();", deepl_one_click_navigation_button)    
                deepl_one_click_navigation_button.click()
            except:
                pass
            
            
            
            #print("Page loaded completed")
        except:
            # print("Waiting for the input_element...")
            var = traceback.format_exc()
            print(var)

        # Close the cookies message box if it is there
        try:
            if closed_cookies_accept_message_bool == False:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Accept all cookies')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 0.02).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                deepl_accept_cookies_button.click()
                closed_cookies_accept_message_bool = True
        except:
            pass
            

        # Close the cookies message box if it is there
        try:
            if closed_cookies_accept_message_bool == False:
                # Accept cookies
                deepl_accept_cookies_element = "//button[contains(.,'Close')]"
                deepl_accept_cookies_button = WebDriverWait(driver, 0.02).until(
                    EC.presence_of_element_located((By.XPATH, deepl_accept_cookies_element)))
                deepl_accept_cookies_button.click()
                closed_cookies_accept_message_bool = True
        except:
            pass
            
        
        # Close the install extension message box if it is there
        try:
            # close install extension message
            if close_install_extension_message_bool == False:
                deepl_close_deepl_extension_element = ".w-6 > .flex"
                deepl_close_deepl_extension_button = WebDriverWait(driver, 0.02).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, deepl_close_deepl_extension_element)))
                deepl_close_deepl_extension_button.click()
                close_install_extension_message_bool = False
        except:
            pass
            
            
        try:
            # Close the translate file button
            deepl_translate_file_dismiss_element = "//button[contains(.,'Dismiss')]"
            deepl_translate_file_dismiss_button = WebDriverWait(driver, 0.01).until(
                EC.presence_of_element_located((By.XPATH, deepl_translate_file_dismiss_element)))
            deepl_translate_file_dismiss_button.click()
                
        except:
            pass       
            
                
        # Wait for copy translation button
        # Removed on 2022-05-25
        found_copy_button = False
        loop_counter_search_button = 4
        while (found_copy_button is False) and (loop_counter_search_button > 0):
            #print(f"loop {loop_counter_search_button}")
            try:
                # Added on 2023-09-26
                copy_translation_element = "//button[contains(@aria-label, 'Copy to clipboard')]" #//svg
                #print(f"Looking for {copy_translation_element}")
                copy_translation_button = WebDriverWait(driver, 0.2).until(
                    EC.presence_of_element_located((By.XPATH, copy_translation_element)))
                
                found_copy_button = True
                #print(f"Loop {loop_counter_search_button}, found xpath button: {copy_translation_element}")
                #print(f"Found xpath button: {copy_translation_element}")
                time.sleep(0.2)
                
            except:
                #print(f"Except loop {loop_counter_search_button}, not found xpath button: {copy_translation_element}")
                try:
                    copy_translation_element = "#dl_translator"
                    #print(f"Looking for {copy_translation_element}")

                    copy_translation_button = WebDriverWait(driver, 0.2).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, copy_translation_element)))
                                                                                                                                      
                    found_copy_button = True
                    #print(f"Found xpath button: {copy_translation_element}")
                    
                except:
                    try:
                        # Version 2022-03-09
                        copy_translation_element = ".lmt__target_toolbar_right > span path:nth-child(2)"
                        copy_translation_element = "div:nth-child(5) > svg"
                        #print(f"Looking for {copy_translation_element}")
                        copy_translation_button = WebDriverWait(driver, 0.2).until(
                            EC.presence_of_element_located((By.XPATH, copy_translation_element)))
                        found_copy_button = True
                    except:
                        # Version 2022-03-30
                        try:
                           copy_translation_element = ".lmt__target_toolbar_right > div > span svg"
                           #print(f"Looking for {copy_translation_element}")       
                           copy_translation_button = WebDriverWait(driver, 0.2).until(
                               EC.presence_of_element_located((By.CSS_SELECTOR, copy_translation_element)))
                           found_copy_button = True
                        except:
                           #print("Copy button not found !!")
                           pass
            #print("Incrementing loop_counter_search_button")
            loop_counter_search_button = loop_counter_search_button - 1
        
        busy_element = ".lmt__textarea_separator__border_inner"
        # busy_element = "//div[@id='dl_translator']/div/div/div[5]"
        sleep(deepl_sleep_wait_translation_seconds)

        busybox_innerhtml = ""
        timeout_busy_translating = 50
        try:
            busybox = WebDriverWait(driver, 0.3).until(EC.presence_of_element_located((By.CSS_SELECTOR, busy_element)))
            attrs = driver.execute_script(
                'var items = {}; for (index = 0; index < arguments[0].attributes.length; ++index) { items[arguments[0].attributes[index].name] = arguments[0].attributes[index].value }; return items;',
                busybox)
            busybox_innerhtml = busybox.get_attribute('innerHTML')
            while busybox_innerhtml != "" and timeout_busy_translating > 0:
                sleep(0.3)
                busybox = WebDriverWait(driver, 15).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, busy_element)))
                busybox_innerhtml = busybox.get_attribute('innerHTML')
                attrs = driver.execute_script(
                    'var items = {}; for (index = 0; index < arguments[0].attributes.length; ++index) { items[arguments[0].attributes[index].name] = arguments[0].attributes[index].value }; return items;',
                    busybox)
                timeout_busy_translating -= 1

                deepl_usage_limit_reached_element = "//button[contains(.,'Back to Translator')]"
                try:
                    deepl_usage_limit_reached_button = WebDriverWait(driver, 0.05).until(
                        EC.presence_of_element_located((By.XPATH, deepl_usage_limit_reached_element)))
                    deepl_usage_limit_reached_button.click()
                    return False, ""
                except:
                    pass
        except:
            #var = traceback.format_exc()
            #print(var)
            limit_reached = False

            # Look for usage limit reached, and try pro for 30 days
            deepl_usage_limit_reached_element = "//button[contains(.,'Back to Translator')]"
            try:
                deepl_usage_limit_reached_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.XPATH, deepl_usage_limit_reached_element)))
                
                limit_reached = True
                #deepl_usage_limit_reached_button.click()
            except:
                
                pass
            # Sometimes the busy element does not show up, just ignore it and continue
            
            

            # Look for usage limit reached, and try pro for 30 days
            deepl_try_pro_free_element = "//button[contains(.,'Try Pro for free')]"
            try:
                deepl_try_pro_free_button = WebDriverWait(driver, 0.05).until(
                    EC.presence_of_element_located((By.XPATH, deepl_try_pro_free_element)))
                print("Deepl is promoting their Pro version")
                limit_reached = True
                #deepl_usage_limit_reached_button.click()
            except:
                
                pass
            # Sometimes the busy element does not show up, just ignore it and continue
            
            if limit_reached:
                try:
                    if deepl_nb_clear_cached_times is None:
                        deepl_nb_clear_cached_times = 0
                except:
                    deepl_nb_clear_cached_times = 0
                    
                if deepl_nb_clear_cached_times > deepl_maximum_clear_cache_retry:
                    return False, ""
                print("Warning : deepl usage limit reached... retrying after cleaning cache.")
                driver.delete_all_cookies()
                driver.get("https://www.deepl.com")
                closed_cookies_accept_message_bool = False
                deepl_nb_clear_cached_times = deepl_nb_clear_cached_times + 1
                logged_into_deepl = selenium_chrome_deepl_log_in()
                return selenium_chrome_deepl_translate(to_translate, retry_count)
                

        #print("Scroll to copy_translation_button")
        actions = ActionChains(driver)
        # actions.move_to_element(copy_translation_button).perform()
        # sleep(0.1)

        # Scroll the browser to the element's Y position
        try:
            driver.execute_script("arguments[0].scrollIntoView();", copy_translation_button)
        except:
            pass

        copy_button_clicked = False
        copy_button_clicked_loop_count = 5
        res = ""
        still_translating_html_str = 'div class="lmt__progress_popup lmt__progress_popup--visible lmt__progress_popup--visible_2" dl-test="translator-progress-popup"'
        
        # When failing to get translation from HTML, use button copy and clipboard and warn user.
        warned_using_clipboard = False
        
        while copy_button_clicked_loop_count > 0 and (res == "" or res is None):
            #print(f"copy_button_clicked_loop_count : {copy_button_clicked_loop_count}")
            try:
                driver.execute_script("scrollBy(0,-1000);")
                # clipboard.copy('')
                try:
                    actions.move_to_element(copy_translation_button).perform()
                except:
                    pass
                sleep(0.2)
                # driver.set_window_size(800, 700)
                page_source_str = driver.page_source
                # print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                # with open('before.html', 'w', encoding="utf-8") as f:
                #    f.write(page_source_str)
                # f.close()
                wait_translation_finish_try = 400
                block_translation_percent_done = 0
                while page_source_str.find(still_translating_html_str) > 0 and wait_translation_finish_try > 0:
                    sleep(0.2)
                    #print("Still translating...")
                    page_source_str = driver.page_source
                    # print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                    # print(driver.page_source)
                    wait_translation_finish_try = wait_translation_finish_try - 1
                    search_percent_re = "of characters translated\">(\d+)\% of characters translated</p>"
                    mo = re.search(search_percent_re, page_source_str)
                    if mo:
                        try:
                            if bar is None:
                                bar = progressbar.ProgressBar().start()
                                bar.maxval = 100

                            block_previous_translation_percent_done = block_translation_percent_done
                            block_translation_percent_done = mo.group(1)
                            if block_previous_translation_percent_done != block_translation_percent_done:
                                # print ("found percent: %s" %block_translation_percent_done)
                                bar.update(int(block_translation_percent_done))
                        except:
                            pass

                    # input("of characters translated")

                if bar is not None:
                    bar.update(100)
                    bar = None
                    print("")

                # print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                # print(driver.page_source)
                # print(":::::::::::::::::::::::::::::::::::::::::::::::::::::::")
                # input("enter to click on button")

                page_source_str = driver.page_source
                # with open('after.html', 'w', encoding="utf-8") as f2:
                #    f2.write(page_source_str)
                # f2.close()
                # input("wait html")

                # print("Done waiting for translation")

                try:
                    # Try to get the translation from the innerhtml of translation button
                    inner_html_plain_text_element = "//button[@class='lmt__translations_as_text__text_btn']"
                    InnerHTMLPlainTextElement = WebDriverWait(driver, 1).until(
                        EC.presence_of_element_located((By.XPATH, inner_html_plain_text_element)))
                    translation_from_plain_text = InnerHTMLPlainTextElement.get_attribute('innerHTML')
                    res = translation_from_plain_text
                except:
                    # if we cannot find translation button with translation the use the copy button
                    # previous_clipbboard = clipboard.paste()
                    # previous_clipbboard = pyperclip.paste()
                    page_source_str = driver.page_source
                    #with open('deepl_page_source.html', 'w', encoding="utf-8") as f:
                    #    f.write(page_source_str)
                    #    f.close()
                    res = ""
                    try:
                        try:
                            #inner_html_translation_xpath_element = '//div[@contenteditable="true" and @role="textbox" and @aria-labelledby="translation-results-heading"]'
                            inner_html_translation_xpath_element = "//div[contains(@aria-labelledby, 'translation-target-heading')]"
                            InnerHTMLTranslationElement = WebDriverWait(driver, 1).until(
                                EC.presence_of_element_located((By.XPATH, inner_html_translation_xpath_element)))
                            
                            if InnerHTMLTranslationElement:
                                # Get the plain text from the element
                                translation_from_plain_text = InnerHTMLTranslationElement.text
                                #print("Plain Text: %s " % (translation_from_plain_text))
                            else:
                                print("Element not found")
                            res = translation_from_plain_text
                        except:
                            var = traceback.format_exc()
                            print(var)
                    
                        # Added on version 2022-05-31
                        #copy_translation_element = '//*[@id="headlessui-tabs-panel-7"]/div/div[1]/section/div/div[2]/div[3]/section/div[2]/div[3]/span[2]/span/span/button'
                        #copy_translation_button = WebDriverWait(driver, 6).until(
                        #    EC.presence_of_element_located((By.XPATH, copy_translation_element)))
                        if not warned_using_clipboard and (res == "" or res == None):
                            print("Warning: Failed to get translation from html, copying from clipboard")
                            warned_using_clipboard = True
                            
                        if warned_using_clipboard and (res == "" or res == None):
                            #return False, None
                            clipboard.copy('')
                            copy_translation_button.click()
                            copy_button_clicked = True
                            res = clipboard.paste()
                            if len(res) == 0 or res == None:
                                print("Error : failed to get translation from Deepl.")
                                return False, ""
                            
                    except:
                        var = traceback.format_exc()
                        print(var)
                        #print("res : %s" %(res))
                        pass
                    #return False, None
                    # res = pyperclip.paste()
                    # print(res)

                # id="target-dummydiv"
                # contains the translation
                res = res.replace("\r", "")
                res = remove_span_tag(res)
                
                input_nb_lines = len(to_translate.replace("\r", "").split("\n"))

                translated_phrases_array = res.split("\n")
                translated_phrases_array_len = len(translated_phrases_array)
                
                translated_phrases_array = translated_phrases_array[:input_nb_lines]
                
                # for pos_remove in range(0,translated_phrases_array_len - to_translate_phrases_array_len):
                if translated_phrases_array_len >= to_translate_phrases_array_len:
                    translated_phrases_array = translated_phrases_array[:input_nb_lines]
                    #print("input_nb_lines: %s" % (input_nb_lines))
                    #print("array: %s" % (translated_phrases_array))
                    res = "\n".join(translated_phrases_array)
                    
                if translated_phrases_array_len < to_translate_phrases_array_len:
                    res = ""

            except:
                #print(f"Found exception on loop {copy_button_clicked_loop_count}")
                if copy_button_clicked_loop_count < 20:
                    print("Waiting for the copy button...")
                    #var = traceback.format_exc()
                    #print(var)
            copy_button_clicked_loop_count = copy_button_clicked_loop_count - 1

        # translation = res
        translation = "\n".join(translated_phrases_array)
        # print("translation=%s" % (translation))
        # input("Press enter to continue")
    except Exception:
        var = traceback.format_exc()
        print(var)
        sleep(1)
        # sys.exit(0)
    return True, translation



def selenium_chrome_chatgpt_translate(to_translate, retry_count):
    global logged_into_chatgpt, src_lang_name, dest_lang_name
    
    translation = ""
    Translated = False
    # Progress bar to show only when deepl also shows it on the browser
    bar = None
    global closed_cookies_accept_message_bool, close_install_extension_message_bool, deepl_nb_clear_cached_times
    global engine_method, end_time, elapsed_time, json_configuration_array
    
    deepl_maximum_clear_cache_retry_key = ['deepl', 'maximum_clear_cache_retry']
    deepl_maximum_clear_cache_retry = get_nested_value_from_json_array(json_configuration_array, deepl_maximum_clear_cache_retry_key)
    
    # Set variable to false if they are not globally defined
    try:
        tmp_var = closed_cookies_accept_message_bool
        tmp_var = close_install_extension_message_bool
    except:
        closed_cookies_accept_message_bool = False
        close_install_extension_message_bool = False

    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)

    driver.set_window_size(800, 700)

    try:
        translation_page_openeing_loop_count = 4
        translation_page_opened = False
        
        # Open ChatGPT translation page
        while translation_page_opened == False and translation_page_openeing_loop_count > 0:
            #print(f"{translation_page_openeing_loop_count} trying left")
            try:
                driver.get("https://chatgpt.com/")
                #sleep(1)

                translation_page_opened = True
            except:
                print("Waiting for https://chatgpt.com/ ...")
                sleep(1)
            translation_page_openeing_loop_count = translation_page_openeing_loop_count - 1

        sleep(1)
        # Locate the contenteditable div
        textarea = driver.find_element(By.XPATH, "//div[@id='prompt-textarea']")

        # Send text to the element
        textarea.click()

        #time.sleep(5)
        #textarea.send_keys("Translate this from English to Persian:")

        # Sending a new line using Keys.RETURN for proper formatting
        #textarea.send_keys(Keys.SHIFT + Keys.ENTER)

        #textarea.send_keys("I was lying with my eyes closed,")
        #textarea.send_keys(Keys.SHIFT + Keys.ENTER)


        # The string that needs to be sent
        # Max 4,096 characters 
        str_prompt = f"""Translate the following text from {src_lang_name} to {dest_lang_name} for Supreme Master Television subtitles:
Each input line must correspond to exactly one output line.
Do not split, merge, or add any lines.
Do not insert any line breaks within a line, even if the line is long.
There should be no formating including URLs and spacing, within each line.
Only use a line break to move to the next input line.
Do not add, remove, or split any lines.
Each output line must contain the full translation of the corresponding input line.

Here is the text to be translated:
"""

        str_prompt = str_prompt + to_translate
        
        lines = str_prompt.split('\n')

        # Split the string on new lines
        lines = str_prompt.splitlines()

        # Wrap each line in <p>...</p>
        wrapped_lines = [f"<p>{line}</p>" for line in lines]

        # Join all wrapped lines into a single string
        output_string = "".join(wrapped_lines)

        # JavaScript to set the content of the contenteditable div
        js_script = """
        var textarea = document.getElementById('prompt-textarea');
        textarea.innerHTML = arguments[0];
        """

        # Execute JavaScript to inject the text into the div
        driver.execute_script(js_script, output_string)

        print("Test 1 completed")

        # Send each line to the textarea
        #for i, line in enumerate(lines):
        #    textarea.send_keys(line)  # Send the current line
            
        #    # If it's not the last line, send SHIFT + ENTER to move to the next line
        #    if i < len(lines) - 1:
        #        textarea.send_keys(Keys.SHIFT + Keys.RETURN)

        # Wait for the button to appear with a timeout of 3 seconds
        sleep(1)
        #button = WebDriverWait(driver, 3).until(
        #    EC.presence_of_element_located((By.XPATH, "//button[@aria-label='Send prompt' and @data-testid='send-button']"))
        #)
        
        #button = WebDriverWait(driver, 3).until(
        #    EC.presence_of_element_located((By.CSS_SELECTOR, "#composer-submit-button > svg.icon > path"))
        #)


        try:
            button = driver.find_element(By.CSS_SELECTOR, 'button[data-testid="close-button"]')
            button.click()
        except:
            pass

        # Locate the button element using its attributes
        button_submit_prompt = driver.find_element(By.ID, "composer-submit-button")
        
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", button_submit_prompt)

        # Click the button
        button_submit_prompt.click()


        # Set a timeout value for waiting for the element
        timeout = 1  # Timeout after 10 seconds if not found

        while True:
            try:
                # Search for the button with aria-label="Stop streaming"
                stop_button = WebDriverWait(driver, timeout).until(
                    EC.presence_of_element_located((By.XPATH, "//button[@aria-label='Stop streaming']"))
                )
                
                # Element found, perform action (if any)
                print("Found the 'Stop streaming' button!")
                
                # Sleep for 0.5 seconds before checking again
                time.sleep(0.25)
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                
            except Exception as e:
                # If the element is no longer found or any other exception occurs
                print("Element not found or timeout reached. Stopping the loop.")
                
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                break

        # Wait for the button to appear with a timeout of 3 seconds
        #button = WebDriverWait(driver, 3).until(
        #    EC.presence_of_element_located((By.XPATH, "//button[@aria-label='Dictate button']"))
        #)

        page_source_str = driver.page_source


        # Parse the HTML with BeautifulSoup
        soup = BeautifulSoup(page_source_str, 'html.parser')

        # Find all the article tags
        articles = soup.find_all('article')

        #print(len(articles))

        second_article_html = str(articles[1])
        #print (second_article_html)
        #print()


        # Get the text of the last article element
        last_article_text = articles[-1].get_text()

        # Print the extracted text
        #print(last_article_text)

        lines = None

        # Find the div with class "markdown"
        markdown_div = articles[1].find('div', class_='markdown')

        # Check if the div exists and then process the text
        if markdown_div:
            # Step 1: Replace all </p><p> with <br/>
            html_str = str(markdown_div)
            html_str = html_str.replace('</p><p>', '<br/><br/>')

            # Reparse the modified HTML to a BeautifulSoup object again
            markdown_div = BeautifulSoup(html_str, 'html.parser')

            # Step 2: Define a complex delimiter for <br/>
            delimiter = 'random_complex_delimiter_123456'
            delimiter_paragraph = f"<p>{delimiter}</p>"

            # Step 3: Replace <br/> tags with the complex delimiter
            for line_break in markdown_div.find_all('br'):
                line_break.insert_before(BeautifulSoup(delimiter_paragraph, 'html.parser'))
                line_break.unwrap()  # Remove the <br> tag after inserting the delimiter
                

            # Get the full text with the complex delimiter and print it
            markdown_text_with_delimiter = markdown_div.get_text()

            # Output the result
            #print(markdown_text_with_delimiter)
            #input("After markdown text split")
            
            lines = markdown_text_with_delimiter.split(delimiter)
            if(len(lines) == 1):
                lines = markdown_text_with_delimiter.split("\n")
            print(lines)
            print("after print lines")

        else:
            print("No div with class 'markdown' found.")

        translated_phrases_array = lines
        translated_phrases_array_len = len(translated_phrases_array)
        
        input_nb_lines = len(to_translate.replace("\r", "").split("\n"))
        # for pos_remove in range(0,translated_phrases_array_len - to_translate_phrases_array_len):
        if translated_phrases_array_len >= to_translate_phrases_array_len:
            print(f"input_nb_lines={input_nb_lines}")
            translated_phrases_array = translated_phrases_array[:input_nb_lines]
            #print("input_nb_lines: %s" % (input_nb_lines))
            #print("array: %s" % (translated_phrases_array))
            res = "\n".join(translated_phrases_array)
            if translated_phrases_array_len > to_translate_phrases_array_len + 1:
                print("Found %s lines out of %s lines" % (translated_phrases_array_len, to_translate_phrases_array_len))

        if translated_phrases_array_len < to_translate_phrases_array_len:
            res = ""
            print("Error, not enough lines")
            print("Cleaning up perplexity cookies...")
            driver.delete_all_cookies()
            sleep(1)

        translation = "\n".join(lines)
        #input("After markdown text")
        # Get the text inside this div
        #if assistant_div:
        #    assistant_text = assistant_div.get_text()
        #    print(assistant_text)
        #else:
        #    print("No div with data-message-author-role='assistant' found.")
        
        # Step 1: Click the 3-dot conversation options button
        menu_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, 'button[data-testid="conversation-options-button"]'))
        )
        menu_button.click()

        # Step 2: Wait for and click the "Delete" button by visible text
        delete_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//*/text()[normalize-space(.)='Delete']/parent::*"))
        )
        delete_button.click()
        
        # Step: Click the red "Delete" confirmation button
        confirm_delete_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, 'button[data-testid="delete-conversation-confirm-button"]'))
        )
        confirm_delete_button.click()
        
        
            
    except Exception:
        var = traceback.format_exc()
        print(var)
        sleep(1)
        # sys.exit(0)
    return True, translation



def selenium_chrome_perplexity_translate(to_translate, retry_count):
    global logged_into_chatgpt, src_lang_name, dest_lang_name, chrome_options, bloc_number, service, chrome_options
    
    translation = ""
    Translated = False
    # Progress bar to show only when deepl also shows it on the browser
    bar = None
    global closed_cookies_accept_message_bool, close_install_extension_message_bool, deepl_nb_clear_cached_times
    global engine_method, end_time, elapsed_time, json_configuration_array
    
    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)


    str_prompt = f"""Translate the following text from {src_lang_name} to {dest_lang_name} for Supreme Master Television subtitles:
Each input line must correspond to exactly one output line.
Do not split, merge, or add any lines.
Do not insert any line breaks within a line, even if the line is long.
Only use a line break to move to the next input line.
Do not add, remove, or split any lines.
If a phrase is on multiple lines, it must remain on multiple lines, no merge.
Do not echo text to be translated in the translation, and do not insert an introduction before the translation:
Each output line must contain the full translation of the corresponding input line. The text has {to_translate_phrases_array_len} lines that must be translated in exactly {to_translate_phrases_array_len} lines.
Your output MUST contain exactly {to_translate_phrases_array_len} lines, not one less, not one more.

The text to be translated start after the first line containing only BEFORETEXTTOTRANSLATE and ends the line before the first occurence if the line containing only AFTERTEXTTOTRANSLATE:
BEFORETEXTTOTRANSLATE
{to_translate}
AFTERTEXTTOTRANSLATE"""
    
    #print(str_prompt)
    
    # Set variable to false if they are not globally defined
    try:
        tmp_var = closed_cookies_accept_message_bool
        tmp_var = close_install_extension_message_bool
    except:
        closed_cookies_accept_message_bool = False
        close_install_extension_message_bool = False

    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)

    driver.set_window_size(800, 700)

    try:
        translation_page_openeing_loop_count = 4
        translation_page_opened = False
        
        # Open ChatGPT translation page
        while translation_page_opened == False and translation_page_openeing_loop_count > 0:
            #print(f"{translation_page_openeing_loop_count} trying left")
            try:
                driver.get("https://www.perplexity.ai/")
                #sleep(1)

                translation_page_opened = True
            except:
                print("Waiting for https://www.perplexity.ai/ ...")
                sleep(1)
            translation_page_openeing_loop_count = translation_page_openeing_loop_count - 1


        # Locate the contenteditable div
        textarea = WebDriverWait(driver, 1).until(
            EC.presence_of_element_located((By.XPATH, "//*[@id='ask-input']"))
        )

        # Send text to the element
        textarea.click()
        
        # Assuming you already have a WebDriver instance (driver)
        #textarea = driver.find_element(By.ID, "ask-input")

        js_script = f"""
        const textarea = document.getElementById('ask-input');

        // Create a clipboard event with the desired text
        const clipboardData = new DataTransfer();
        clipboardData.setData('text/plain', `{str_prompt}`);
        const pasteEvent = new ClipboardEvent('paste', {{
          bubbles: true,
          cancelable: true,
          clipboardData: clipboardData
        }});

        // Focus and dispatch paste
        textarea.focus();
        textarea.dispatchEvent(pasteEvent);
        """

        driver.execute_script(js_script)
        
        
        #Click accept all cookies if found
        try:
            accept_all_cookies = WebDriverWait(driver, 0.1).until(
                EC.presence_of_element_located((By.XPATH, "//*/text()[normalize-space(.)='Accept All Cookies']/parent::*"))
            )
        except TimeoutException:
            accept_all_cookies = None  # Element not found, continue program

        # Now you can check:
        if accept_all_cookies:
            accept_all_cookies.click()
        
        #Close login dialog if found
        try:
            close_signin_button = WebDriverWait(driver, 0.1).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button[data-testid='floating-signup-close-button']"))
            )
        except TimeoutException:
            close_signin_button = None  # Element not found, continue program

        # Now you can check:
        if close_signin_button:
            close_signin_button.click()
        

        submit_button = WebDriverWait(driver, 1).until(
            EC.presence_of_element_located((By.XPATH, '//button[@data-testid="submit-button"]'))
        )
        submit_button.click()

        time.sleep(1)

        timeout = 300  # seconds
        poll_interval = 1  # seconds
        start_time = time.time()

       # print("⏳ Waiting for stop button to disappear", end='')
        while True:
            try:
                stop_button = driver.find_element(By.CSS_SELECTOR,  '[data-testid="stop-generating-response-button"]')
                if stop_button:
                    try:
                        if stop_button.is_displayed():
                            #print("⏳ Waiting for stop button to disappear...")
                            #print('.', end='')
                                        
                            # Sleep for 0.5 seconds before checking again
                            time.sleep(0.25)
                            # Locate the div by its class
                            try:
                                prose_div = driver.find_element(By.CSS_SELECTOR, "div.prose")
                                prose_div.click()
                                prose_div.send_keys(Keys.PAGE_DOWN)
                                
                                body = driver.find_element(By.TAG_NAME, "body")
                                body.send_keys(Keys.PAGE_DOWN)
                                
                            except:
                                try:
                                    body = driver.find_element(By.TAG_NAME, "body")
                                    body.send_keys(Keys.PAGE_DOWN)
                                except:
                                    print("Cannot find html body...")
                                    pass
                                pass
                            #driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                            
                            pass
                        else:
                            #print("\n✅ Stop button is no longer visible.")
                            #print("\n")
                            time.sleep(1)
                            break
                    except:
                        break
                
            except NoSuchElementException:
                #print("✅ Stop button has been removed from the DOM.")
                #print("\n")
                break

            # Timeout check
            if time.time() - start_time > timeout:
                print("⚠️ Timed out waiting for stop button to disappear.")
                break

            time.sleep(poll_interval)
            
        time.sleep(1)

        input_nb_lines = len(to_translate.replace("\r", "").split("\n"))

        
        # Get the div with class "prose"
        try:
            prose_div = WebDriverWait(driver, 2.5).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "div.prose"))
            )
        except:
            pass

        try:
            prose_div = driver.find_element(By.CSS_SELECTOR, "div.prose")
            prose_div.click()
            prose_div.send_keys(Keys.PAGE_DOWN)
            
        except:
            try:
                body = driver.find_element(By.TAG_NAME, "body")
                body.send_keys(Keys.PAGE_DOWN)
            except:
                print("Cannot find html body...")
                pass
                                
        # Try to get the div again (big)
        time.sleep(0.25)
        try:
            prose_div = driver.find_element(By.CSS_SELECTOR, "div.prose")
        except:
            pass
        
        # Extract all visible text content
        text = prose_div.text

        # Split into lines and strip empty ones
        result_lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        translated_phrases_array = result_lines
        translated_phrases_array_len = len(translated_phrases_array)
        
        #print("result_lines:")
        #print(result_lines)
        
        res = None
                

        # for pos_remove in range(0,translated_phrases_array_len - to_translate_phrases_array_len):
        if translated_phrases_array_len >= to_translate_phrases_array_len:
            #print(f"input_nb_lines={input_nb_lines}")
            translated_phrases_array = translated_phrases_array[:input_nb_lines]
            #print("input_nb_lines: %s" % (input_nb_lines))
            #print("array: %s" % (translated_phrases_array))
            res = "\n".join(translated_phrases_array)
            if translated_phrases_array_len > to_translate_phrases_array_len + 1:
                print("Found %s lines out of %s lines" % (translated_phrases_array_len, to_translate_phrases_array_len))

        if translated_phrases_array_len < to_translate_phrases_array_len:
            res = ""
            print(f"Error, not enough lines : {translated_phrases_array_len} out of {to_translate_phrases_array_len}")
            sleep(0.3)

        #print(res)
        
        ##################################################
        # Delete this chat from perplexity AI history
        wait = WebDriverWait(driver, 10)

        # 1. Click the three-dot (⋯) menu icon
        dots_button = wait.until(EC.element_to_be_clickable((
            By.CSS_SELECTOR,
            'svg.tabler-icon-dots'
        )))
        dots_button.click()

        # 2. Click the Delete option (with trash icon and text "Delete")
        delete_button = wait.until(EC.element_to_be_clickable((
            By.XPATH,
            '//div[contains(@class, "cursor-pointer")]//span[text()="Delete"]'
        )))
        delete_button.click()

        # Wait for the Confirm button and click it
        confirm_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, 'button[data-testid="thread-delete-confirm"]'))
        )
        confirm_button.click()
        
        translation = res
            
    except Exception:
        var = traceback.format_exc()
        print(var)
        sleep(1)
        # sys.exit(0)
    return True, translation

total_cost = 0

def perplexity_api_translate(to_translate, retry_count):
    global src_lang_name, dest_lang_name, total_cost
    
    translation = ""
    Translated = False
    # Progress bar to show only when deepl also shows it on the browser
    bar = None
    global closed_cookies_accept_message_bool, close_install_extension_message_bool, deepl_nb_clear_cached_times
    global engine_method, end_time, elapsed_time, json_configuration_array
    
    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)


    str_prompt = f"""Translate the following text from {src_lang_name} to {dest_lang_name} for Supreme Master Television subtitles:
Each input line must correspond to exactly one output line.
Do not split, merge, or add any lines.
Do not insert any line breaks within a line, even if the line is long.
Only use a line break to move to the next input line.
Do not add, remove, or split any lines.
If a phrase is on multiple lines, it must remain on multiple lines, no merge.
Do not echo text to be translated in the translation, and do not insert an introduction before the translation:
Each output line must contain the full translation of the corresponding input line. The text has {to_translate_phrases_array_len} lines that must be translated in exactly {to_translate_phrases_array_len} lines.
Your output MUST contain exactly {to_translate_phrases_array_len} lines, not one less, not one more.

The text to be translated start after the first line containing only BEFORETEXTTOTRANSLATE and ends the line before the first occurence if the line containing only AFTERTEXTTOTRANSLATE:
BEFORETEXTTOTRANSLATE
{to_translate}
AFTERTEXTTOTRANSLATE"""
    
    #print(str_prompt)

    to_translate_phrases_array = to_translate.split("\n")
    to_translate_phrases_array_len = len(to_translate_phrases_array)

    try:
        translation_page_openeing_loop_count = 4

        input_nb_lines = len(to_translate.replace("\r", "").split("\n"))
        
        response = requests.post(
            'https://api.perplexity.ai/chat/completions',
            headers={
                'Authorization': 'Bearer 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',
                'Content-Type': 'application/json'
            },
            json={
                'model': 'sonar',
                'messages': [
                    {
                        'role': 'user',
                        'content': str_prompt
                    }
                ]
            }
        )

        # Pretty-print the full JSON response
        try:
            print(json.dumps(response.json(), indent=2))
        except json.JSONDecodeError:
            print("Response is not JSON:", response.text)
        
        # Parse the JSON string
        data = response.json()

        # Get original content
        original_content = data["choices"][0]["message"]["content"]

        # Remove all whitespace before newlines
        modified_content = re.sub(r'[ \t]+\n', '\n', original_content)

        # Print the "content" value
        print("Content:")
        print(modified_content)

        print("\nToken and cost values:")

        # Print the specified usage and cost values line by line
        print("prompt_tokens:", data["usage"]["prompt_tokens"])
        print("completion_tokens:", data["usage"]["completion_tokens"])
        print("total_tokens:", data["usage"]["total_tokens"])
        print("input_tokens_cost:", data["usage"]["cost"]["input_tokens_cost"])
        print("output_tokens_cost:", data["usage"]["cost"]["output_tokens_cost"])
        print("total_cost:", data["usage"]["cost"]["total_cost"])
        total_cost = total_cost + data["usage"]["cost"]["total_cost"]


        # Split into lines and strip empty ones
        result_lines = [line.strip() for line in modified_content.splitlines() if line.strip()]
        
        translated_phrases_array = result_lines
        translated_phrases_array_len = len(translated_phrases_array)
        
        #print("result_lines:")
        #print(result_lines)
        
        res = None

        # for pos_remove in range(0,translated_phrases_array_len - to_translate_phrases_array_len):
        if translated_phrases_array_len >= to_translate_phrases_array_len:
            #print(f"input_nb_lines={input_nb_lines}")
            translated_phrases_array = translated_phrases_array[:input_nb_lines]
            #print("input_nb_lines: %s" % (input_nb_lines))
            #print("array: %s" % (translated_phrases_array))
            res = "\n".join(translated_phrases_array)
            if translated_phrases_array_len > to_translate_phrases_array_len + 1:
                input("Found %s lines out of %s lines" % (translated_phrases_array_len, to_translate_phrases_array_len))

        if translated_phrases_array_len < to_translate_phrases_array_len:
            res = ""
            print(f"Error, not enough lines : {translated_phrases_array_len} out of {to_translate_phrases_array_len}")
            #input(str_prompt)
            sleep(3)

        #print("Translation:")
        #print(res)

            
    except Exception:
        var = traceback.format_exc()
        print(var)
        sleep(1)
        # sys.exit(0)
    return True, res



def set_translation_function():
    global selenium_chrome_machine_translate_once
    if not splitonly:
        print("\ntranslation_engine=%s" % (translation_engine))
        print("engine_method=%s" % (engine_method))
        if (engine_method == "phrasesblock"):
            print("maximum number of characters per block: %d" % MAX_TRANSLATION_BLOCK_SIZE)

    if translation_engine == 'yandex':
        print("Using translation_engine=%s" % (translation_engine))
        selenium_chrome_machine_translate_once = selenium_chrome_yandex_translate
    elif translation_engine == 'deepl':
        if engine_method == 'phrasesblock':
            selenium_chrome_machine_translate_once = selenium_chrome_translate_get_from_text_array
        else:
            selenium_chrome_machine_translate_once = selenium_chrome_deepl_translate
    elif translation_engine == 'deepl':
        if engine_method == 'api':
            selenium_chrome_machine_translate_once = perplexity_api_translate 
        else:
            selenium_chrome_machine_translate_once = selenium_chrome_deepl_translate
    else:
        if engine_method == 'textfile':
            selenium_chrome_machine_translate_once = selenium_chrome_translate_get_from_text_array
        elif engine_method == 'singlephrase':
            selenium_chrome_machine_translate_once = selenium_chrome_google_translate
        else:
            selenium_chrome_machine_translate_once = selenium_chrome_translate_get_from_text_array


def selenium_chrome_machine_translate(to_translate, index):
    global selenium_chrome_machine_translate_once
    translation = ""
    translation_try_count = 1
    max_try_count = 15
    global translation_errors_count
    global deepl_sleep_wait_translation_seconds
    #print("--index-- : %d" % index)
    #to_translate_str = str(to_translate)
    try:
        while translation_try_count < max_try_count and translation == "":
            if translation_try_count > 1:
                print("Retrying to translate again (%d)..." % (translation_try_count))
                translation_errors_count = translation_errors_count + 1
                deepl_sleep_wait_translation_seconds  = deepl_sleep_wait_translation_seconds * 1.1
                print("%d translation retry so far..." % (translation_errors_count))
            if translation_engine == 'deepl':
                if engine_method == 'phrasesblock':
                    translation = selenium_chrome_machine_translate_once(to_translate, index)
                else:
                    translation = selenium_chrome_machine_translate_once(to_translate, translation_try_count - 1)
            elif engine_method == 'textfile':
                translation = selenium_chrome_machine_translate_once(to_translate, index)
            elif engine_method == 'xlsxfile':
                translation = selenium_chrome_machine_translate_once(to_translate, index)
            elif engine_method == 'phrasesblock':
                translation = selenium_chrome_machine_translate_once(to_translate, index)
            elif engine_method == 'javascript':
                translation = selenium_chrome_machine_translate_once(to_translate, index)
            else:            
                translation = selenium_chrome_machine_translate_once(to_translate)
            translation_try_count = translation_try_count + 1
    except:
        print("Error in selenium_chrome_machine_translate")
    return translation
    
def initialize_translation_memory_xlsx():
    global xtm
    # If --xlsxreplacefile was provided in the command line
    if xlsxreplacefile is not None:
        print("xlsxreplacefile: %s" % (xlsxreplacefile))
        xtm = xlsx_translation_memory.xlsx_translation_memory(xlsxreplacefile)
        print("")
    else:
        xtm = xlsx_translation_memory.xlsx_translation_memory(None)


def is_end_of_line(line):
    for eol in eol_array:
        #print("Testing is_end_of_line '%s' on string '%s'" %(eol, line))
        if re.search(eol, line):
            #print("Found is_end_of_line '%s' on string '%s'" %(eol, line))
            return 1
    return 0


def is_conditional_end_of_line(line):
    for ceol in eol_conditional_array:
        #print("Testing is_conditional_end_of_line '%s' on string '%s'" %(ceol, line))
        if re.search(ceol, line):
            return 1
    return 0

def is_beginning_of_line(line):
    for bol in bol_array:
        if re.search(bol, line):
            return 1
    return 0

def is_empty_line(line):
    line_trimmed = re.sub(' +', '', line)
    length = len(line_trimmed)
    if length == 0:
        return 1
    return 0

def get_paragraph_shading_color(xml_paragraph_str):
    paragraph_xml = etree.fromstring(xml_paragraph_str)
    attrib_fill = None
    
    namespaces = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        namespaces = {paragraph_xml.prefix : paragraph_xml.nsmap[paragraph_xml.prefix]}
    except:
        #print("Could not determine namespace")
        pass
    attrib_fill = None
    
    for e in paragraph_xml.findall('.//w:pPr/w:shd', namespaces):
        #print("e:", etree.tostring(e, pretty_print=True))
        try:
            attrib_val = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            attrib_color = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            attrib_fill = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            #print(f"attrib_color : {attrib_color}")
            #print(f"attrib_fill : {attrib_fill}")
            #print(f"attrib_val : {attrib_val}")
        except:
            pass
    return attrib_fill


def get_run_shading_color(xml_run_str):
    run_xml = etree.fromstring(xml_run_str)
    
    namespaces = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        namespaces = {run_xml.prefix : run_xml.nsmap[run_xml.prefix]}
    except:
        #print("Could not determine namespace")
        pass
    attrib_fill = None
    
    for e in run_xml.findall('.//w:rPr/w:shd', namespaces):
        #print("e:", etree.tostring(e, pretty_print=True))
        try:
            attrib_val = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            attrib_color = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            attrib_fill = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        except:
            pass
    return attrib_fill

# Return cell_non_greyed_text (string), cell_is_gray (integer for boolean)
def get_cell_data(cell,row_n):
    global from_text_nb_lines_in_cell
    cell_is_gray = None
    cell_is_red = None
    cell_non_greyed_text = ''
    
    re_enter = re.compile('enter')
    re_newline = re.compile('\n')
    
    n_paragraph = 0
    n_cell_lines = 1
    
    
    for paragraph in cell.paragraphs:
        paragraphs_text = ""
        n_paragraph = n_paragraph + 1
        
        #print("paragraph:", paragraph._p.xml)
            
        root = etree.fromstring(paragraph._p.xml)
        p_shading_color = get_paragraph_shading_color(paragraph._p.xml)
        
        p_text = paragraph.text
        nb_pause = len(re.findall('(?i)(<pause>)', p_text))
        nb_enter = len(re.findall('(?i)(<enter>)', p_text))
            
        n_cell_lines = n_cell_lines + nb_pause + nb_enter
        
        if p_shading_color is not None:
            #print(paragraph.text)
            #input("Found a shaded paragraph")
            if p_shading_color in shading_color_ignore_text:
                continue
        
        #if n_paragraph > 1:
        #    print("paragraph %d" % (n_paragraph))
        previous_run_text = ""
        for run in paragraph.runs:
            current_run_text = run.text
            
            #print("cell row %d has %d runs," % (row_n, len(paragraph.runs) ))
            #print(f"current_run_text : '{current_run_text
            
            root = etree.fromstring(run.element.xml)
            run_shading_color = get_run_shading_color(run.element.xml)
            
            if run_shading_color is not None:
                #print(f"run.element.xml : {run.element.xml}")
                #print(f"current_run_text : {current_run_text}")
                #input(f"Found a shaded run {run_shading_color}")
                if run_shading_color in shading_color_ignore_text:
                    #print(f"Color {run_shading_color} in the list of colors to ignore text")
                    pass
            
            # if re_enter.match(current_run_text):
                # print("found enter")

            if str(run.font.color.rgb) == "FF0000":
                if cell_is_red == None:
                    cell_is_red = 1
            else:
                if current_run_text != "":
                    if cell_is_red == None:
                        cell_is_red = 0
                    else:
                        cell_is_red = cell_is_red * 0
                
            if run.font.highlight_color == WD_COLOR_INDEX.RED :
                pass

            if run.font.highlight_color == WD_COLOR_INDEX.GRAY_25 or run.font.highlight_color == WD_COLOR_INDEX.GRAY_50 or run.font.strike or run.font.double_strike or run.font.highlight_color == WD_COLOR_INDEX.PINK or run.font.highlight_color == WD_COLOR_INDEX.RED or run_shading_color in shading_color_ignore_text:
                #print("Found GRAY_25")
                cell_non_greyed_text = cell_non_greyed_text + ' '
                if cell_is_gray == None:
                    cell_is_gray = 1
                
            else:
                #print("Not gray")
                if current_run_text != "":
                    cell_non_greyed_text = cell_non_greyed_text + current_run_text
                    if cell_is_gray == None:
                        cell_is_gray = 0
                    else:
                        cell_is_gray = cell_is_gray * 0
                    #return cell_is_gray
            previous_run_text = current_run_text
        #if (paragraphs_text.upper() == '<ENTER>' or paragraphs_text.upper() == '<PAUSE>'):
        #    print("Found <ENTER> or <PAUSE>")
        #    #input("press enter")
        
    
    from_text_nb_lines_in_cell[row_n-1] = n_cell_lines
    #if n_cell_lines > 1:
    #    print("%d lines" % (n_cell_lines))
    #    #input("here")

    cell_non_greyed_text = cell_non_greyed_text.replace('’', "'")
    cell_non_greyed_text = cell_non_greyed_text.replace("\n", " ")
    cell_non_greyed_text = cell_non_greyed_text.replace("\r", " ")
    cell_non_greyed_text = re.sub(r'[\r\n\u2028\u2029]+', ' ', cell_non_greyed_text)
    
    cell_non_greyed_text = re.sub("(?i)<pause>", "", cell_non_greyed_text) #'remove <pause> case insensitive
    cell_non_greyed_text = re.sub("(?i)<enter>", "", cell_non_greyed_text) #'remove <pause> case insensitive
    
    cell_non_greyed_text = re.sub(' +', ' ', cell_non_greyed_text)
    cell_non_greyed_text = cell_non_greyed_text.strip()
    

    #if cell_is_gray == 1:
    #    print("FOUND A GRAY CELL")
    #time.sleep(4)
    return cell_non_greyed_text, cell_is_gray, cell_is_red



def change_cell_font(cell):

    #print("cell has %d runs," % (len(paragraph[0].runs) ))
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = dest_font
    return

def join_from_lines(line_start, line_end, separator_str):
    joined_str = ""
    row_n = line_start
    joined_str = from_text_table[row_n]
    row_n = row_n + 1
    while row_n <= line_end:
        joined_str += from_text_table[row_n]
        row_n += 1
    #print "joined_str (%d, %d)=%s<br>" % (line_start, line_end, joined_str)
    return joined_str

def tokenize_text_to_array(text, lang_code):
    lang_code = lang_code + ""
    lang_code = lang_code.lower()

    words = []
    # In japanese tokenize words
    if lang_code == 'ja' or lang_code== 'zh-cn' or lang_code == 'zh' or lang_code == 'zh-tw' or lang_code == 'ko':
        words = cjk_segmenter.tokenize(text)
    # In other languages, just use spaces
    elif lang_code == 'th':
        #words = thai_segmenter(text)
        words = word_tokenize(text)
    # In other languages, just use spaces
    else:
        #xtm.tokenize_phrase(text, dest_lang)

        # search do not split here
        #xtm.pprint_translation_memory_list()

        # Old simple split method replaced by tokenize_phrase method having do not split
        # words = text.split()

        words = xtm.tokenize_phrase(text, lang_code)
        #input("Wait, remove tokenize_phrase here..")


    return words

def divide_array(words_array, dest_lang, width):
    dest_lang = dest_lang.lower()

    #print("Divide into max %d size lines" % (width))
    count = len(words_array)
    offsets = [0]
    for w in words_array:
        offsets.append(offsets[-1] + len(w))

    minima = [0] + [10 ** 20] * count
    breaks = [0] * (count + 1)

    def cost(i, j):
        w = offsets[j] - offsets[i] + j - i - 1
        if w > width:
            return 10 ** 10
        return minima[i] + (width - w) ** 2

    def search(i0, j0, i1, j1):
        stack = [(i0, j0, i1, j1)]
        while stack:
            i0, j0, i1, j1 = stack.pop()
            if j0 < j1:
                j = (j0 + j1) // 2
                for i in range(i0, i1):
                    c = cost(i, j)
                    if c <= minima[j]:
                        minima[j] = c
                        breaks[j] = i
                stack.append((breaks[j], j+1, i1, j1))
                stack.append((i0, j0, breaks[j]+1, j))

    n = count + 1
    i = 0
    offset = 0
    while True:
        r = min(n, 2 ** (i + 1))
        edge = 2 ** i + offset
        search(0 + offset, edge, edge, r + offset)
        x = minima[r - 1 + offset]
        for j in range(2 ** i, r - 1):
            y = cost(j + offset, r - 1 + offset)
            if y <= x:
                n -= j
                i = 0
                offset += j
                break
        else:
            if r == n:
                break
            i = i + 1

    lines = []
    j = count
    while j > 0:
        i = breaks[j]
        # In japanese just join words_array without adding any spaces
        if dest_lang == 'ja' or dest_lang == 'zh-cn' or dest_lang == 'zh-tw' or dest_lang == 'ko' \
                or dest_lang == 'th':
            lines.append(''.join(words_array[i:j]))
        # In other languages, join words_array using a space
        else:
            lines.append(' '.join(words_array[i:j]))

        j = i
    lines.reverse()
    return lines

def split_phrases():
    n_last_row_phrase = 3
    last_table_row = word_translation_table_length
    cur_row_n = 2
    while cur_row_n < (last_table_row):
        if from_text_nb_lines_in_cell[cur_row_n] > 1:
            #print("from_text_nb_lines_in_cell[%d]=%d" % (cur_row_n, from_text_nb_lines_in_cell[cur_row_n]))
            #input("nb lines here")
            pass
        if from_text_is_beginning_of_line_table[cur_row_n] == 1:
            n_last_row_phrase = cur_row_n
            nb_lines_in_phrase = 1
            from_text_nb_lines_in_phrase[cur_row_n] = from_text_nb_lines_in_cell[cur_row_n]
            #print "cur_row_n=%s<br>" % (cur_row_n)
            while from_text_is_end_of_line_table[n_last_row_phrase] != 1 \
                and n_last_row_phrase < (last_table_row - 1):
                if from_text_by_phrase_separator_table[cur_row_n] == "":
                    from_text_by_phrase_separator_table[cur_row_n] = from_text_table[n_last_row_phrase]
                    from_text_by_phrase_table[cur_row_n] = from_text_table[n_last_row_phrase]
                else:
                    from_text_by_phrase_separator_table[cur_row_n] = from_text_by_phrase_separator_table[cur_row_n] + line_separator_str + from_text_table[n_last_row_phrase]
                    from_text_by_phrase_table[cur_row_n] = from_text_by_phrase_table[cur_row_n] + ' ' + from_text_table[n_last_row_phrase]
                    nb_lines_in_phrase += 1
                    #from_text_nb_lines_in_phrase[cur_row_n] += 1
                    if from_text_nb_lines_in_cell[cur_row_n] > 1:
                        pass
                        #print("from_text_nb_lines_in_cell[%d]=%d" % (cur_row_n, from_text_nb_lines_in_cell[cur_row_n]))
                        #input("nb lines here")
                    from_text_nb_lines_in_phrase[cur_row_n] += from_text_nb_lines_in_cell[n_last_row_phrase]
                n_last_row_phrase += 1
            if from_text_by_phrase_separator_table[cur_row_n] == "":
                from_text_by_phrase_separator_table[cur_row_n] = from_text_table[n_last_row_phrase]
                from_text_by_phrase_table[cur_row_n] = from_text_table[n_last_row_phrase]
            else:
                from_text_by_phrase_separator_table[cur_row_n] = from_text_by_phrase_separator_table[cur_row_n] + line_separator_str+ from_text_table[n_last_row_phrase]
                nb_lines_in_phrase += 1
                #from_text_nb_lines_in_phrase[cur_row_n] += 1
                from_text_nb_lines_in_phrase[cur_row_n] += from_text_nb_lines_in_cell[n_last_row_phrase]
                from_text_by_phrase_table[cur_row_n] = from_text_by_phrase_table[cur_row_n] + ' ' + from_text_table[n_last_row_phrase]
            if use_html:
                print("(%d)from_text_by_phrase_table[%d]=%s<br>" % (n_last_row_phrase, cur_row_n, from_text_by_phrase_table[cur_row_n]))
            nb_lines_in_phrase_str = "[%s]" % (nb_lines_in_phrase)
            #from_text_by_phrase_separator_table[cur_row_n] = from_text_by_phrase_separator_table[cur_row_n]

            #nb_character_total = nb_character_total + 1#len(from_text_by_phrase_table[cur_row_n])
            cur_row_n = n_last_row_phrase + 1
        else:
            cur_row_n += 1

    return 0

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def generate_tmx_file():
    print("In generate_tmx_file")

    try:
        f = open(tmx_file_path, 'w', encoding='utf-8')

        # Writing TMX Header
        username = getpass.getuser()
        datenow = datetime.datetime.now()
        creation_date = "%s%0.2d%0.2dT%0.2d%0.2d%0.2dZ" % (datenow.year, datenow.month, datenow.day, datenow.hour, datenow.minute,
 datenow.second)
        header = """<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE tmx PUBLIC "-//LISA OSCAR:1998//DTD for Translation Memory eXchange//EN" "tmx14.dtd" >
<tmx version="1.4">
<header
	creationtool="SMTV translation robot"
	creationtoolversion="1.0"
	srclang="%s"
	adminlang=%s
	datatype="unknown"
	o-tmf="unknown"
	segtype="sentence"
	creationid="%s"
	creationdate="%s">
</header>
<body>\n""" % (src_lang, src_lang, username,creation_date)
        f.write(header)

        for i, line in enumerate(from_text_table):
            item = from_text_by_phrase_separator_table[i]
            item.strip()
            from_language = src_lang
            phrase_separator_removed_str = ''

            p_remove_separator = re.compile(line_separator_regex_str)
            p_remove_double_spaces = re.compile(' +')
            p_remove_parenthesis_spaces = re.compile('\( +')

            item = from_text_by_phrase_table[i]
            item_escaped = from_text_by_phrase_table[i].replace("&", "&amp;")
            item_escaped = item_escaped.replace("<", "&lt;")
            item_escaped = item_escaped.replace(">", "&gt;")

            item_translation = to_text_by_phrase_separator_table[i].replace("&", "&amp;")
            item_translation = item_translation.replace("<", "&lt;")
            item_translation = item_translation.replace(">", "&gt;")
            if item_escaped.strip() != "":
                segment = """<tu changeid="french user 1" changedate="%s" creationid="Black Mamba RS7" creationdate="%s" creationtool="SMTV translation robot" creationtoolversion="1.0.0">
<tuv xml:lang="en-US"><seg>%s</seg></tuv>
<tuv xml:lang="%s"><seg>%s</seg></tuv>
</tu>""" % (creation_date, creation_date, item_escaped, dest_lang, item_translation)
                f.write(segment)
                f.write("\n")

        # Writing TMX Footer
        footer = """</body>
</tmx>\n"""
        f.write(footer)
    except Exception:
        var = traceback.format_exc()
        print(var)


def prepare_and_clear_cell_for_writing(row_n, translation_cell_text):
    global table_cells
    paragraph_no = 0
    current_cell = table_cells[row_n][2]

    # Clear paragraphs in the cell
    for paragraph in current_cell.paragraphs:
        if paragraph_no != 0:
            delete_paragraph(paragraph)
        else:
            paragraph.text = ''
        paragraph_no += 1

    # Ensure there's at least one paragraph in the cell
    if len(current_cell.paragraphs) == 0:
        cell_paragraph = current_cell.add_paragraph("")
    else:
        cell_paragraph = current_cell.paragraphs[0]

    # Add orientation for Right-to-Left (RTL) languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(translation_cell_text)
        run.style = rtlstyle  # Ensure `rtlstyle` exists in the document
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = translation_cell_text
        # Check if the 'Normal' style exists before applying
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Apply font changes if necessary
    if dest_font != "":
        change_cell_font(current_cell)

    table_cells[row_n][2] = current_cell
    
        
def cell_set_1st_paragraph(row_n, paragraph_text):
    paragraph_no = 0
    current_cell = table_cells[row_n][2]
    
    #print("cell_add_paragraph")
    #print("paragraph[%d]: %s" % (row_n,paragraph_text))
    cell_paragraph = cell_paragraph = current_cell.paragraphs[0]

    # Add orientation from Right To Left (RTL) for specific languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text,style = "rtlstyle")
        run.style = rtlstyle
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if dest_font != "":
        change_cell_font (current_cell)
        
    table_cells[row_n][2] = current_cell
        

def cell_add_paragraph(row_n, paragraph_text):
    paragraph_no = 0
    current_cell = table_cells[row_n][2]
    
    #print("cell_add_paragraph")
    #print("paragraph[%d]: %s" % (row_n,paragraph_text))
    cell_paragraph = current_cell.add_paragraph("")

    # Add orientation from Right To Left (RTL) for specific languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text,style = "rtlstyle")
        run.style = rtlstyle
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if dest_font != "":
        change_cell_font (current_cell)
        
    table_cells[row_n][2] = current_cell

def read_and_parse_docx_document():
    global from_text_table
    global from_text_is_greyed_table
    global from_text_is_red_color_table
    global from_text_is_end_of_line_table
    global from_text_is_beginning_of_line_table
    global from_text_is_empty_line_table
    global from_text_is_conditional_end_of_line_table
    global from_text_by_phrase_separator_table
    global from_text_by_phrase_table
    global from_text_nb_lines_in_phrase
    global from_text_nb_lines_in_cell
    global to_text_by_phrase_separator_table
    global to_text_by_phrase_separator_removed_table
    global to_text_splited_table1
    global to_text_by_phrase_table
    global to_text_table
    global to_raw_translated_table
    global to_text_removed_line_separator
    global translation_result_using_separator
    global translation_result_phrase_array
    global translation_result
    global from_text_is_read

    global table_cells

    global word_translation_table_length

    global table, numrows, numcols
    global E_mail_str


    start = timeit.timeit()

    if use_html:
        print("Content-Type: text/html\n")

    word_translation_table_length = len(docxdoc.tables[0].rows)

    nb_tables = len(docxdoc.tables)

    nb_character_total = 0

    if use_html:
        print(
            "<!doctype html><head><meta http-equiv=""Content-Type"" content=""text/html"" charset=utf-8 /><title>Winword in python</title></head><h2>tables</h2><span style=""font-family:monospace,monospace;"">")

    # Number of tables</h2>nb_tables=", nb_tables

    numerrors = 0
    # print("word_translation_table_length=%d" %(word_translation_table_length))
    # print("docx_translation_table_length=%d" %(docx_translation_table_length))

    table = docxdoc.tables[0]
    table_cells = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]

    numrows = len(table.rows)
    numcols = len(table.columns)

    if numcols <= 2:
        print("ERROR : The table has %s column but expected 3" % (numcols))
        print("Exiting\n")

        print("\nDeveloper: %s" %(E_mail_str))
        print("Program version: %s\n" % (PROGRAM_VERSION))
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(11)

    rownum = 0

    from_text_table = [''] * (numrows + 1)
    from_text_is_greyed_table = [0] * (numrows + 1)
    from_text_is_red_color_table = [0] * (numrows + 1)
    from_text_is_end_of_line_table = [0] * (numrows + 1)
    from_text_is_beginning_of_line_table = [0] * (numrows + 1)
    from_text_is_empty_line_table = [0] * (numrows + 1)
    from_text_is_conditional_end_of_line_table = [0] * (numrows + 1)
    from_text_by_phrase_separator_table = [''] * (numrows + 1)
    from_text_by_phrase_table = [''] * (numrows + 1)
    #number of lines in per phrase
    from_text_nb_lines_in_phrase = [0] * (numrows + 1)
    from_text_nb_lines_in_cell = [0] * (numrows + 1)
   #input(numrows)
    #
    to_text_by_phrase_separator_table = [''] * (numrows + 1)
    to_text_by_phrase_separator_removed_table = [''] * (numrows + 1)
    to_text_splited_table1 = [''] * (numrows + 1)
    to_text_by_phrase_table = [''] * (numrows + 1)
    to_text_table = [''] * (numrows + 1)
    to_raw_translated_table = [''] * (numrows + 1)
    to_text_removed_line_separator = [''] * (numrows + 1)
    translation_result_using_separator = [''] * (numrows + 1)
    translation_result_phrase_array = [[]] * (numrows + 1)
    translation_result = [''] * (numrows + 1)
    from_text_is_read = [0] * (numrows + 1)

    if use_html :
        print("<br>%s rows.<br>%d colums.<br>" % (numrows, numcols))

    for i, row in enumerate(table.rows):
        col_no = 1
        row_n = i + 1
        
        p_remove_pause = re.compile('(?i)<pause>')
        p_remove_double_spaces = re.compile(' +')
        p_remove_parenthesis_spaces = re.compile('\( +')
        
        try:
            for j, cell in enumerate(row.cells):
                #if cell.text:
                #    df[i][j] = cell.text
                table_cells[i][j] = cell
                # XML is ._tc
                #df[i][j] = cell._tc
                if col_no == 2:
                
                    #from_text_is_greyed_table[row_n] = is_greyed_line(cell)
                    #cellvalue = cell.text.replace('’', "'").strip()
                    #print(from_text_is_greyed_table)
                    #print(from_text_is_red_color_table)
                    #print("row_n=%d" % (row_n))
                    cellvalue, from_text_is_greyed_table[i], from_text_is_red_color_table[i] = get_cell_data(cell,row_n)
                    p_remove_pause
                    cellvalue = p_remove_pause.sub(' ', cellvalue)
                    cellvalue = p_remove_double_spaces.sub(' ', cellvalue)
                    cellvalue = p_remove_parenthesis_spaces.sub('(', cellvalue)
                    length = len(cellvalue)

                    try:
                        print("%d : %s" % (i, cellvalue), flush=True)
                    except Exception:
                        try:
                            print("%d : %s" % (i, cellvalue.encode("utf-8")))
                        except Exception:
                            print("%d : (unable to print content to screen)" )

                    from_text_is_end_of_line_table[i] = is_end_of_line(cellvalue) or from_text_is_red_color_table[i]
                    from_text_is_empty_line_table[i] = is_empty_line(cellvalue)
                    from_text_is_beginning_of_line_table[i] = is_beginning_of_line(cellvalue)
                    from_text_is_conditional_end_of_line_table[i] = is_conditional_end_of_line(cellvalue)

                    if from_text_is_greyed_table[i] == 1:
                        from_text_is_beginning_of_line_table[i] = 0
                        from_text_is_end_of_line_table[i] = 0
                        
                    if i == 2 and len(cellvalue) > 0:
                        from_text_is_beginning_of_line_table[i] = 1

                    if i > 1:
                        #Test conditionel de fin de ligne
                        if from_text_is_conditional_end_of_line_table[i - 1] == 1 \
                            and from_text_is_beginning_of_line_table[i] == 1:
                            from_text_is_end_of_line_table [i - 1] = 1
                            from_text_is_beginning_of_line_table [i] = 1

                        # Verifier debut de ligne special
                        # Si ligne precedente est vide ou grisee:
                        #    Si ligne courante est non vide et non grisee
                        #        ligne courante est debut de ligne
                        if (from_text_is_empty_line_table[i - 1] == 1 \
                            or from_text_is_greyed_table[i - 1] == 1):
                            if (from_text_is_empty_line_table[i] == 1 \
                                and from_text_is_greyed_table[i] == 1):
                                from_text_is_beginning_of_line_table[i] = 1

                        # Verifier la ligne precedente est fin de ligne
                        # Si ligne precedente est non vide et non grisee
                        #    Si ligne courante est vide ou grisee
                        #        la ligne precedente est fin de ligne
                        if (from_text_is_empty_line_table[i - 1] == 0 \
                            and from_text_is_greyed_table[i - 1] == 0):
                            if (from_text_is_empty_line_table[i] == 1 \
                                or from_text_is_greyed_table[i] == 1):
                                from_text_is_end_of_line_table[i - 1] = 1


                        # Verifier que c'est vraiment un debut de ligne suivant une fin de ligne
                        # Si ligne precedente n'est pas fin de ligne
                        #    et ligne oourante est debut de ligne
                        #        la ligne courante n'est pas un debut de ligne
                        if from_text_is_beginning_of_line_table[i] == 1 and \
                            from_text_is_end_of_line_table[i - 1] == 0 \
                            and from_text_is_greyed_table[i - 1] == 0 \
                            and i > 2:
                            from_text_is_beginning_of_line_table[i] = 0


                        # Verifier qu'on a pas loupe un debut de ligne
                        # Si ligne precedente est fin de ligne
                        #    et ligne oourante n'est pas grisee et pas debut de ligne
                        #        la ligne courante est un debut de ligne
                        if from_text_is_end_of_line_table[i - 1] == 1 \
                            and from_text_is_greyed_table[i] == 0 \
                            and from_text_is_beginning_of_line_table[i] == 0:
                            from_text_is_beginning_of_line_table[i] = 1

                        if (from_text_is_empty_line_table[i - 1] == 1 \
                            or from_text_is_greyed_table[i - 1] == 1) \
                            and (from_text_is_empty_line_table[i] == 0 \
                            and from_text_is_greyed_table[i] == 0):
                            from_text_is_beginning_of_line_table[i] = 1

                        if from_text_is_empty_line_table[i - 1] == 1:
                            from_text_is_beginning_of_line_table[i - 1] = 0

                        if i == numrows:
                            from_text_is_end_of_line_table[i - 1] = 1

                    from_text_table[i] = cellvalue
                col_no = col_no + 1
            
            if not splitonly and i > 1:
                prepare_and_clear_cell_for_writing (i, '')
            from_text_is_read[i] = 1
        except Exception:
            var = traceback.format_exc()
            print(var)
            numerrors = numerrors + 1

    if from_text_is_greyed_table[numrows] == 0 \
        and from_text_is_empty_line_table[numrows] == 0:
        from_text_is_end_of_line_table[numrows] = 1

    split_phrases()

    if use_html :
        print("<table border=1 width=800>")

    for row_n in range(1, len(from_text_table)):
        try:
            if use_html :
                print("<tr>")
                print("<td width=50>", row_n)
                print("<td width=250>")

            if from_text_is_beginning_of_line_table[row_n] == 1:
                if use_html :
                    print("<hr style=\"height:5px;border:none;color:#ffff00;background-color:#ffff00;\" />")
            
            if from_text_is_greyed_table[row_n] == 1:
                if use_html :
                    print("'<span style=\"background-color: #DCDCDC\">%s</span>' (%s)" % (from_text_table[row_n], len(from_text_table[row_n])))
                    print("<hr style=\"height:5px;border-top: dotted 2px;color:##DCDCDC;background-color:#DCDCDC;\" />")
            else:

                if use_html :
                    print("'%s' (%s)" % (from_text_table[row_n], len(from_text_table[row_n])))

            if from_text_is_end_of_line_table[row_n] == 1:
                
                if use_html :
                    print("<hr style=\"height:5px;border:none;color:#333;background-color:#333;\" />")
            
            if from_text_is_empty_line_table[row_n] == 1:

                if use_html :
                    print("<hr style=\"height:5px;border-top: dotted 2px;color:##DCDCDC;background-color:#DCDCDC;\" />")
                    print("<td>is_greyed=%s<br>is_end_of_line=%s<br>is_empty_line=%s<br>is_beginning_of_line=%s<br>is_conditional_end_of_line=%s" %(
                from_text_is_greyed_table[row_n], \
                from_text_is_end_of_line_table[row_n], \
                from_text_is_empty_line_table[row_n], \
                from_text_is_beginning_of_line_table[row_n], \
                from_text_is_conditional_end_of_line_table[row_n]))

            if use_html :
                print("<td>'%s' (%d)<td>'%s' (%d)" % (from_text_by_phrase_table[row_n], len(from_text_by_phrase_table[row_n]), \
                                                  from_text_by_phrase_separator_table[row_n], len(from_text_by_phrase_separator_table[row_n])))
        except Exception:
            var = traceback.format_exc()
            print(var)
            numerrors = numerrors + 1

def clean_up_previous_chrome_selenium_drivers(current_driver_full_path):
    found_previous_chrome_driver = False
    
    try:
        list_driver_path = []
        
        if platform.system().lower() == 'windows':
            userprofile_path = os.environ.get('USERPROFILE')
            selenium_cache_folder = f"{userprofile_path}\\.cache\\selenium"
            list_driver_path = glob.glob(f"{selenium_cache_folder}\\**\\chromedriver.exe", recursive=True)
        else:
            home_path = os.environ.get('HOME')
            selenium_cache_folder = f"{home_path}/.cache/selenium"
            list_driver_path = glob.glob(f"{selenium_cache_folder}/*/**/chromedriver", recursive=True)
            
        for driver_path in list_driver_path:
            if driver_path == current_driver_full_path:
                pass # Latest version of the driver, keep the file
            else:
                if os.path.exists(driver_path):
                    try:
                        if found_previous_chrome_driver == False:
                            print("\nCleaning up old chrome driver files")
                            found_previous_chrome_driver = True
                        print(f"Removing previous chrome driver at {driver_path}")
                        os.remove(driver_path)
                    except:
                        print(f"Unable to cleanup chrome driver at {driver_path}")
                        
        if len(list_driver_path) >= 2:
            print(f"Keeping current chrome driver at {current_driver_full_path}")
                
    except:
        var = traceback.format_exc()
        print(var)
        

def create_webdriver():
    global driver, chromedriverpath, translation_engine
    if not splitonly:
        print("\nStarting translation using engine : %s" % (translation_engine.title()))


    if use_api == False and not splitonly:
        print(f"Starting Chrome browser\n")
        service = Service()
        
        try:
            driver = uc.Chrome(service=service, options=chrome_options)
        except:
            print("An error occured during launching chrome. This may happen during google chrome automatic updates or if Google Chrome is not installed.")
            print("You may start google chrome and open the menu Help -> About Google Chrome to see if there is an update running and retry machine translation after the update.")
            print("Exiting, please retry.")
            
            print("\nDeveloper: %s" % (E_mail_str))
            print("Program version: %s\n" % (PROGRAM_VERSION))
            if not exitonsuccess:
                input("Enter to close program")
            
            sys.exit(12)
        
        print("\nChrome started using driver at %s\n" % (driver.service.path))

        #input("driver loaded and running")
        #driver.set_window_position(0, 350)
        if translation_engine == 'yandex' or translation_engine == 'deepl':
            driver.set_window_position(0, 100)
            driver.set_window_size(800, 700)
        else:
            driver.set_window_size(800, 700)
            #driver.set_window_size(400, 650)


    numerrors_deepl = 0
    numerrors_googletranslate= 0

# Reverse a string
def reverse_string(s):
    return s[::-1]


def generate_html_file_from_phrases_for_google_translate_javascript():
    #input("Here")
    global dest_lang_name, html_file_path, docxfile_table_number_of_phrases
    print("Generating html page.")

    docxfile_table_number_of_phrases = 0
    html_to_translate = '''<html lang=%s >
<head>
  <meta charset="UTF-8">
  <title>machine-translate-docx - %s - %s</title>
</head> 
<body>
<h3 translate="no">%s - %s</h3>
<p lang=%s translate="yes" id="%s">%s</p>
<table border=1 CELLPADDING=5 CELLSPACING=0>
<tr><td translate="no">Line No</td><td translate="no">%s<td translate="no">%s</td></tr>
''' % (src_lang, docx_file_name, dest_lang_name, docx_file_name, dest_lang_name, src_lang_name, dest_lang_name, dest_lang_name, src_lang_name, dest_lang_name)
    
    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item, count=False)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        item_html_escaped = html.escape(item_searched_and_replaced_before.strip())  
        
        if item_searched_and_replaced_before != '':
            docxfile_table_number_of_phrases = docxfile_table_number_of_phrases + 1
            html_to_translate = html_to_translate + '''
<tr>
    <td>%s</td>
    <td><p id="phrase_%s_%s" lang="%s" translate="no" class="source">%s</p></td>
    <td id="%s"><p id="phrase_%s_%s" lang=%s class="translation">%s</td>
</tr>
''' % (i+1, i, src_lang, src_lang, item_html_escaped, i, i, dest_lang, src_lang, item_html_escaped)
    html_to_translate = html_to_translate + '''
</table>
<div id="google_translate_element"></div>

<script type="text/javascript">
function googleTranslateElementInit() {
  new google.translate.TranslateElement({pageLanguage: '%s'}, 'google_translate_element');
}
</script>
<script>
    function googleTranslateElementInit() {
        new google.translate.TranslateElement({pageLanguage: '%s', includedLanguages: '%s', autoDisplay: true}, 'google_translate_element');
        var a = document.querySelector("#google_translate_element select");
        a.selectedIndex=1;
        a.dispatchEvent(new Event('change'));
    }
</script>

<script type = "text/javascript">  
window.onload = function(){  
        var a = document.querySelector("#google_translate_element select");
        a.selectedIndex.value = '%s';
		//document.querySelector('#google_translate_element select').value = '%s';
        a.dispatchEvent(new Event('change')); 
}  
</script>

<script type="text/javascript" src="https://translate.google.com/translate_a/element.js?cb=googleTranslateElementInit"></script>
</body>
''' % (src_lang, src_lang, dest_lang,dest_lang,dest_lang)
    #print (html_to_translate)
    try:
        if(platform.system() == "Darwin"):
            # Write to TMPDIR or /tmp folder
            try:
                tmpdir = os.environ['TMPDIR']
                if(tmpdir is None or tmpdir == ""):
                    tmpdir = '/tmp/'
            except:
                tmpdir = '/tmp/'
            html_file_path = tmpdir + docx_file_name + '.' + str(os.getpid()) + '.' + dest_lang + '.html'
        else:
            # Windows, write to file at the same location
            html_file_path = os.path.abspath(os.path.expanduser(os.path.expandvars(word_file_to_translate))) + '.' + str(os.getpid()) + '.' + dest_lang + '.html'
        
        print(f"Writing temporary html file to : {html_file_path}")
        html_file_to_translate = open(html_file_path, 'w', encoding='utf-8')
        html_file_to_translate.write(html_to_translate)
        html_file_to_translate.close()
        #print("HTML FILE WRITTEN !")
    except Exception:
        var = traceback.format_exc()
        print(var)


def generate_text_file_from_phrases(text_file_path):
    global dest_lang_name
    global docxfile_table_number_of_phrases
    global xtm
    docxfile_table_number_of_phrases = 0
    print("Generating text file for google translation...")
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing text before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    
    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
#''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            docxfile_table_number_of_phrases = docxfile_table_number_of_phrases + 1
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    for index in range(len(text_to_translate_array)):
        #print("%d : '%s'" % (index, text_to_translate_array[index]))
        if index == (len_text_to_translate_array - 1):
            text_to_translate = text_to_translate + '%s' % (text_to_translate_array[index])
        else:
            text_to_translate = text_to_translate + '%s\n' % (text_to_translate_array[index])

    #print("text_to_translate=\n%s" % (text_to_translate))
    try:
        #text_file_path = docx_file_name + '.txt'
        text_file_to_translate = open(text_file_path, 'w', encoding='utf-8')
        text_file_to_translate.write(text_to_translate)
        text_file_to_translate.close()
    except Exception:
        var = traceback.format_exc()
        print(var)
        

def generate_xlsx_file_from_phrases(xlsx_file_path):
    global dest_lang_name
    global docxfile_table_number_of_phrases
    global xtm
    docxfile_table_number_of_phrases = 0
    print("Generating xlsx file for google translation...")
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing xlsx before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    
    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
            #''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            docxfile_table_number_of_phrases = docxfile_table_number_of_phrases + 1
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "English"
    except Exception:
        print ("Error creating empty xlsx workbook")
        var = traceback.format_exc()
        print("ERROR: %s" % (var))
        self.wb = None
        self.ws = None
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(13)
    
    index_current_row = 1
    max_col_length = 0
    for index in range(len(text_to_translate_array)):
        ws.cell(row=index_current_row, column=1, value=text_to_translate_array[index])
        if len(text_to_translate_array[index]) > max_col_length:
            max_col_length = len(text_to_translate_array[index])
        index_current_row = index_current_row + 1
    
    ws.column_dimensions['A'].width = max_col_length + 1000
    
    file_saved = 0
    while file_saved == 0:
        try:
            wb.save(xlsx_file_path)
            print ("Excel XLSX english text to translate file \"%s\" saved..." %(xlsx_file_path))
            file_saved=1
        except Exception:
            var = traceback.format_exc()
            txt_readline = input("\n\nERROR: File saving failed. Please close microsoft excel or other program and press enter to save the xlsx document.\n")
        
        
def generate_char_blocks_array_from_phrases(text_file_path):
    global dest_lang_name
    global docxfile_table_number_of_phrases
    global xtm
    global blocks_nchar_max_to_translate_array
    docxfile_table_number_of_phrases = 0
    print("Generating %d character blocks for translation..." % (MAX_TRANSLATION_BLOCK_SIZE))
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing text before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    blocks_nchar_max_to_translate_array = []
    
    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
#''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            docxfile_table_number_of_phrases = docxfile_table_number_of_phrases + 1
                
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    
    current_text_block = ""
    current_text_block_len = len(current_text_block)
    
    
    #print(text_to_translate_array)
    #print(len_text_to_translate_array)
    #input("len_text_to_translate_array")
    #print("Before current_text_block generation")
    for index in range(0, len(text_to_translate_array)):
        #print("%d : '%s'" % (index, text_to_translate_array[index]))
        current_phrase_str = text_to_translate_array[index]
        
        if current_text_block_len + len(current_phrase_str) <= MAX_TRANSLATION_BLOCK_SIZE:
            if len(current_text_block) == 0:
                current_text_block = current_phrase_str
                current_text_block_len = len(current_text_block)
                #print(current_phrase_str)
                #input("adding first phrase")
            else:
                current_text_block = current_text_block + "\n" + current_phrase_str
                current_text_block_len = current_text_block_len + len(current_phrase_str) + 1
                #print(current_phrase_str)
                #print("current_text_block")
                #print(current_text_block)
        else:
            blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------" % (MAX_TRANSLATION_BLOCK_SIZE))
            #print("current_text_block")
            #print(current_text_block)
            #print("end")
            #input("OK Here")
            #input("adding more phrase")
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            
            current_text_block = current_phrase_str
            current_text_block_len = len(current_phrase_str)
            #input("------------------------------------------\nType enter to continue")
        
        if index == (len(text_to_translate_array) - 1):
            blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------"  % (MAX_TRANSLATION_BLOCK_SIZE))
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            #input("------------------------------------------\nType enter to continue")
    
    #print("blocks_nchar_max_to_translate_array:")
    #print("\n******************************************\n".join(blocks_nchar_max_to_translate_array))
    #print ("len(text_to_translate_array) = %d " % (len(text_to_translate_array) - 1)) 

    #print("text_to_translate=\n%s" % (text_to_translate))
    try:
        #text_file_path = docx_file_name + '.txt'
        #text_file_to_translate = open(text_file_path, 'w', encoding='utf-8')
        #text_file_to_translate.write(text_to_translate)
        #text_file_to_translate.close()
        pass
    except Exception:
        var = traceback.format_exc()
        print(var)


def google_translate_from_text_file():
    global docx_file_name, translation_array
    #word_file_to_translate
    text_file_path = docx_file_name + '.txt'
    text_file_full_path = os.path.realpath(text_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    generate_text_file_from_phrases(text_file_full_path)
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using text file...")
    translation_array = selenium_chrome_google_translate_text_file(text_file_full_path)
    try:
        os.remove(text_file_path)
        pass
    except:
        pass

def google_translate_from_html_javascript():
    global translation_array
    global html_file_path
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using html file...")
    
    generate_html_file_from_phrases_for_google_translate_javascript()
    
    translation_array = selenium_chrome_google_translate_html_javascript_file(html_file_path)
    
    try:
        #input("before remove html file")
        os.remove(html_file_path)
        pass
    except:
        pass

    return translation_array

def google_translate_from_html_xlsxfile():
    global word_file_to_translate, translation_array
    xlsx_file_path = word_file_to_translate + '.xlsx'
    xlsx_file_full_path = os.path.realpath(xlsx_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    generate_xlsx_file_from_phrases(xlsx_file_full_path)
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using text file...")
    translation_array = selenium_chrome_google_translate_xlsx_file(xlsx_file_full_path)
    try:
        os.remove(xlsx_file_full_path)
        pass
    except:
        pass

def translate_from_phrasesblock():
    global docx_file_name, translation_array, translation_engine
    text_file_path = docx_file_name + '.txt'
    text_file_full_path = os.path.realpath(text_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    #generate_text_file_from_phrases(text_file_full_path)
    generate_char_blocks_array_from_phrases(text_file_full_path)

    translation_succeded = True

    #input("phrasesblock")
    print("Starting translation in %s using phrase blocks of %d characters..." % (translation_engine, MAX_TRANSLATION_BLOCK_SIZE))

    translation_succeded, translation_array = selenium_chrome_translate_maxchar_blocks()
    try:
        os.remove(text_file_path)
        pass
    except:
        pass
    return translation_succeded



def translate_docx():
    translation_array = []
    translation_succeded = True

    if engine_method == 'textfile':
        google_translate_from_text_file()

    if engine_method == 'javascript':
        translation_array = google_translate_from_html_javascript()
        #print (translation_array)

    if engine_method == 'xlsxfile':
        google_translate_from_html_xlsxfile()

    # For both deepl and google translate
    if engine_method == "phrasesblock":
        # For both deepl and google translate
        if translation_engine == "deepl" or translation_engine == "chatgpt" or translation_engine == "perplexity":
            translation_succeded = translate_from_phrasesblock()

    return translation_succeded

def get_translation_and_replace_after():
    global from_text_by_phrase_separator_table, to_text_by_phrase_separator_table, numerrors_deepl, use_api
    phrase_no = 0

    p_remove_pause = re.compile('(?i)<pause>')
    p_remove_double_spaces = re.compile(' +')
    p_remove_parenthesis_spaces = re.compile('\( +')

    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item = item.strip()
        from_language = src_lang
        phrase_separator_removed_str = ''

        p_remove_separator = re.compile(line_separator_regex_str)
        p_remove_double_spaces = re.compile(' +')

        # Avec separateurs ()

        try:
            web_translation_separators = ''
            if item.strip() != '':
                phrase_no = phrase_no + 1
                print("\n%d/%d" % (i, word_translation_table_length))
                print("Phrase to translate :'%s'\n" % (item.strip()))
                item = item.strip()

                item_searched_and_replaced_before = item
                if xlsxreplacefile is not None:
                    if xtm.wb is not None:
                        item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                        if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                            continue
                if splitonly:
                    web_translation_separators = get_translated_cells_content (i, item_searched_and_replaced_before)
                elif use_api:
                    try:
                        web_translation_separators = ""
                        if use_api:
                            translation = translator.translate(item_searched_and_replaced_before, src=src_lang, dest=dest_lang)
                            web_translation_separators = translation.text
                        if not len(web_translation_separators) > 0:
                            use_api = False
                            # Faster google Chrome translate failed, using Selenium as backup

                            if translation_engine != 'yandex' and driver is not None:
                                print(f"[Line {inspect.currentframe().f_lineno}] Starting Chrome browser\n")
                                
                                service = Service()                                
                                driver = uc.Chrome(service=service, options=chrome_options)
                                
                                driver.set_window_position(100, 100)
                                driver.set_window_size(800, 700)
                                #driver.set_window_size(400, 650)

                            print("phrase_no=%d" % phrase_no)
                            web_translation_separators = selenium_chrome_machine_translate(item_searched_and_replaced_before, phrase_no)
                    except Exception:
                        use_api = False
                        # Faster google Chrome translate failed, using Selenium as backup

                        if driver is not None:
                            print(f"Starting Chrome browser\n")
                            
                            service = Service()                                
                            driver = uc.Chrome(service=service, options=chrome_options)

                        if translation_engine == 'google' and driver is not None:
                            driver.set_window_position(100, 100)
                            driver.set_window_size(800, 700)

                        if translation_engine == 'yandex' and driver is not None:
                            driver.set_window_position(100, 100)
                            driver.set_window_size(800, 700)

                        if translation_engine == 'deepl' and driver is not None:
                            driver.set_window_position(100, 100)
                            driver.set_window_size(800, 700)

                        print("phrase_no = %d" % phrase_no)
                        web_translation_separators = selenium_chrome_machine_translate(item_searched_and_replaced_before, phrase_no)
                else:
                    if engine_method == "singlephrase" and translation_engine == 'deepl':
                        translation_succeded, web_translation_separators  = selenium_chrome_machine_translate(item_searched_and_replaced_before, phrase_no)
                    else:
                        web_translation_separators = selenium_chrome_machine_translate(item_searched_and_replaced_before, phrase_no)

                #web_translation_separators = translation.text
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_separators)

                #print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8')))
                if xlsxreplacefile is not None:
                    nb_searched_and_replaced = 0
                    web_translation_separators_searched_and_replaced, nb_searched_and_replaced = xtm.search_and_replace_text('after', phrase_separator_removed_str)
                    if nb_searched_and_replaced > 0:
                        #print("\nPhrase %d replacements :\n'%s'" % (nb_searched_and_replaced, web_translation_separators))
                        #print("Replaced phrase :\n'%s'" % (web_translation_separators_searched_and_replaced))
                        phrase_separator_removed_str = web_translation_separators_searched_and_replaced

                if dest_lang in right_to_left_languages_list.keys():
                    phrase_separator_removed_aligned_str = reverse_string (phrase_separator_removed_str)
                else:
                    phrase_separator_removed_aligned_str = phrase_separator_removed_str
                try:
                    if splitonly:
                        print("Translated text :'%s'\n" % (phrase_separator_removed_aligned_str))
                    else:
                        print("%s translation (%s):'%s'" % (translation_engine.title() ,dest_lang_name, phrase_separator_removed_aligned_str))
                except Exception:
                    print("")
                    print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8').decode('utf8')))
                if web_translation_separators.strip() == '' and not splitonly:
                    print("Error translating='%s'" % (item))
                to_text_by_phrase_separator_table[i] = phrase_separator_removed_str
                phrase_separator_removed_str = p_remove_separator.sub(' ', phrase_separator_removed_str)
                phrase_separator_removed_str.strip()
                to_text_by_phrase_separator_removed_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            numerrors_deepl = numerrors_deepl + 1
            web_translation_separators = var
            print("ERROR:%s" % (var))

        item = from_text_by_phrase_table[i]
        try:
            web_translation_no_separators = ''
            if item.strip() != '':
                #google_translation_res = translator.translate(item, src=src_lang, dest='fr')
                #time.sleep(5)
                #web_translation_no_separators = pydeepl.translate(item, to_language)
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_no_separators)
                phrase_separator_removed_str = p_remove_parenthesis_spaces.sub('(', phrase_separator_removed_str)
                to_text_by_phrase_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            numerrors_googletranslate = numerrors_googletranslate + 1
            web_translation_no_separators = var
        Identical_with_without_separators = 'DIFFERENT<BR>'
        if phrase_separator_removed_str == web_translation_no_separators:
            Identical_with_without_separators = 'SAME<BR>'


def minimize_browser():
    if not use_api and not splitonly:
        # Minimize browser
        #print("Minimizing browser...")
        try:
            driver.minimize_window()
        except:
            pass


def document_split_phrases():
    # Split phrases into multiple lines to match source language number of lines
    global docxfile_table_number_of_phrases, docxfile_table_number_of_characters, phrase_number_of_words, docxfile_table_number_of_words
    for i, line in enumerate(from_text_table):
        if to_text_by_phrase_separator_table[i] != '':
            #docxfile_table_number_of_phrases = docxfile_table_number_of_phrases + 1
            docxfile_table_number_of_characters = docxfile_table_number_of_characters + len(from_text_by_phrase_separator_table[i])
            phrase_number_of_words = len(from_text_by_phrase_separator_table[i].strip().split(" "))
            #print("Phrase to split: %s" % (from_text_by_phrase_separator_table[i]))
            #print("number of words: %d" % (phrase_number_of_words))
            docxfile_table_number_of_words = docxfile_table_number_of_words + phrase_number_of_words
            try:
                current_line = to_text_by_phrase_separator_table[i]
                # Using () as separator for splitting phrases, not used anymore
                #lines = current_line.split(line_separator_nospace_str)
                str_translation_len = len(current_line)

                try:
                    if str_translation_len <= 0:
                        str_phrase_stats = ""
                    else:
                        str_nb_lines = from_text_nb_lines_in_phrase[i]
                        if str_nb_lines > 0:
                            str_line_average = str_translation_len / str_nb_lines
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        else:
                            str_line_average = 0
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        #print("str_phrase_stats=%s" % (str_phrase_stats))
                except Exception:
                    var = traceback.format_exc()
                    print("  ERROR:%s<br>" % (var))

                if str_line_average > MAX_LINE_SIZE:
                    #input("str_line_average > MAX_LINE_SIZE : %s > %s" % (str_line_average, MAX_LINE_SIZE))
                    #str_line_average = MAX_LINE_SIZE
                    str_line_average = math.ceil(str_line_average)
                current_phrase_tokenized_array = tokenize_text_to_array(current_line, dest_lang)
                lines_divided = divide_array(current_phrase_tokenized_array, dest_lang, str_line_average + 4)

                #print "lines(%d)=%s<br>" % (len(lines), lines)
                number_lines = len(lines_divided)

                divide_max_try = MAX_LINE_SIZE
                while (number_lines > str_nb_lines) and (divide_max_try > 0):
                    str_line_average += 1
                    #print("Too many lines in split : %d, max %d ..... increasing line size to max %d" % (number_lines,str_nb_lines, str_line_average))
                    lines_divided_attempt = divide_array(current_phrase_tokenized_array, dest_lang, str_line_average + 4)
                    lines_divided = divide_array(current_phrase_tokenized_array, dest_lang, str_line_average + 4)
                    number_lines = len(lines_divided_attempt)
                    #print("   lines in split : %d, max %d ..... reducing line size to max %d" % (number_lines,str_nb_lines, str_line_average))
                    divide_max_try = divide_max_try - 1

                #print("Before increasing line size -- %s (%d): %d " % (to_text_by_phrase_separator_table[i], i, str_nb_lines + 0))
                number_lines = len(lines_divided)
                divide_max_try = MAX_LINE_SIZE
                while (number_lines < str_nb_lines) and (number_lines > 1) and (divide_max_try > 0):
                    str_line_average = str_line_average - 1
                    #print("Too few lines in split : %d, max %d ..... reducing line size to max %d" % (number_lines,str_nb_lines, str_line_average))
                    lines_divided_attempt = divide_array(current_line, dest_lang, str_line_average + 4)
                    number_lines = len(lines_divided_attempt)
                    if number_lines <= str_nb_lines:
                        lines_divided = lines_divided_attempt
                    divide_max_try = divide_max_try - 1

                print("number_lines=%d  ; str_nb_lines=%d  ; divide_max_try=%d" % (number_lines, str_nb_lines, divide_max_try))
                number_lines = len(lines_divided)
                translation_result_phrase_array[i] = lines_divided
                for line_no in range (0, number_lines):
                    translation_result_using_separator[line_no+i] = lines_divided[line_no].rstrip().lstrip()
                    #if (line_no > 2):
                    #    if (translation_result_using_separator[line_no+i][:1] == ','):
                    #        translation_result_using_separator[line_no] = translation_result_using_separator[line_no] + ','
                    #        translation_result_using_separator[line_no+i] = translation_result_using_separator[line_no+i][1:].lstrip()
                number_lines = len(lines_divided)

                try:
                    print("Splitting phrase : %s (%d) = %d lines" % (to_text_by_phrase_separator_table[i], i, str_nb_lines + 0))
                except Exception:
                    try:
                        print("%s (%d): %d " % (to_text_by_phrase_separator_table[i].encode("utf-8"), i, str_nb_lines + 0))
                    except Exception:
                        print("(unable to print content to screen) (%d): %d : " % (i, str_nb_lines + 0))

                if number_lines != str_nb_lines:
                    print("Error in number of line %d, expected %d." % (number_lines, str_nb_lines))
                    #frequency = 2500  # Set Frequency To 2500 Hertz
                    #duration = 600  # Set Duration To 1000 ms == 1 second
                    #winsound.Beep(frequency, duration)
            except Exception:
                var = traceback.format_exc()
                print("  ERROR:%s<br>" % (var))


def print_html_program_result():
    if use_html :
        print("<table border=1 bgcolor=""#EEEEEE"">")

    for i, line in enumerate(from_text_table):
        Identical_with_without_separators = 'DIFFERENT<BR>'
        if to_text_by_phrase_separator_removed_table[i] == to_text_by_phrase_table[i]:
            Identical_with_without_separators = 'SAME<BR>'
        #print "<tr><td>%s<td>%s<td>%s<td>%s<td>%s%s" % (i, from_text_table[i], from_text_by_phrase_separator_table[i].encode('utf8'), to_text_by_phrase_separator_table[i].encode('utf8'), Identical_with_without_separators.encode('utf8'), to_text_by_phrase_separator_removed_table[i].encode('utf8') )
        if len(from_text_by_phrase_separator_table[i]) == 0:
            Identical_with_without_separators = ''
        if use_html :
            print("<tr><td>%d<td>'%s'<td>%s<td>%s<td>%s<td>%s%s" % (i, from_text_table[i], translation_result_using_separator[i].encode('utf8'), to_text_by_phrase_separator_table[i].encode('utf8'), to_text_by_phrase_table[i].encode('utf8'), Identical_with_without_separators.encode('utf8'), to_text_by_phrase_table[i].encode('utf8') ))
        #sys.exit(0)

    if use_html :
        print("</table><br>elapsedtime = ", elapsedtime)
        print("</span>")

def write_destination_language_in_docx_cell():
    if not splitonly:
        docxdoc.tables[0].cell(1, 2).text = dest_lang_name


def print_console_docx_file_translated():
    print("\nTranslated text:\n")
    numrows = len(table.rows)
    numcols = len(table.columns)
    current_cell_row = 2
    for row_n in range(2, (numrows)):

        str_translation_len = len(translation_result_using_separator[row_n])
        translation_phrase_lines_len = len(translation_result_phrase_array[row_n])
        if translation_phrase_lines_len == 0 and current_cell_row < row_n:
            print("%d :" % row_n)
        #print("row_n = %d" %  row_n)
        if translation_phrase_lines_len >= 1 :
            #print("%d : %s" % (row_n,' '.join(translation_result_phrase_array[row_n])))

            if not split_translation:
                translation_cell_text = to_text_by_phrase_separator_table[row_n]
                prepare_and_clear_cell_for_writing(row_n, translation_cell_text)
                if dest_lang in right_to_left_languages_list.keys():
                    translation_cell_aligned_text = reverse_string (translation_cell_text)
                else:
                    translation_cell_aligned_text = translation_cell_text
                print("%d : %s" % (row_n, translation_cell_aligned_text))
            else:
                #translation_cell_text = translation_result_using_separator[row_n]
                #print("len array: %d" % (translation_phrase_lines_len))
                #print("translation_result_phrase_array[%d] : %s" % (row_n,'\n'.join(translation_result_phrase_array[row_n])))

                translation_phrase_line_pos = 0
                translation_phrase_cell_pos = 0

                while translation_phrase_line_pos < translation_phrase_lines_len:
                    current_cell_row = row_n + translation_phrase_cell_pos
                    cell_lines_len = from_text_nb_lines_in_cell[row_n + translation_phrase_cell_pos]
                    cell_line_pos = 0
                    current_cell = table_cells[current_cell_row][2]
                    while cell_line_pos < cell_lines_len \
                        and translation_phrase_line_pos < translation_phrase_lines_len:

                        translation_phrase_line_str = translation_result_phrase_array[row_n][translation_phrase_line_pos]
                        if dest_lang in right_to_left_languages_list.keys():
                            translation_cell_aligned_text = reverse_string (translation_phrase_line_str)
                        else:
                            translation_cell_aligned_text = translation_phrase_line_str
                        if cell_lines_len > 1:
                            print("%d-%d : %s" % (current_cell_row, cell_line_pos + 1, translation_cell_aligned_text))
                        else:
                            print("%d : %s" % (current_cell_row, translation_cell_aligned_text))
                        if cell_line_pos == 0:
                            #print("cell_line_pos=%d" % cell_line_pos)
                            if splitonly:
                                prepare_and_clear_cell_for_writing(current_cell_row, translation_phrase_line_str)
                            else:
                            #prepare_and_clear_cell_for_writing(current_cell_row, translation_phrase_line_str)
                                cell_set_1st_paragraph(current_cell_row, translation_phrase_line_str)
                            # Not needed
                            #current_cell.paragraphs[0].text = translation_phrase_line_str
                        else:
                            # Add empty paragraph between translation lines
                            cell_add_paragraph(current_cell_row, "")
                            # Add the translation line
                            cell_add_paragraph(current_cell_row, translation_phrase_line_str)
                        cell_line_pos = cell_line_pos + 1
                        translation_phrase_line_pos = translation_phrase_line_pos + 1
                        #input("press enter")
                    translation_phrase_cell_pos = translation_phrase_cell_pos + 1

        try:
            if str_translation_len <= 0:
                str_phrase_stats = ""
            else:
                str_translation_len = len(translation_result_using_separator[row_n])
                str_nb_lines = from_text_nb_lines_in_phrase[row_n]
                if str_nb_lines > 0:
                    str_line_average = str_translation_len / str_nb_lines
                    str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                else:
                    str_line_average = 0
                    str_phrase_stats = ""

        except Exception:
            var = traceback.format_exc()
            print("  ERROR:%s<br>" % (var))


#print("Generating TMX file for translation comparison")
#generate_tmx_file ()
#word.Application.ActiveWindow.Close()
#word.Application.Quit()

def set_docx_properties_comment_for_history():
    now = datetime.datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    docxdoc.core_properties.comments = "Document translated by SMTV Robot version %s using %s engine on %s." % (PROGRAM_VERSION, translation_engine, dt_string)



def local_time_offset(t=None):
    """Return offset of local zone from GMT, either at present or at time t."""
    localtimezone = 0
    # python2.3 localtime() can't take None
    if t is None:
        t = time.time()
    localtimezone = -time.altzone / 3600
    if (localtimezone - int(localtimezone)) == 0:
        localtimezone = int(localtimezone)
    if time.localtime(t).tm_isdst == False or time.daylight != 1:
        localtimezone = -localtimezone
    return localtimezone


def run_statistics():
    global use_api
    global splitonly, driver
    global engine_method, end_time, elapsed_time, json_configuration_array
    
    statistics_html_statistics_form_url_key = ['statistics', 'html_statistics_form_url']
    statistics_html_statistics_form_url = get_nested_value_from_json_array(json_configuration_array, statistics_html_statistics_form_url_key)
    
    bool_print_stats = False
    
    try:
        if splitonly:
            action = "splitonly"
        else:
            action = "translate"
        
        docxfile_size = os.path.getsize(word_file_to_translate)
        if use_api == True:
            engine_method = "api"
        elif engine_method is None or engine_method == "":
            engine_method = "web"
        
        if xlsxreplacefile is not None:
            xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
            xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
            xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size-2], xlsxreplacefile_splitted[xlsxreplacefile_filename_size-1])
        else:
            xlsxreplacefile_name = ""
        
        if xlsxreplacefile_name != "":
            replacebeforelistsize = xtm.get_sheet_number_lines('before')
            replacebeforelistreplaced = xtm.get_sheet_number_of_replacements('before')
            replaceafterlistsize = xtm.get_sheet_number_lines('after')
            replaceafterlistreplaced = xtm.get_sheet_number_of_replacements('after')
            donotsplitlistsize = xtm.get_sheet_number_lines('keep_on_same_line')
            donotsplitfound = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
        else:
            replacebeforelistsize = ""
            replacebeforelistreplaced = ""
            replaceafterlistsize = ""
            replaceafterlistreplaced = ""
            donotsplitlistsize = ""
            donotsplitfound = ""
        
        platform_uname = platform.uname()
        platform_system = platform.system()
        platform_node = platform.node()
        platform_release = platform.release()
        platform_version = platform.version()
        platform_machine = platform.machine()
        platform_processor = platform.processor()
        
        cpu_count = psutil.cpu_count()
        mem_total = psutil.virtual_memory().total
        
        cost_google_translate = 20 * docxfile_table_number_of_characters / 1000000
        
        local_time_offset_str = local_time_offset()
        
        docxfile_page_count = None
        try:
            document = zipfile.ZipFile(word_file_to_translate)
            dxml = document.read('docProps/app.xml')
            uglyXml = xml.dom.minidom.parseString(dxml)
            docxfile_page_count = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
        except:
            if bool_print_stats:
                print("Unable to get number of pages from document. You can ignore this.")
        
        try:
            archive = zipfile.ZipFile("myDocxOrPptxFile.docx", "r")
            ms_data = archive.read("docProps/app.xml")
            archive.close()
            app_xml = ms_data.decode("utf-8")
            
            regex = r"<(Pages|Slides)>(\d)</(Pages|Slides)>"
            
            matches = re.findall(regex, app_xml, re.MULTILINE)
            match = matches[0] if matches[0:] else [0, 0]
            page_count = match[1]
        except:
            if bool_print_stats:
                print("Unable to get number of pages from document. You can ignore this.")
        
        if bool_print_stats:
            print("Statistics:")
            print("program_version: %s" % (PROGRAM_VERSION))
            #https://stackoverflow.com/questions/1695183/how-to-percent-encode-url-parameters-in-python
            
            print("docxfile: %s" % (word_file_to_translate))
            print("action: %s" % (action))
            print("destlang_code: %s" % (dest_lang))
            print("destlang_name: %s" % (dest_lang_name))
            print("docxfile: %s" % (docx_file_name))
            print("docxfile_page_count: %s" % docxfile_page_count)
            print("docxfile_size: %s" % (docxfile_size))
            print("docxfile_table_number_of_lines: %s" % (numrows))
            print("docxfile_table_number_of_phrases: %s" % (docxfile_table_number_of_phrases))
            print("docxfile_table_number_of_words: %s" % (docxfile_table_number_of_words))
            print("docxfile_table_number_of_characters: %s" % (docxfile_table_number_of_characters))
            print("engine: %s" % (translation_engine))
            print("xlsxreplacefile: %s" % (xlsxreplacefile_name))
            print("destfont: %s" % (dest_font))
            print("splitonly: %s" % (splitonly))
            print("split_translation: %s" % (split_translation))
            print("showbrowser: %s" % (showbrowser))
            print("start_time: %s" % (start_time))
            print("end_time: %s" % (end_time))
            print("elapsed_time: %s" % ((elapsed_time)))
            
            if xlsxreplacefile_name != "":
                print("replacebeforelistsize: %s" % (replacebeforelistsize))
                print("replacebeforelistreplaced: %s" % (replacebeforelistreplaced))
                print("replaceafterlistsize: %s" % (replaceafterlistsize))
                print("replaceafterlistreplaced: %s" % (replaceafterlistreplaced))
                print("donotsplitlistsize: %s" % (donotsplitlistsize))
                print("donotsplitfound: %s" % (donotsplitfound))
            
            print("str_uname : %s" % (str(platform_uname)))
            # print("platform_uname: %s" % (platform_uname))
            print("platform_system: %s" % (platform_system))
            print("platform_node: %s" % (platform_node))
            print("platform_release: %s" % (platform_release))
            print("platform_version: %s" % (platform_version))
            print("platform_machine: %s" % (platform_machine))
            print("platform_processor: %s" % (platform_processor))
            print("cpu_count: %s" % (cpu_count))
            print("mem_total: %s" % (mem_total))
            print("local_time_offset: %s" % (local_time_offset_str))
            print(f"cost_google_translate: {cost_google_translate:.2f}$")
            print("")
        
        #if use_api == False and not splitonly:
        chrome_options = Options()
        chrome_options.add_argument("--disable-web-security")
        chrome_options.add_argument("--disable-xss-auditor")
        chrome_options.add_argument("--log-level=3")  # fatal
        chrome_options.add_argument("--lang=en-GB")
        chrome_options.add_argument("--password-store=basic")
        
        if not showbrowser:
            chrome_options.add_argument("--headless")
        
        docxfile_table_number_of_lines = numrows
        if use_api or splitonly:
            print("\nCreating a new browser for stats")
            
                                                               
            driver = uc.Chrome(service=service, options=chrome_options)
            service = Service()                                
        
        query_params = {
            "program_version" : PROGRAM_VERSION,
            "engine" : translation_engine,
            "engine_method" : engine_method,
            "action" : action,
            "destlang_code" : dest_lang,
            "destlang_name" : dest_lang_name,
            "docxfile_size" : docxfile_size,
            "docxfile_table_number_of_lines" : docxfile_table_number_of_lines,
            "docxfile_table_number_of_phrases" : docxfile_table_number_of_phrases,
            "docxfile_table_number_of_words" : docxfile_table_number_of_words,
            "docxfile_table_number_of_characters" : docxfile_table_number_of_characters,
            "xlsxreplacefile" : xlsxreplacefile_name,
            "destfont" : dest_font,
            "split_translation" : split_translation,
            "showbrowser" : showbrowser,
            "start_time" : start_time,
            "end_time" : end_time,
            "elapsed_time" : elapsed_time,
            "replacebeforelistsize" : replacebeforelistsize,
            "replacebeforelistreplaced" : replacebeforelistreplaced,
            "replaceafterlistsize" : replaceafterlistsize,
            "replaceafterlistreplaced" : replaceafterlistreplaced,
            "replaceafterlistsize" : replaceafterlistsize,
            "replaceafterlistreplaced" : replaceafterlistreplaced,
            "donotsplitlistsize" : donotsplitlistsize,
            "donotsplitfound" : donotsplitfound,
            "platform_uname" : platform_uname,
            "platform_system" : platform_system,
            "platform_release" : platform_release,
            "platform_version" : platform_version,
            "platform_machine" : platform_machine,
            "platform_processor" : platform_processor,
            "cpu_count" : cpu_count,
            "platform_processor" : platform_processor,
            "mem_total" : mem_total,
            "elapsed_time" : elapsed_time,
            "local_time_offset" : local_time_offset_str,
            "docxfile_page_count" : docxfile_page_count,
            "platform_node" : platform_node,
            "docxfile" : docx_file_name
        }
        
        base_url = statistics_html_statistics_form_url
        encoded_params = urlencode(query_params, quote_via=quote_plus)
        url = f"{base_url}?{encoded_params}"
        
        driver.get(url)
        
        #time.sleep(20)
        
        submit_stats_element = "//input[@value='Submit']"
        try:
            submit_stats_button = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.XPATH, submit_stats_element)))
            submit_stats_button.submit()
            #time.sleep(1)
            submited_div_element = "//div[@id='form_post_submitted']"
            submited_div = WebDriverWait(driver, 4).until(EC.presence_of_element_located((By.XPATH, submited_div_element)))
            #print("statistics updated")
        except:
            print("Warning failed to update stats, you can ignore this.")
            #pass

    except:
        #var = traceback.format_exc()
        #print(var)
        print("Warning failed to update stats, you can ignore this...")
    
    #time.sleep(10)


def browser_fill_form_field_value(field_css_id, field_value):
    try:
        input_field = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.CSS_SELECTOR, field_css_id)))
        input_field.send_keys (field_value)
    except:
        var = traceback.format_exc()
        print(var)


def get_robot_usage_comment():
    global use_api
    global splitonly, driver
    global engine_method, end_time, elapsed_time, json_configuration_array, str_needs_update
    
    javascript_json_version_checker_url_key = ['version_checker', 'javascript_json_version_checker_url']
    javascript_json_version_checker_url = get_nested_value_from_json_array(json_configuration_array, javascript_json_version_checker_url_key)
        
    if use_api == True:
        engine_method = "api"
    elif engine_method is None or engine_method == "":
        engine_method = "web"
    
    if xlsxreplacefile is not None:
        xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
        xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
        xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size-2], xlsxreplacefile_splitted[xlsxreplacefile_filename_size-1])
    else:
        xlsxreplacefile_name = ""
    
    if xlsxreplacefile_name != "":
        replacebeforelistsize = xtm.get_sheet_number_lines('before')
        replacebeforelistreplaced = xtm.get_sheet_number_of_replacements('before')
        replaceafterlistsize = xtm.get_sheet_number_lines('after')
        replaceafterlistreplaced = xtm.get_sheet_number_of_replacements('after')
        donotsplitlistsize = xtm.get_sheet_number_lines('keep_on_same_line')
        donotsplitfound = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
    else:
        replacebeforelistsize = ""
        replacebeforelistreplaced = ""
        replaceafterlistsize = ""
        replaceafterlistreplaced = ""
        donotsplitlistsize = ""
        donotsplitfound = ""
    

    try:
        driver.get(javascript_json_version_checker_url)
        bool_print_stats = False

        json_obj = json.loads("{}")

        json_obj["program_version"] = PROGRAM_VERSION
        json_obj["docxfile"] = word_file_to_translate

        if splitonly:
            json_obj['action'] = "splitonly"
        else:
            json_obj['action'] = "translate"

        json_obj["destlang_code"] = dest_lang
        json_obj["destlang_name"] = dest_lang_name
        json_obj["docxfile"] = "%s%s" % (docx_file_name,"'")
        json_obj["docxfile_table_number_of_lines"] = numrows
        json_obj["docxfile_table_number_of_phrases"] = docxfile_table_number_of_phrases
        json_obj["docxfile_table_number_of_words"] = docxfile_table_number_of_words
        json_obj["docxfile_table_number_of_characters"] = docxfile_table_number_of_characters
        json_obj["engine"] = translation_engine
        json_obj["engine_method"] = engine_method
        json_obj["xlsxreplacefile"] = xlsxreplacefile_name
        if dest_font is not None:
            json_obj["destfont"] = "%s" % dest_font
        json_obj["splitonly"] = splitonly
        json_obj["split_translation"] = split_translation
        json_obj["showbrowser"] = showbrowser
        json_obj["start_time"] = "%s" % start_time
        json_obj["end_time"] = "%s" % end_time
        json_obj["elapsed_time"] = "%s" % elapsed_time

        try:
            docxfile_size = os.path.getsize(word_file_to_translate)
            json_obj["docxfile_size"] = docxfile_size
            if use_api == True:
                json_obj['engine_method'] = "api"
            elif engine_method is None or engine_method == "":
                json_obj['engine_method'] = "web"

            if xlsxreplacefile is not None:
                xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
                xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
                xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size - 2],
                                                 xlsxreplacefile_splitted[xlsxreplacefile_filename_size - 1])
                json_obj['xlsxreplacefile'] = xlsxreplacefile_name
            else:
                json_obj['xlsxreplacefile'] = ""
                json_obj['xlsxreplacefile_filename_size'] = ""

            if xlsxreplacefile_name != "":
                json_obj['replacebeforelistsize'] = xtm.get_sheet_number_lines('before')
                json_obj['replacebeforelistreplaced'] = xtm.get_sheet_number_of_replacements('before')
                json_obj['replaceafterlistsize'] = xtm.get_sheet_number_lines('after')
                json_obj['replaceafterlistreplaced'] = xtm.get_sheet_number_of_replacements('after')
                json_obj['donotsplitlistsize'] = xtm.get_sheet_number_lines('keep_on_same_line')
                json_obj['donotsplitfound'] = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
            else:
                json_obj['replacebeforelistsize'] = ""
                json_obj['replacebeforelistreplaced'] = ""
                json_obj['replaceafterlistsize'] = ""
                json_obj['replaceafterlistreplaced'] = ""
                json_obj['donotsplitlistsize'] = ""
                json_obj['donotsplitfound'] = ""

            json_obj['platform_uname'] = platform.uname()
            json_obj['platform_system'] =  platform.system()
            json_obj['platform_node'] = platform.node()
            json_obj['platform_release'] = platform.release()
            json_obj['platform_version'] = platform.version()
            json_obj['platform_machine'] = platform.machine()
            json_obj['platform_processor'] = platform.processor()

            json_obj['cpu_count'] = psutil.cpu_count()
            json_obj['mem_total'] = psutil.virtual_memory().total

            cost_google_translate = 20 * docxfile_table_number_of_characters / 1000000

            local_time_offset_str = local_time_offset()

            docxfile_page_count = None
            try:
                document = zipfile.ZipFile(word_file_to_translate)
                dxml = document.read('docProps/app.xml')
                uglyXml = xml.dom.minidom.parseString(dxml)
                docxfile_page_count = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
                json_obj['docxfile_page_count'] = docxfile_page_count
            except:
                if bool_print_stats:
                    json_obj['docxfile_page_count'] = ""

            try:
                archive = zipfile.ZipFile("myDocxOrPptxFile.docx", "r")
                ms_data = archive.read("docProps/app.xml")
                archive.close()
                app_xml = ms_data.decode("utf-8")

                regex = r"<(Pages|Slides)>(\d)</(Pages|Slides)>"

                matches = re.findall(regex, app_xml, re.MULTILINE)
                match = matches[0] if matches[0:] else [0, 0]
                page_count = match[1]
                json_obj['docxfile_page_count'] = page_count
            except:
                if bool_print_stats:
                    print("Unable to get number of pages from document. You can ignore this.")

            #print(json.dumps(json_obj, indent=4))
            print("\n-------------------------------")
            print("Checking for program updates...")
            print("-------------------------------\n")

            element_json_robot = WebDriverWait(driver, 1).until(
                    EC.presence_of_element_located((By.ID, "json_robot")))
            driver.execute_script("arguments[0].innerText = arguments[1]", element_json_robot, json.dumps(json_obj))

            element_submit = WebDriverWait(driver, 1).until(
                    EC.presence_of_element_located((By.ID, "submit")))
            element_submit.click()

            html_translation = driver.page_source
            # soup = BeautifulSoup(html_translation)
            soup = BeautifulSoup(html_translation, features="lxml")
            soup_div_text = soup.find('div', id='message_text')
            available_updates_message = ''.join(map(str, soup_div_text.text))
            
            try:
                soup_div_needs_update = soup.find('div', id='needs_update')
                str_needs_update = ''.join(map(str, soup_div_needs_update.text))
            except:
                pass
                
            if available_updates_message != "":
                print (''.join(map(str, soup_div_text.text)))
                print("\n-------------------------------")

            return 0;
            try:
                print(driver.capabilities['browserVersion'])
            except:
                pass
            print(driver.name)

            if bool_print_stats:
                print("Statistics:")
                print("program_version: %s" % (PROGRAM_VERSION))
                # https://stackoverflow.com/questions/1695183/how-to-percent-encode-url-parameters-in-python

                print("docxfile: %s" % (word_file_to_translate))
                print("action: %s" % (action))
                print("destlang_code: %s" % (dest_lang))
                print("destlang_name: %s" % (dest_lang_name))
                print("docxfile: %s" % (docx_file_name))
                print("docxfile_page_count: %s" % docxfile_page_count)
                print("docxfile_size: %s" % (docxfile_size))
                print("docxfile_table_number_of_lines: %s" % (numrows))
                print("docxfile_table_number_of_phrases: %s" % (docxfile_table_number_of_phrases))
                print("docxfile_table_number_of_words: %s" % (docxfile_table_number_of_words))
                print("docxfile_table_number_of_characters: %s" % (docxfile_table_number_of_characters))
                print("engine: %s" % (translation_engine))
                print("xlsxreplacefile: %s" % (xlsxreplacefile_name))
                print("destfont: %s" % (dest_font))
                print("splitonly: %s" % (splitonly))
                print("split_translation: %s" % (split_translation))
                print("showbrowser: %s" % (showbrowser))
                print("start_time: %s" % (start_time))
                print("end_time: %s" % (end_time))
                print("elapsed_time: %s" % ((elapsed_time)))

                if xlsxreplacefile_name != "":
                    print("replacebeforelistsize: %s" % (replacebeforelistsize))
                    print("replacebeforelistreplaced: %s" % (replacebeforelistreplaced))
                    print("replaceafterlistsize: %s" % (replaceafterlistsize))
                    print("replaceafterlistreplaced: %s" % (replaceafterlistreplaced))
                    print("donotsplitlistsize: %s" % (donotsplitlistsize))
                    print("donotsplitfound: %s" % (donotsplitfound))

                print("str_uname : %s" % (str(platform_uname)))
                # print("platform_uname: %s" % (platform_uname))
                print("platform_system: %s" % (platform_system))
                print("platform_node: %s" % (platform_node))
                print("platform_release: %s" % (platform_release))
                print("platform_version: %s" % (platform_version))
                print("platform_machine: %s" % (platform_machine))
                print("platform_processor: %s" % (platform_processor))
                print("cpu_count: %s" % (cpu_count))
                print("mem_total: %s" % (mem_total))
                print("local_time_offset: %s" % (local_time_offset_str))
                print(f"cost_google_translate: {cost_google_translate:.2f}$")
                print("")

            # if use_api == False and not splitonly:
            chrome_options = Options()
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--disable-xss-auditor")
            chrome_options.add_argument("--log-level=3")  # fatal
            chrome_options.add_argument("--lang=en-GB")
            chrome_options.add_argument("--password-store=basic")

            driver.get("https://forms.gle/YeYYUYY5SNo6MKkB8")
            browser_fill_form_field_value(".Qr7Oae:nth-child(1) .whsOnd", "REMOTE_ADDR")
            browser_fill_form_field_value(".Qr7Oae:nth-child(2) .whsOnd", "country_name")
            browser_fill_form_field_value(".Qr7Oae:nth-child(3) .whsOnd", "remote_location_text")
            browser_fill_form_field_value(".Qr7Oae:nth-child(4) .whsOnd", "HTTP_USER_AGENT")
            browser_fill_form_field_value(".Qr7Oae:nth-child(5) .whsOnd", "program_version")
            browser_fill_form_field_value(".Qr7Oae:nth-child(6) .whsOnd", "docxfile")
            browser_fill_form_field_value(".Qr7Oae:nth-child(7) .whsOnd", "docxfile_page_count")
            browser_fill_form_field_value(".Qr7Oae:nth-child(8) .whsOnd", "docxfile_size")
            browser_fill_form_field_value(".Qr7Oae:nth-child(9) .whsOnd", "docxfile_table_number_of_lines")
            browser_fill_form_field_value(".Qr7Oae:nth-child(10) .whsOnd", "docxfile_table_number_of_words")
            browser_fill_form_field_value(".Qr7Oae:nth-child(11) .whsOnd", "docxfile_table_number_of_characters")
            browser_fill_form_field_value(".Qr7Oae:nth-child(12) .whsOnd", "action")
            browser_fill_form_field_value(".Qr7Oae:nth-child(13) .whsOnd", "destlang_code")
            browser_fill_form_field_value(".Qr7Oae:nth-child(14) .whsOnd", "destlang_name")
            browser_fill_form_field_value(".Qr7Oae:nth-child(15) .whsOnd", "engine")
            browser_fill_form_field_value(".Qr7Oae:nth-child(16) .whsOnd", "engine_method")
            browser_fill_form_field_value(".Qr7Oae:nth-child(17) .whsOnd", "xlsxreplacefile")
            browser_fill_form_field_value(".Qr7Oae:nth-child(18) .whsOnd", "destfont")
            browser_fill_form_field_value(".Qr7Oae:nth-child(19) .whsOnd", "split_translation")
            browser_fill_form_field_value(".Qr7Oae:nth-child(20) .whsOnd", "splitonly")
            browser_fill_form_field_value(".Qr7Oae:nth-child(21) .whsOnd", "showbrowser")
            browser_fill_form_field_value(".Qr7Oae:nth-child(22) .whsOnd", "server_time")
            browser_fill_form_field_value(".Qr7Oae:nth-child(23) .whsOnd", "start_time")
            browser_fill_form_field_value(".Qr7Oae:nth-child(24) .whsOnd", "end_time")
            browser_fill_form_field_value(".Qr7Oae:nth-child(25) .whsOnd", "elapsed_time")
            browser_fill_form_field_value(".Qr7Oae:nth-child(26) .whsOnd", "replacebeforelistsize")
            browser_fill_form_field_value(".Qr7Oae:nth-child(27) .whsOnd", "replacebeforelistreplaced")
            browser_fill_form_field_value(".Qr7Oae:nth-child(28) .whsOnd", "replaceafterlistsize")
            browser_fill_form_field_value(".Qr7Oae:nth-child(29) .whsOnd", "replaceafterlistreplaced")
            browser_fill_form_field_value(".Qr7Oae:nth-child(30) .whsOnd", "donotsplitlistsize")
            browser_fill_form_field_value(".Qr7Oae:nth-child(31) .whsOnd", "donotsplitfound")
            browser_fill_form_field_value(".Qr7Oae:nth-child(32) .whsOnd", "platform_uname")
            browser_fill_form_field_value(".Qr7Oae:nth-child(33) .whsOnd", "platform_system")
            browser_fill_form_field_value(".Qr7Oae:nth-child(34) .whsOnd", "platform_node")
            browser_fill_form_field_value(".Qr7Oae:nth-child(35) .whsOnd", "platform_release")
            browser_fill_form_field_value(".Qr7Oae:nth-child(36) .whsOnd", "platform_version")
            browser_fill_form_field_value(".Qr7Oae:nth-child(37) .whsOnd", "platform_machine")
            browser_fill_form_field_value(".Qr7Oae:nth-child(38) .whsOnd", "platform_processor")
            browser_fill_form_field_value(".Qr7Oae:nth-child(39) .whsOnd", "cpu_count")
            browser_fill_form_field_value(".Qr7Oae:nth-child(40) .whsOnd", "mem_total")
            browser_fill_form_field_value(".Qr7Oae:nth-child(41) .whsOnd", "local_time_offset")
            browser_fill_form_field_value(".Qr7Oae:nth-child(42) .whsOnd", "docxfile_table_number_of_phrases")

            if not showbrowser:
                chrome_options.add_argument("--headless")

            docxfile_table_number_of_lines = numrows
            if use_api or splitonly:
                print("\nCreating a new browser for stats")
                
                service = Service()                                
                driver = uc.Chrome(service=service, options=chrome_options)

            query_params = {
                "program_version": PROGRAM_VERSION,
                "engine": translation_engine,
                "engine_method": engine_method,
                "action": action,
                "destlang_code": dest_lang,
                "destlang_name": dest_lang_name,
                "docxfile_size": docxfile_size,
                "docxfile_table_number_of_lines": docxfile_table_number_of_lines,
                "docxfile_table_number_of_phrases": docxfile_table_number_of_phrases,
                "docxfile_table_number_of_words": docxfile_table_number_of_words,
                "docxfile_table_number_of_characters": docxfile_table_number_of_characters,
                "xlsxreplacefile": xlsxreplacefile_name,
                "destfont": dest_font,
                "split_translation": split_translation,
                "showbrowser": showbrowser,
                "start_time": start_time,
                "end_time": end_time,
                "elapsed_time": elapsed_time,
                "replacebeforelistsize": replacebeforelistsize,
                "replacebeforelistreplaced": replacebeforelistreplaced,
                "replaceafterlistsize": replaceafterlistsize,
                "replaceafterlistreplaced": replaceafterlistreplaced,
                "replaceafterlistsize": replaceafterlistsize,
                "replaceafterlistreplaced": replaceafterlistreplaced,
                "donotsplitlistsize": donotsplitlistsize,
                "donotsplitfound": donotsplitfound,
                "platform_uname": platform_uname,
                "platform_system": platform_system,
                "platform_release": platform_release,
                "platform_version": platform_version,
                "platform_machine": platform_machine,
                "platform_processor": platform_processor,
                "cpu_count": cpu_count,
                "platform_processor": platform_processor,
                "mem_total": mem_total,
                "elapsed_time": elapsed_time,
                "local_time_offset": local_time_offset_str,
                "docxfile_page_count": docxfile_page_count,
                "platform_node": platform_node,
                "docxfile": docx_file_name
            }

            base_url = javascript_json_version_checker_url
            encoded_params = urlencode(query_params, quote_via=quote_plus)
            url = f"{base_url}?{encoded_params}"

            driver.get(url)
            # time.sleep(20)

            submit_stats_element = "//input[@value='Submit']"
            try:
                submit_stats_button = WebDriverWait(driver, 1).until(
                    EC.presence_of_element_located((By.XPATH, submit_stats_element)))
                submit_stats_button.submit()
                # time.sleep(1)
                submited_div_element = "//div[@id='form_post_submitted']"
                submited_div = WebDriverWait(driver, 5).until(
                    EC.presence_of_element_located((By.XPATH, submited_div_element)))
                print("statistics updated")
            except:
                #var = traceback.format_exc()
                #print(var)
                print("Warning failed to get available updates status, you can ignore this.")
                # pass

        except:
            #var = traceback.format_exc()
            #print(var)
            print("Warning failed to get available updates status, you can ignore this.")

        # time.sleep(10)

    except:
        var = traceback.format_exc()
        #print(var)
        print("Warning failed to get available updates status, you can ignore this.")


# Open the default app for the docx file
def open_app_docx_file():
    global word_file_to_translate_save_as_path
    
    try:
        if platform.system() == 'Windows':
            subprocess.Popen(["start", "", word_file_to_translate_save_as_path], shell=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", word_file_to_translate_save_as_path])
        elif platform.system() == "Linux":  # Linux
            subprocess.Popen(["xdg-open", word_file_to_translate_save_as_path])
        else:
            print("Unsupported operating system.")
            
    except Exception as e:
        print("Error:", e)
        print("Warning, unable to open file %s." % (word_file_to_translate_save_as_path))
def save_docx_file():
    global docxdoc, word_file_to_translate, word_file_to_translate_save_as_path
    
    lang_name = ""
    lang_code = dest_lang
    
    # Find valid two letter code (Norwegian is invalid nb, but should be no)
    try:
        lang_name = google_translate_lang_codes[lang_code]
    except:
        try:
            lang_name = deepl_translate_lang_codes[lang_code]
            for google_lang_code in google_translate_lang_codes.keys():
                try:
                    if deepl_translate_lang_codes[lang_code].lower() == google_translate_lang_codes[google_lang_code].lower() and lang_code != google_lang_code:
                        lang_code = google_lang_code
                except:
                    pass
        except:
            pass
    
    language_alpha_extension = None
    lang_alpha3b_code = None
        
    try:
        lang_alpha3_code = Language.get(lang_code).to_alpha3()
        lang_alpha3b_code = Language.get(lang_code).to_alpha3(variant='B')
        pass
    except:
        lang_alpha3b_code = None

    word_file_to_translate_save_as_path = word_file_to_translate
    if lang_alpha3b_code is not None:
        find_alpha3_code_suffix = f"(?i)_{lang_alpha3b_code}.docx$"
        if not re.search(find_alpha3_code_suffix, word_file_to_translate):
            word_file_to_translate_save_as_path = re.sub("(?i)_{lang_alpha3b_code}.docx$", f".docx", word_file_to_translate)
            lang_alpha3b_code = lang_alpha3b_code.upper()
            word_file_to_translate_save_as_path = re.sub("(?i).docx$", f"_{lang_alpha3b_code}.docx", word_file_to_translate)
            print(f"\nAdding file name suffix _{lang_alpha3b_code}.")

    local_time_offset()

    file_saved = 0
    while file_saved == 0:
        try:
            docxdoc.save(word_file_to_translate_save_as_path)
            file_saved = 1
        except Exception:
            var = traceback.format_exc()
            txt_readline = input(
                "\n\nERROR: File saving failed. Please close microsoft word or other program and press enter to save the translated document.\n")

def cleanup_selenium_chrome_temp_folders():
    root_path = r"C:\\Program Files"

    delete_patterns = [
        r"scoped_dir\d{3,}_\d{6,}",
        r"chrome_BITS_\d{3,}_\d{6,}",
        r"chrome_PuffinComponentUnpacker_BeginUnzipping\d{3,}_\d{7,}",
        r"chrome_url_fetcher_\d{3,}_\d{7,}"
    ]

    if platform.system().lower() == 'windows':
        print("\nCleaning selenium chrome temporary folders")    
    
        # 24 hours ago
        cutoff_time = time.time() - 24 * 60 * 60

        # Get only immediate subfolders in root_path
        folders = [f for f in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, f))]

        def is_folder_inactive(folder_path):
            for dirpath, _, filenames in os.walk(folder_path):
                for filename in filenames:
                    filepath = os.path.join(dirpath, filename)
                    try:
                        if os.path.getmtime(filepath) > cutoff_time:
                            return False  # File modified recently
                    except Exception:
                        continue  # Skip inaccessible files
            return True

        for folder in folders:
            folder_path = os.path.join(root_path, folder)

            if any(re.fullmatch(pattern, folder) for pattern in delete_patterns):
                if is_folder_inactive(folder_path):
                    print(f"⚠️ DELETING (unsused >24h): {folder_path}")
                    try:
                        shutil.rmtree(folder_path, ignore_errors=True)
                    except Exception as e:
                        print(f"Error deleting {folder_path}: {e}")
                else:
                    print(f"⏳ ACTIVE: Skipping recently modified folder {folder_path}")
            else:
                #print(f"✅ SAFE: Skipping non-matching folder {folder_path}")
                pass

def main() -> int:
    global E_mail_str, end_time, elapsed_time, translation_engine, engine_method, tried_login_in_deepl, viewdocx, word_file_to_translate_save_as_path
    global logged_into_deepl, deepl_nb_clear_cached_times, version_checker_sleep_seconds_on_update
    translation_succeded = False

    set_translation_function()
    initialize_translation_memory_xlsx()

    read_and_parse_docx_document()

    create_webdriver()
    
    if translation_engine == 'deepl':
        logged_into_deepl = selenium_chrome_deepl_log_in()
        
    if translation_engine == 'perplexity':
        logged_into_perplexity = selenium_chrome_perplexity_wait_log_in
        if not logged_into_perplexity:
            print("Failed to login into perplexity")
            exit(100)

    translation_succeded = translate_docx()
    
    if logged_into_deepl:
        selenium_chrome_deepl_log_off()

    if translation_succeded == False and translation_engine == 'deepl' and engine_method == 'phrasesblock':
        engine_method = 'singlephrase'
        set_translation_function()
        try:
            driver.close()
            driver.quit()
        except:
            pass
        create_webdriver()

    get_translation_and_replace_after()

    minimize_browser()

    document_split_phrases()

    write_destination_language_in_docx_cell()

    print_console_docx_file_translated()
    set_docx_properties_comment_for_history()

    end_time = datetime.datetime.now()

    elapsed_time = end_time - start_time

    run_statistics()
    save_docx_file()
    
    if viewdocx:
        print(f"Opening document : {word_file_to_translate_save_as_path}")
        open_app_docx_file()
    end_time = datetime.datetime.now()

    elapsed_time = end_time - start_time

    if xlsxreplacefile is not None:
        xtm.print_replaced_items_number_of_replacements('before')
        xtm.print_replaced_items_number_of_replacements('after')
        xtm.print_do_not_split_number_of_matches('keep_on_same_line')

    if driver is not None:
        clean_up_previous_chrome_selenium_drivers(driver.service.path)
        
    print("\nTranslation ended, file saved. Elasped time: %s (h:mm:ss.mmm)" % (elapsed_time))
    print("\nSaved file name: %s" % (word_file_to_translate_save_as_path))
    
    cleanup_selenium_chrome_temp_folders()
    
    get_robot_usage_comment()

    if translation_engine == 'perplexity':
        if engine_method == 'api':
            print(f"Total cost: {total_cost}")

    try:
        #driver.maximize_window()
        print("\nClosing chrome browser...")
        
        driver_before_close_time = datetime.datetime.now()
        driver.close()
        driver_after_close_time = datetime.datetime.now()
        driver.quit()
        
        driver_after_quit_time = datetime.datetime.now()

        driver_close_time = driver_after_close_time - driver_before_close_time
        driver_quit_time = driver_after_quit_time - driver_after_close_time
        driver_close_quit_time = driver_after_quit_time - driver_before_close_time
        
        #print("\nDriver close time: %s (h:mm:ss.mmm)" % (driver_close_time))
        #print("\nDriver quit time: %s (h:mm:ss.mmm)" % (driver_quit_time))
        #print("\nDriver close and quit time: %s (h:mm:ss.mmm)" % (driver_close_quit_time))
    except:
        var = traceback.format_exc()
        print(var)
        pass

    if dest_lang_name is None or dest_lang_name == "":
        if not splitonly:
            print("\n*********************************************************************************")
            print("WARNING: Target language name for %s not found. Translation may have have failed." % (dest_lang))
            print("*********************************************************************************\n")

    print("\nDeveloper: %s" % (E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not exitonsuccess and not silent:
        input("Enter to close program")
    else:
        if str_needs_update == "1":
            print(f"Please download and install the program update (message will be shown for {version_checker_sleep_seconds_on_update} seconds).")
            time.sleep(version_checker_sleep_seconds_on_update)
        print("Program ended")
    
    # Suppress any error message from undetected_chromedriver cleanup
    devnull = open(os.devnull, 'w')
    sys.stderr = devnull
    sys.__stderr__ = devnull
    
    return 0

if __name__ == '__main__':
    main()  # next section explains the use of sys.exit
    # Redirect all stderr output to null (silences destructor error messages)
    sys.exit(0)