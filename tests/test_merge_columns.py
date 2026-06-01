"""Unit tests for merge_columns — the pure-XML 4th-column merge behind google+deepl.
No Selenium, no globals: build synthetic 3-column docx in memory, merge, assert shape.
Run: pip install python-docx pytest ; python -m pytest tests/test_merge_columns.py
"""
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))
from merge_columns import merge_second_engine_column  # noqa: E402


def _make_docx(path, rows, rtl_last=False):
    doc = Document()
    t = doc.add_table(rows=len(rows), cols=3)
    for ri, (c0, c1, c2) in enumerate(rows):
        cells = t.rows[ri].cells
        cells[0].text, cells[1].text, cells[2].text = c0, c1, c2
    if rtl_last:
        tc = t.rows[-1].cells[2]._tc
        p = tc.findall(qn("w:p"))[0]
        p.get_or_add_pPr().append(OxmlElement("w:bidi"))
    doc.save(str(path))
    return path


def _rows(path):
    t = Document(str(path)).tables[0]
    return [[c.text for c in r.cells] for r in t.rows]


def test_merge_appends_fourth_column(tmp_path):
    g = _make_docx(tmp_path / "g.docx", [("", "", "Persian"), ("", "Hello", "G-x"), ("", "Bye", "G-y")])
    d = _make_docx(tmp_path / "d.docx", [("", "", "Persian"), ("", "Hello", "D-x"), ("", "Bye", "D-y")])
    out = merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"), landscape=False)
    t = Document(str(out)).tables[0]
    assert len(t.columns) == 4
    rows = _rows(out)
    assert rows[1][2] == "G-x" and rows[1][3] == "D-x"
    assert rows[2][2] == "G-y" and rows[2][3] == "D-y"


def test_engine_labels_below_language(tmp_path):
    # The language name sits at cell (1, 2); each engine's name goes in the EMPTY cell
    # directly BELOW it (row 2) -- "Google" under col 2, "DeepL" under the new last col.
    base = [("", "", ""), ("", "", "Persian"), ("", "", ""), ("", "Hello", "X")]
    g = _make_docx(tmp_path / "g.docx", base)
    d = _make_docx(tmp_path / "d.docx", base)
    out = merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"),
                                     primary_label="Google", secondary_label="DeepL", landscape=False)
    rows = _rows(out)
    assert rows[1][2] == "Persian"
    assert rows[2][2] == "Google"
    assert rows[2][3] == "DeepL"
    assert rows[3][1] == "Hello"


def test_landscape_default_on_and_toggle_off(tmp_path):
    from docx.enum.section import WD_ORIENT
    rows = [("", "", "Persian"), ("", "Hi", "G1")]
    g = _make_docx(tmp_path / "g.docx", rows)
    d = _make_docx(tmp_path / "d.docx", rows)
    on = merge_second_engine_column(str(g), str(d), str(tmp_path / "on.docx"))
    sec = Document(str(on)).sections[0]
    assert sec.orientation == WD_ORIENT.LANDSCAPE
    assert sec.page_width > sec.page_height
    g2 = _make_docx(tmp_path / "g2.docx", rows)
    d2 = _make_docx(tmp_path / "d2.docx", rows)
    off = merge_second_engine_column(str(g2), str(d2), str(tmp_path / "off.docx"), landscape=False)
    assert Document(str(off)).sections[0].page_width < Document(str(off)).sections[0].page_height


def test_source_columns_preserved(tmp_path):
    g = _make_docx(tmp_path / "g.docx", [("X", "src-a", "G1"), ("Y", "src-b", "G2")])
    d = _make_docx(tmp_path / "d.docx", [("X", "src-a", "D1"), ("Y", "src-b", "D2")])
    out = merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"), primary_label=None, secondary_label=None, landscape=False)
    rows = _rows(out)
    assert [r[0] for r in rows] == ["X", "Y"]
    assert [r[1] for r in rows] == ["src-a", "src-b"]
    assert [r[2] for r in rows] == ["G1", "G2"]


def test_row_count_mismatch_raises(tmp_path):
    g = _make_docx(tmp_path / "g.docx", [("", "a", "G1"), ("", "b", "G2")])
    d = _make_docx(tmp_path / "d.docx", [("", "a", "D1")])
    with pytest.raises(ValueError):
        merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"))


def test_picks_largest_table_not_first(tmp_path):
    def _two_table(path, rows):
        doc = Document()
        head = doc.add_table(rows=1, cols=3)
        head.rows[0].cells[0].text = "meta"
        t = doc.add_table(rows=len(rows), cols=3)
        for ri, (c0, c1, c2) in enumerate(rows):
            cells = t.rows[ri].cells
            cells[0].text, cells[1].text, cells[2].text = c0, c1, c2
        doc.save(str(path))
        return path

    g = _two_table(tmp_path / "g.docx", [("", "", "Persian"), ("", "Hi", "G1"), ("", "Bye", "G2")])
    d = _two_table(tmp_path / "d.docx", [("", "", "Persian"), ("", "Hi", "D1"), ("", "Bye", "D2")])
    out = merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"), primary_label=None, secondary_label=None, landscape=False)
    content = max(Document(str(out)).tables, key=lambda t: len(t.rows))
    rows = [[c.text for c in r.cells] for r in content.rows]
    assert len(content.columns) == 4
    assert rows[1][2] == "G1" and rows[1][3] == "D1"


def test_rtl_bidi_survives_into_new_column(tmp_path):
    g = _make_docx(tmp_path / "g.docx", [("", "a", "G1")])
    d = _make_docx(tmp_path / "d.docx", [("", "a", "D1")], rtl_last=True)
    out = merge_second_engine_column(str(g), str(d), str(tmp_path / "out.docx"), primary_label=None, secondary_label=None, landscape=False)
    t = Document(str(out)).tables[0]
    new_cell_tc = t.rows[0]._tr.findall(qn("w:tc"))[3]
    p = new_cell_tc.findall(qn("w:p"))[0]
    pPr = p.find(qn("w:pPr"))
    assert pPr is not None and pPr.find(qn("w:bidi")) is not None
